from __future__ import annotations

from datetime import datetime, time, timedelta
import hashlib
from io import BytesIO
import json
from pathlib import Path
import re
import subprocess
import unicodedata
import zipfile

import matplotlib.pyplot as plt
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageChops
from streamlit.errors import StreamlitSecretNotFoundError

from app_auth import (
    ROLE_DEFINITIONS,
    clear_authenticated_user,
    describe_access_window,
    get_authenticated_user,
    set_authenticated_user,
    user_has_permission,
)
from app_storage import build_storage_backend


APP_TITLE = "Procesador TEC de Peajes"
APP_NAV_KEY = "app_selected_page"
TEC_RESULT_STATE_KEY = "tec_last_processing_result"
PROCESSING_SIGNATURE_VERSION = "2026-03-20-v3"
CONTRACTOR_LOGO_PATH = Path(__file__).parent / "ChatGPT Image 18 mar 2026, 03_37_00 a.m..png"
INFORME_TEMPLATE_CANDIDATES = (
    Path(__file__).parent / "templates" / "Informe TEC NORVIAL - 2022.docx",
    Path(r"C:\Users\chrys\OneDrive\Dic Virtual D\work\2026\CIDATT\TEC\Archivos Originales\Informe TEC NORVIAL - 2022.docx"),
)
CLIENT_LOGO_DIR_CANDIDATES = (
    Path(__file__).parent / "templates" / "client_logos",
    Path(__file__).parent / "client_logos",
    Path(__file__).parent / "logos",
    Path(__file__).parent,
)

MODULE_CATALOG = [
    {
        "page": "TEC",
        "title": "TEC",
        "eyebrow": "Modulo operativo",
        "description": "Procesamiento de bases de peaje, limpieza de placas, recuperacion de tiempos y generacion de entregables.",
        "status": "Activo",
        "accent": "#0f3d91",
        "icon_svg": '<path d="M4 19h16" /><path d="M7 16V9" /><path d="M12 16V5" /><path d="M17 16v-3" />',
    },
    {
        "page": "Relevamientos",
        "title": "Relevamientos",
        "eyebrow": "Modulo planificado",
        "description": "Registro estructurado de levantamiento de campo, evidencias, hallazgos y seguimiento por estacion o activo.",
        "status": "Proximamente",
        "accent": "#1849a9",
        "icon_svg": '<path d="M9 11l3 3L22 4" /><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11" />',
    },
    {
        "page": "Auditorias",
        "title": "Auditorias",
        "eyebrow": "Modulo planificado",
        "description": "Control de auditorias operativas, observaciones, estados de cumplimiento y trazabilidad de acciones correctivas.",
        "status": "Proximamente",
        "accent": "#245cc6",
        "icon_svg": '<path d="M9 3h6l5 5v10a2 2 0 0 1-2 2H6a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h3" /><path d="M9 3v5h5" /><path d="M9 14h6" /><path d="M9 18h6" />',
    },
    {
        "page": "Satisfaccion",
        "title": "Satisfaccion",
        "eyebrow": "Modulo planificado",
        "description": "Seguimiento de experiencia de usuario, encuestas, indicadores de percepcion y resumenes ejecutivos.",
        "status": "Proximamente",
        "accent": "#2f6ddc",
        "icon_svg": '<path d="M8 14s1.5 2 4 2 4-2 4-2" /><path d="M9 9h.01" /><path d="M15 9h.01" /><path d="M12 21c4.97 0 9-3.582 9-8s-4.03-8-9-8-9 3.582-9 8c0 2.394 1.184 4.542 3.06 6.008L6 21l2.563-1.025A10.11 10.11 0 0 0 12 21Z" />',
    },
    {
        "page": "Flujogramas",
        "title": "Flujogramas",
        "eyebrow": "Modulo planificado",
        "description": "Biblioteca de procesos, mapas operativos, versionado documental y vistas navegables de flujo de trabajo.",
        "status": "Proximamente",
        "accent": "#3b7be6",
        "icon_svg": '<path d="M12 3v4" /><path d="M12 17v4" /><path d="M5 12H3" /><path d="M21 12h-2" /><rect x="8" y="7" width="8" height="10" rx="2" /><path d="M8 12H5" /><path d="M19 12h-3" />',
    },
]

MODULE_PLACEHOLDERS = {
    "Relevamientos": {
        "headline": "Levantamiento de campo con estructura comun",
        "summary": "Este espacio quedara listo para programar formularios, evidencias y seguimiento de hallazgos por sede o peaje.",
        "items": [
            "fichas por punto relevado",
            "adjuntos fotograficos y comentarios",
            "estado de observaciones y responsables",
        ],
    },
    "Auditorias": {
        "headline": "Control operativo con trazabilidad",
        "summary": "Aqui podras concentrar revisiones, no conformidades, evidencias y cierres por auditoria.",
        "items": [
            "programacion de auditorias",
            "matriz de hallazgos y severidad",
            "seguimiento de planes de accion",
        ],
    },
    "Satisfaccion": {
        "headline": "Indicadores de experiencia y percepcion",
        "summary": "El modulo servira para centralizar encuestas, cortes por canal y tableros de satisfaccion.",
        "items": [
            "carga de encuestas y bases",
            "resumenes por periodo o sede",
            "hallazgos de experiencia del usuario",
        ],
    },
    "Flujogramas": {
        "headline": "Mapa visual de procesos y documentos",
        "summary": "Quedara como repositorio de diagramas, procedimientos y navegacion por proceso.",
        "items": [
            "catalogo de procesos por area",
            "versionado de diagramas y anexos",
            "consulta rapida para operacion y control",
        ],
    },
}


@st.cache_data(show_spinner=False)
def get_runtime_version_label() -> str:
    repo_dir = Path(__file__).parent
    try:
        short_sha = subprocess.run(
            ["git", "-C", str(repo_dir), "rev-parse", "--short", "HEAD"],
            capture_output=True,
            text=True,
            check=True,
        ).stdout.strip()
    except Exception:
        short_sha = "sin-git"
    return f"{PROCESSING_SIGNATURE_VERSION} | {short_sha}"

EXPECTED_COLUMNS = [
    "PEAJE",
    "CASETA",
    "SENTIDO",
    "FECHA",
    "VEHICULO",
    "PLACA",
    "LLEGADA COLA",
    "LLEGADA CASETA",
    "SALIDA CASETA",
    "T. TEC",
    "T. CASETA",
]

TIME_ALIASES = ["T1", "T2", "T3"]
TIME_COLUMNS = {
    "LLEGADA COLA": "T1",
    "LLEGADA CASETA": "T2",
    "SALIDA CASETA": "T3",
}
TIME_GROUP_KEYS = ["PEAJE", "CASETA", "SENTIDO", "FECHA_DIA"]
UMBRAL_SWAP_TIEMPOS_NEGATIVOS_SEGUNDOS = 10

EXPECTED_PLATE_LENGTHS = {6}
COMMON_PATTERNS = {"AAA999", "A9A999", "A99999", "AA9999", "AAAA99", "AA999"}
PLACEHOLDER_PLATES = {"AAA111", "BBB999", "ABC123", "XXX", "X"}

CONFUSION_LOOKUP = {
    "0": {"O"},
    "O": {"0"},
    "1": {"I"},
    "I": {"1"},
    "5": {"S"},
    "S": {"5"},
    "8": {"B"},
    "B": {"8"},
    "2": {"Z"},
    "Z": {"2"},
    "6": {"G"},
    "G": {"6"},
}

ACCIONES_PLACA_VALIDAS = {
    "corregir_a_normalizada",
    "corregir_a_sugerida",
    "corregir_a_coincidencia_lista",
    "corregir_a_recorte_sufijo_x",
    "consolidar_duplicado_cercano",
    "corregir_manual",
    "excluir_analisis_placa",
    "excluir_duplicado_cercano",
    "mantener_observada",
    "sin_cambio",
}

ACCIONES_PLACA_AJUSTE = {
    "corregir_a_normalizada",
    "corregir_a_sugerida",
    "corregir_a_coincidencia_lista",
    "corregir_a_recorte_sufijo_x",
    "consolidar_duplicado_cercano",
    "corregir_manual",
}
DEDUPLICACION_DUPLICADO_CERCANO_SEGUNDOS = 300
DEFAULT_MANUAL_RULES = [
    {
        "activo": True,
        "tipo_regla": "corregir_placa",
        "placa_origen": "D4P750UN",
        "placa_destino": "D4P750",
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Correccion manual heredada del caso Norvial.",
    },
    {
        "activo": True,
        "tipo_regla": "corregir_placa",
        "placa_origen": "TDA824UN",
        "placa_destino": "TDA824",
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Correccion manual heredada del caso Norvial.",
    },
    {
        "activo": True,
        "tipo_regla": "corregir_placa",
        "placa_origen": "D00191X",
        "placa_destino": "D0O191",
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Correccion manual heredada del caso Norvial.",
    },
    {
        "activo": True,
        "tipo_regla": "corregir_placa",
        "placa_origen": "BNI87D",
        "placa_destino": "BNI870",
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Correccion manual heredada del caso Norvial.",
    },
    {
        "activo": True,
        "tipo_regla": "corregir_placa",
        "placa_origen": "2PL067",
        "placa_destino": "ZPL067",
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Correccion manual heredada del caso Norvial.",
    },
    {
        "activo": True,
        "tipo_regla": "eliminar_placa",
        "placa_origen": "VCJV27CH",
        "placa_destino": pd.NA,
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Excluir del analisis.",
    },
    {
        "activo": True,
        "tipo_regla": "eliminar_placa",
        "placa_origen": "VCJV28CH",
        "placa_destino": pd.NA,
        "longitud_objetivo": pd.NA,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Excluir del analisis.",
    },
    {
        "activo": True,
        "tipo_regla": "eliminar_placa",
        "placa_origen": "VHT567",
        "placa_destino": pd.NA,
        "longitud_objetivo": pd.NA,
        "peaje": "VARIANTE",
        "caseta": "2",
        "sentido": "ASCENDENTE",
        "fecha": "2026-02-20",
        "comentario": "Excluir registro de prueba identificado en la toma preliminar.",
    },
    {
        "activo": True,
        "tipo_regla": "eliminar_por_longitud",
        "placa_origen": pd.NA,
        "placa_destino": pd.NA,
        "longitud_objetivo": 1,
        "peaje": pd.NA,
        "caseta": pd.NA,
        "sentido": pd.NA,
        "fecha": pd.NA,
        "comentario": "Excluir placas de largo 1.",
    },
]

DEFAULT_CONFIG = {
    "aplicar_limpieza_placa": True,
    "aplicar_ruido_con_respaldo": True,
    "aplicar_ruido_sin_respaldo": True,
    "aplicar_coincidencia_unica_lista": True,
    "aplicar_confusion_visual": True,
    "aplicar_recorte_sufijo_x": True,
    "aplicar_exclusion_placeholders": True,
    "aplicar_reglas_manuales": True,
    "aplicar_consolidacion_duplicados_cercanos": True,
    "eliminar_bordes_caseta": True,
    "aplicar_interpolacion": True,
    "aplicar_mediana_local": True,
    "aplicar_donantes": True,
    "aplicar_swap_final_t2_t3": True,
    "aplicar_swap_tiempos_completos_cortos": True,
    "modo_contraste_estricto": False,
}


def inject_global_styles() -> None:
    st.markdown(
        """
        <style>
            :root {
                --app-bg: #f3f7ff;
                --app-surface: rgba(255, 255, 255, 0.92);
                --app-surface-strong: #ffffff;
                --app-ink: #102347;
                --app-muted: #5c6f93;
                --app-line: rgba(15, 61, 145, 0.12);
                --app-accent: #0f3d91;
                --app-accent-soft: #e9f0ff;
                --app-shadow: 0 18px 45px rgba(9, 29, 74, 0.12);
            }

            .stApp {
                background:
                    radial-gradient(circle at top left, rgba(50, 100, 204, 0.18), transparent 28%),
                    linear-gradient(180deg, #eef4ff 0%, #f7faff 42%, #eef3fb 100%);
                color: var(--app-ink);
            }

            [data-testid="stSidebar"] {
                background: linear-gradient(180deg, #07162f 0%, #0d2448 100%);
                border-right: 1px solid rgba(255, 255, 255, 0.08);
            }

            [data-testid="stSidebar"] * {
                color: #f4f8ff;
            }

            [data-testid="stSidebar"] .stRadio label,
            [data-testid="stSidebar"] .stCheckbox label,
            [data-testid="stSidebar"] .stMarkdown,
            [data-testid="stSidebar"] .stCaption {
                color: #f4f8ff;
            }

            [data-testid="stSidebar"] .stButton > button {
                border-radius: 18px;
                border: 1px solid rgba(255, 255, 255, 0.18);
                background: rgba(255, 255, 255, 0.08);
                color: #f4f8ff;
                box-shadow: none;
            }

            [data-testid="stSidebar"] .stButton > button:hover {
                border-color: rgba(255, 255, 255, 0.3);
                background: rgba(255, 255, 255, 0.16);
                color: #ffffff;
            }

            [data-testid="stSidebar"] .stButton > button[kind="primary"] {
                background: linear-gradient(135deg, #2a7de1 0%, #59b2ff 100%);
                border-color: transparent;
                color: #ffffff;
            }

            [data-testid="stSidebar"] .stButton > button[kind="secondary"] {
                background: rgba(255, 255, 255, 0.1);
                color: #f4f8ff;
            }

            [data-testid="stSidebar"] .stButton > button p {
                font-size: 1.05rem;
                font-weight: 600;
            }

            [data-testid="stSidebar"] .stButton > button#sidebar_users_button {
                min-height: 3.25rem;
            }

            [data-testid="stSidebar"] .stButton > button#sidebar_users_button p {
                font-size: 1.45rem;
                line-height: 1;
            }

            .sidebar-brand-wrap {
                padding: 0.35rem 0 1rem;
            }

            .sidebar-brand-kicker {
                text-transform: uppercase;
                letter-spacing: 0.14em;
                font-size: 0.68rem;
                font-weight: 700;
                color: rgba(214, 230, 255, 0.7);
                margin-bottom: 0.35rem;
            }

            .sidebar-brand-note {
                font-size: 0.82rem;
                line-height: 1.45;
                color: rgba(230, 239, 255, 0.76);
                margin-top: 0.45rem;
            }

            .sidebar-text-link {
                margin-top: 0.55rem;
            }

            .sidebar-text-link a {
                color: #8fd0ff;
                text-decoration: none;
                font-size: 0.98rem;
                font-weight: 700;
            }

            .sidebar-text-link a:hover {
                color: #d6efff;
                text-decoration: underline;
            }

            .auth-layout-title {
                font-size: 1.1rem;
                font-weight: 700;
                color: var(--app-ink);
                margin: 0.4rem 0 0.25rem;
            }

            .auth-layout-copy {
                color: var(--app-muted);
                line-height: 1.7;
                margin-bottom: 1rem;
            }

            .auth-card {
                border-radius: 26px;
                border: 1px solid rgba(15, 61, 145, 0.12);
                background: linear-gradient(180deg, rgba(255,255,255,0.98), rgba(242,247,255,0.96));
                box-shadow: 0 18px 40px rgba(8, 29, 74, 0.09);
                padding: 1.4rem 1.4rem 1rem;
                margin-bottom: 1rem;
            }

            .auth-card-kicker {
                text-transform: uppercase;
                letter-spacing: 0.14em;
                font-size: 0.72rem;
                color: var(--app-accent);
                font-weight: 700;
                margin-bottom: 0.55rem;
            }

            .auth-card-title {
                font-size: 1.45rem;
                font-weight: 700;
                color: var(--app-ink);
                margin-bottom: 0.35rem;
            }

            .auth-card-copy {
                color: var(--app-muted);
                line-height: 1.65;
                margin-bottom: 0;
            }

            .brand-panel {
                border-radius: 24px;
                border: 1px solid rgba(255,255,255,0.14);
                background: rgba(255,255,255,0.08);
                padding: 1rem;
                backdrop-filter: blur(8px);
                margin-top: 1rem;
            }

            .brand-panel-copy {
                color: rgba(240, 246, 255, 0.82);
                font-size: 0.92rem;
                line-height: 1.6;
                margin-top: 0.7rem;
            }

            .hero-panel {
                padding: 2rem 2.2rem;
                border-radius: 28px;
                background:
                    linear-gradient(135deg, rgba(5, 20, 44, 0.96) 0%, rgba(15, 61, 145, 0.94) 55%, rgba(74, 125, 218, 0.88) 100%);
                box-shadow: var(--app-shadow);
                position: relative;
                overflow: hidden;
                margin-bottom: 1.2rem;
                color: #f6f9ff;
            }

            .hero-panel::after {
                content: "";
                position: absolute;
                inset: auto -90px -90px auto;
                width: 240px;
                height: 240px;
                border-radius: 50%;
                background: radial-gradient(circle, rgba(255,255,255,0.22), rgba(255,255,255,0.02) 70%);
            }

            .hero-kicker {
                text-transform: uppercase;
                letter-spacing: 0.18em;
                font-size: 0.72rem;
                color: rgba(235, 243, 255, 0.78);
                margin-bottom: 0.7rem;
                font-weight: 700;
            }

            .hero-title {
                font-size: clamp(2.5rem, 4vw, 4rem);
                line-height: 1.08;
                font-weight: 700;
                margin: 0;
                max-width: none;
                width: min(100%, 28ch);
                text-wrap: pretty;
            }

            .hero-copy {
                font-size: 1rem;
                line-height: 1.7;
                max-width: 72ch;
                margin-top: 0.9rem;
                color: rgba(245, 248, 255, 0.86);
            }

            .metrics-strip {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
                gap: 0.9rem;
                margin: 1rem 0 1.3rem;
            }

            .metric-tile {
                border: 1px solid rgba(255, 255, 255, 0.16);
                background: rgba(255, 255, 255, 0.08);
                backdrop-filter: blur(8px);
                border-radius: 18px;
                padding: 1rem 1rem 0.9rem;
            }

            .metric-value {
                font-size: 1.7rem;
                font-weight: 700;
                color: #ffffff;
            }

            .metric-label {
                margin-top: 0.2rem;
                font-size: 0.85rem;
                color: rgba(238, 245, 255, 0.78);
            }

            .section-heading {
                font-size: 1.25rem;
                font-weight: 700;
                color: var(--app-ink);
                margin: 1rem 0 0.3rem;
            }

            .section-copy {
                color: var(--app-muted);
                margin-bottom: 1rem;
            }

            .module-card {
                min-height: 240px;
                border-radius: 24px;
                padding: 1.2rem;
                background: linear-gradient(180deg, rgba(255,255,255,0.98), rgba(244,248,255,0.92));
                border: 1px solid var(--app-line);
                box-shadow: 0 14px 35px rgba(12, 31, 76, 0.08);
                position: relative;
                overflow: hidden;
                margin-bottom: 0.75rem;
            }

            .module-card::before {
                content: "";
                position: absolute;
                inset: 0 0 auto 0;
                height: 5px;
                background: var(--module-accent);
            }

            .module-icon {
                width: 52px;
                height: 52px;
                border-radius: 16px;
                display: grid;
                place-items: center;
                color: var(--module-accent);
                background: rgba(15, 61, 145, 0.08);
                border: 1px solid rgba(15, 61, 145, 0.12);
            }

            .module-icon svg {
                width: 26px;
                height: 26px;
                stroke: currentColor;
                stroke-width: 1.7;
                fill: none;
                stroke-linecap: round;
                stroke-linejoin: round;
            }

            .module-eyebrow {
                margin-top: 1rem;
                text-transform: uppercase;
                letter-spacing: 0.12em;
                font-size: 0.72rem;
                font-weight: 700;
                color: var(--module-accent);
            }

            .module-title {
                font-size: 1.35rem;
                font-weight: 700;
                margin: 0.45rem 0 0.35rem;
                color: var(--app-ink);
            }

            .module-description {
                color: var(--app-muted);
                line-height: 1.6;
                font-size: 0.95rem;
                min-height: 74px;
            }

            .module-status {
                display: inline-flex;
                align-items: center;
                gap: 0.4rem;
                border-radius: 999px;
                padding: 0.38rem 0.72rem;
                background: rgba(15, 61, 145, 0.08);
                color: var(--module-accent);
                font-size: 0.82rem;
                font-weight: 700;
            }

            .placeholder-card {
                border-radius: 24px;
                border: 1px solid var(--app-line);
                background: rgba(255, 255, 255, 0.95);
                padding: 1.45rem;
                box-shadow: 0 16px 34px rgba(12, 31, 76, 0.08);
            }

            .placeholder-title {
                font-size: 1.55rem;
                font-weight: 700;
                color: var(--app-ink);
                margin-bottom: 0.35rem;
            }

            .placeholder-copy {
                color: var(--app-muted);
                line-height: 1.7;
            }

            .placeholder-list {
                margin: 0.9rem 0 0;
                padding-left: 1rem;
                color: var(--app-ink);
            }

            .placeholder-list li {
                margin-bottom: 0.45rem;
            }

            div[data-testid="stMetric"] {
                background: rgba(255,255,255,0.92);
                border: 1px solid var(--app-line);
                border-radius: 20px;
                padding: 0.8rem 1rem;
                box-shadow: 0 10px 28px rgba(12, 31, 76, 0.07);
            }

            @media (max-width: 900px) {
                .hero-title {
                    font-size: 2rem;
                    width: min(100%, 11ch);
                    text-wrap: pretty;
                }

                .hero-panel {
                    padding: 1.5rem;
                }
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


def build_hero_panel(title: str, copy: str, kicker: str, metrics: list[tuple[str, str]] | None = None) -> str:
    metrics_markup = ""
    if metrics:
        metric_items = []
        for value, label in metrics:
            metric_items.append(
                f'<div class="metric-tile"><div class="metric-value">{value}</div><div class="metric-label">{label}</div></div>'
            )
        metrics_markup = f'<div class="metrics-strip">{"".join(metric_items)}</div>'

    return (
        '<section class="hero-panel">'
        f'<div class="hero-kicker">{kicker}</div>'
        f'<h1 class="hero-title">{title}</h1>'
        f'<div class="hero-copy">{copy}</div>'
        f'{metrics_markup}'
        '</section>'
    )


def build_module_card(module: dict) -> str:
    return (
        f'<section class="module-card" style="--module-accent:{module["accent"]};">'
        f'<div class="module-icon"><svg viewBox="0 0 24 24" aria-hidden="true">{module["icon_svg"]}</svg></div>'
        f'<div class="module-eyebrow">{module["eyebrow"]}</div>'
        f'<div class="module-title">{module["title"]}</div>'
        f'<div class="module-description">{module["description"]}</div>'
        f'<div class="module-status">{module["status"]}</div>'
        '</section>'
    )


@st.cache_data(show_spinner=False)
def load_contractor_logo() -> Image.Image | None:
    if not CONTRACTOR_LOGO_PATH.exists():
        return None

    image = Image.open(CONTRACTOR_LOGO_PATH).convert("RGBA")
    alpha_channel = image.getchannel("A")
    bbox = alpha_channel.point(lambda value: 255 if value > 20 else 0).getbbox()
    if bbox is None:
        background = Image.new("RGBA", image.size, (255, 255, 255, 255))
        diff = ImageChops.difference(image, background)
        bbox = diff.getbbox()
    if bbox:
        left, top, right, bottom = bbox
        padding = 8
        left = max(0, left - padding)
        top = max(0, top - padding)
        right = min(image.width, right + padding)
        bottom = min(image.height, bottom + padding)
        image = image.crop((left, top, right, bottom))

    return image


def render_contractor_branding() -> None:
    logo_image = load_contractor_logo()
    if logo_image is None:
        return

    left_col, center_col, right_col = st.columns([0.12, 0.76, 0.12])
    with center_col:
        st.image(logo_image, use_container_width=True)


def navigate_to(page: str) -> None:
    if page == "Inicio":
        st.query_params.clear()
    else:
        st.query_params["page"] = page
    st.session_state[APP_NAV_KEY] = page
    st.rerun()


def render_back_to_home_button(page_key: str) -> None:
    if st.button(f"← Volver a Inicio", key=f"back_{page_key}", use_container_width=False):
        navigate_to("Inicio")


def get_requested_page() -> str | None:
    raw_page = st.query_params.get("page")
    if raw_page is None:
        return None
    if isinstance(raw_page, list):
        raw_page = raw_page[0] if raw_page else None
    if raw_page is None:
        return None
    page = str(raw_page).strip()
    return page or None


def render_home_page(current_user: dict | None) -> None:
    user_label = current_user["full_name"] if current_user else "Equipo CIDATT"
    st.markdown(
        build_hero_panel(
            title="Centro de control operativo y analitico",
            copy=(
                "Una portada unificada para operar los modulos del aplicativo con una interfaz sobria, "
                "clara y lista para crecer. Desde aqui entras a TEC y a los espacios que luego completaremos."
            ),
            kicker="Suite Operativa",
        ),
        unsafe_allow_html=True,
    )

    st.markdown('<div class="section-heading">Modulos disponibles</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="section-copy">Sesion iniciada como {user_label}. Selecciona un modulo para continuar.</div>',
        unsafe_allow_html=True,
    )

    columns_per_row = 3
    for start_index in range(0, len(MODULE_CATALOG), columns_per_row):
        row_modules = MODULE_CATALOG[start_index : start_index + columns_per_row]
        columns = st.columns(len(row_modules))
        for column, module in zip(columns, row_modules):
            with column:
                st.markdown(build_module_card(module), unsafe_allow_html=True)
                button_label = "Abrir modulo" if module["page"] == "TEC" else "Ver portada"
                if st.button(button_label, key=f'home_{module["page"]}', use_container_width=True):
                    navigate_to(module["page"])


def render_placeholder_module_page(page: str) -> None:
    render_back_to_home_button(page)
    module = next((item for item in MODULE_CATALOG if item["page"] == page), None)
    placeholder = MODULE_PLACEHOLDERS[page]
    st.markdown(
        build_hero_panel(
            title=module["title"] if module else page,
            copy=placeholder["summary"],
            kicker=module["eyebrow"] if module else "Modulo",
            metrics=[
                (module["status"] if module else "Proximamente", "estado"),
                ("Azul marino", "linea visual"),
                ("En definicion", "siguiente etapa"),
            ],
        ),
        unsafe_allow_html=True,
    )
    items_markup = "".join(f"<li>{item}</li>" for item in placeholder["items"])
    st.markdown(
        (
            '<section class="placeholder-card">'
            f'<div class="placeholder-title">{placeholder["headline"]}</div>'
            f'<div class="placeholder-copy">{placeholder["summary"]}</div>'
            f'<ul class="placeholder-list">{items_markup}</ul>'
            '</section>'
        ),
        unsafe_allow_html=True,
    )


def suggest_column(columns: list[str], patterns: list[str]) -> str | None:
    normalized = {col: re.sub(r"[^a-z0-9]+", "", col.lower()) for col in columns}
    for pattern in patterns:
        target = re.sub(r"[^a-z0-9]+", "", pattern.lower())
        for col, normalized_col in normalized.items():
            if target == normalized_col or target in normalized_col:
                return col
    return None


def patron_alfanumerico(plate: str) -> str:
    return "".join("A" if char.isalpha() else "9" if char.isdigit() else char for char in plate)


def generar_candidatas_confusion(plate: str) -> list[str]:
    candidates = set()
    for idx, char in enumerate(plate):
        for alternative in CONFUSION_LOOKUP.get(char, set()):
            candidates.add(plate[:idx] + alternative + plate[idx + 1 :])
    return sorted(candidates)


def normalizar_hora(valor):
    if pd.isna(valor):
        return pd.NaT
    if isinstance(valor, pd.Timedelta):
        return valor
    if isinstance(valor, timedelta):
        return pd.to_timedelta(valor)
    if isinstance(valor, pd.Timestamp):
        return (
            pd.to_timedelta(valor.hour, unit="h")
            + pd.to_timedelta(valor.minute, unit="m")
            + pd.to_timedelta(valor.second, unit="s")
        )
    if isinstance(valor, time):
        return (
            pd.to_timedelta(valor.hour, unit="h")
            + pd.to_timedelta(valor.minute, unit="m")
            + pd.to_timedelta(valor.second, unit="s")
        )
    if isinstance(valor, (int, float)) and not isinstance(valor, bool) and 0 <= float(valor) < 2:
        return pd.to_timedelta(float(valor), unit="D")
    texto = str(valor).strip()
    if not texto or texto.lower() in {"nat", "nan", "none"}:
        return pd.NaT
    como_timedelta = pd.to_timedelta(texto, errors="coerce")
    if pd.notna(como_timedelta):
        return como_timedelta
    como_datetime = pd.to_datetime(texto, errors="coerce")
    if pd.notna(como_datetime):
        return (
            pd.to_timedelta(como_datetime.hour, unit="h")
            + pd.to_timedelta(como_datetime.minute, unit="m")
            + pd.to_timedelta(como_datetime.second, unit="s")
        )
    return pd.NaT


def formatear_hora(valor):
    if pd.isna(valor):
        return pd.NA
    total_segundos = int(valor.total_seconds())
    horas = total_segundos // 3600
    minutos = (total_segundos % 3600) // 60
    segundos = total_segundos % 60
    return f"{horas:02}:{minutos:02}:{segundos:02}"


def segundos_a_timedelta(valor):
    if pd.isna(valor):
        return pd.NaT
    return pd.to_timedelta(float(valor), unit="s")


def timedelta_a_segundos(valor):
    if pd.isna(valor):
        return pd.NA
    return int(valor.total_seconds())


def timedelta_a_minutos(valor):
    if pd.isna(valor):
        return pd.NA
    return round(valor.total_seconds() / 60, 2)


def parse_manual_rules(rules_df: pd.DataFrame) -> pd.DataFrame:
    reglas = rules_df.copy()
    expected_columns = [
        "activo",
        "tipo_regla",
        "placa_origen",
        "placa_destino",
        "longitud_objetivo",
        "peaje",
        "caseta",
        "sentido",
        "fecha",
        "comentario",
    ]
    for column_name in expected_columns:
        if column_name not in reglas.columns:
            reglas[column_name] = pd.NA
    if reglas.empty:
        return reglas[expected_columns]
    reglas = reglas[reglas["activo"].fillna(False)].copy()
    reglas["tipo_regla"] = reglas["tipo_regla"].astype(str).str.strip()
    reglas["placa_origen"] = reglas["placa_origen"].fillna("").astype(str).str.upper().str.strip()
    reglas["placa_destino"] = reglas["placa_destino"].fillna("").astype(str).str.upper().str.strip()
    reglas["longitud_objetivo"] = pd.to_numeric(reglas["longitud_objetivo"], errors="coerce")
    for scope_column in ["peaje", "caseta", "sentido"]:
        reglas[scope_column] = reglas[scope_column].fillna("").astype(str).str.upper().str.strip()
    reglas["fecha"] = pd.to_datetime(reglas["fecha"], errors="coerce").dt.normalize()
    return reglas[expected_columns]


def manual_rule_matches_row(rule: pd.Series, row: pd.Series, normalized_plate: str) -> bool:
    rule_type = str(rule.get("tipo_regla") or "").strip()

    if rule_type in {"corregir_placa", "eliminar_placa"}:
        if not str(rule.get("placa_origen") or "").strip():
            return False
        if normalized_plate != str(rule.get("placa_origen") or "").strip():
            return False
    elif rule_type == "eliminar_por_longitud":
        length_target = rule.get("longitud_objetivo")
        if pd.isna(length_target) or len(normalized_plate) != int(length_target):
            return False
    else:
        return False

    peaje_rule = str(rule.get("peaje") or "").strip()
    if peaje_rule and str(row.get("PEAJE") or "").upper().strip() != peaje_rule:
        return False

    caseta_rule = str(rule.get("caseta") or "").strip()
    if caseta_rule and str(row.get("CASETA") or "").upper().strip() != caseta_rule:
        return False

    sentido_rule = str(rule.get("sentido") or "").strip()
    if sentido_rule and str(row.get("SENTIDO") or "").upper().strip() != sentido_rule:
        return False

    fecha_rule = rule.get("fecha")
    if pd.notna(fecha_rule):
        row_date = pd.to_datetime(row.get("FECHA"), errors="coerce")
        if pd.isna(row_date) or row_date.normalize() != fecha_rule:
            return False

    return True


def apply_manual_rules_to_df(df: pd.DataFrame, rules_df: pd.DataFrame) -> pd.DataFrame:
    reglas = parse_manual_rules(rules_df)
    if reglas.empty:
        return df

    for idx, row in df.iterrows():
        normalized = str(row["PLACA_NORMALIZADA"])
        for _, rule in reglas.iterrows():
            if not manual_rule_matches_row(rule, row, normalized):
                continue

            rule_type = str(rule["tipo_regla"])
            comentario = str(rule.get("comentario") or "").strip()
            scope_parts = []
            for scope_column, row_column in [("peaje", "PEAJE"), ("caseta", "CASETA"), ("sentido", "SENTIDO")]:
                rule_value = str(rule.get(scope_column) or "").strip()
                if rule_value:
                    scope_parts.append(f"{scope_column}={row.get(row_column)}")
            if pd.notna(rule.get("fecha")):
                scope_parts.append(f"fecha={pd.to_datetime(row.get('FECHA'), errors='coerce').strftime('%Y-%m-%d')}")
            scope_text = f" [{', '.join(scope_parts)}]" if scope_parts else ""

            if rule_type == "corregir_placa":
                target = str(rule.get("placa_destino") or "").strip()
                if not target:
                    continue
                df.at[idx, "PLACA_ACCION_FINAL"] = "corregir_manual"
                df.at[idx, "PLACA_FINAL_DECIDIDA"] = target
                df.at[idx, "PLACA_EXCLUIR_ANALISIS"] = False
                df.at[idx, "PLACA_AJUSTE_MANUAL"] = f"Correccion manual{scope_text}: {normalized} -> {target}. {comentario}".strip()
            elif rule_type == "eliminar_placa":
                df.at[idx, "PLACA_ACCION_FINAL"] = "excluir_analisis_placa"
                df.at[idx, "PLACA_EXCLUIR_ANALISIS"] = True
                df.at[idx, "PLACA_AJUSTE_MANUAL"] = f"Eliminacion manual{scope_text}: {normalized}. {comentario}".strip()
            elif rule_type == "eliminar_por_longitud" and not bool(df.at[idx, "PLACA_EXCLUIR_ANALISIS"]):
                df.at[idx, "PLACA_ACCION_FINAL"] = "excluir_analisis_placa"
                df.at[idx, "PLACA_EXCLUIR_ANALISIS"] = True
                df.at[idx, "PLACA_AJUSTE_MANUAL"] = f"Eliminacion manual por longitud {len(normalized)}{scope_text}. {comentario}".strip()
            break

    return df


def load_input_dataframe(uploaded_file, sheet_name: str | None) -> pd.DataFrame:
    name = uploaded_file.name.lower()
    uploaded_file.seek(0)
    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file)
    if name.endswith(".xlsx") or name.endswith(".xlsm") or name.endswith(".xls"):
        return pd.read_excel(uploaded_file, sheet_name=sheet_name or 0)
    raise ValueError("Formato no soportado. Usa CSV o Excel.")


def list_excel_sheets(uploaded_file) -> list[str]:
    name = uploaded_file.name.lower()
    if not (name.endswith(".xlsx") or name.endswith(".xlsm") or name.endswith(".xls")):
        return []
    uploaded_file.seek(0)
    excel = pd.ExcelFile(uploaded_file)
    return list(excel.sheet_names)


def build_standardized_df(df_raw: pd.DataFrame, mapping: dict[str, str | None]) -> pd.DataFrame:
    df_std = pd.DataFrame(index=df_raw.index)
    for target in EXPECTED_COLUMNS:
        source = mapping.get(target)
        if source:
            df_std[target] = df_raw[source]
        elif target == "VEHICULO":
            df_std[target] = pd.RangeIndex(1, len(df_raw) + 1)
        elif target in {"PEAJE", "CASETA", "SENTIDO"}:
            df_std[target] = f"SIN_{target}"
        elif target == "FECHA":
            df_std[target] = pd.NaT
        else:
            df_std[target] = pd.NA
    return df_std


def combine_action(placa_action: str, tiempo_action: str) -> str:
    partes = []
    if placa_action and placa_action != "sin_cambio":
        partes.append(f"placa:{placa_action}")
    if tiempo_action and tiempo_action != "sin_cambio":
        partes.append(f"tiempo:{tiempo_action}")
    return "sin_cambio" if not partes else " | ".join(partes)


def append_note(existing_note, new_note: str) -> str:
    if not new_note:
        return existing_note
    if pd.isna(existing_note) or not str(existing_note).strip():
        return new_note
    existing_text = str(existing_note).strip()
    if new_note in existing_text:
        return existing_text
    return f"{existing_text} | {new_note}"


def build_duplicate_reference_time(df: pd.DataFrame) -> pd.Series:
    t1 = df["LLEGADA COLA"].map(normalizar_hora)
    t2 = df["LLEGADA CASETA"].map(normalizar_hora)
    t3 = df["SALIDA CASETA"].map(normalizar_hora)
    return t2.combine_first(t1).combine_first(t3)


def consolidate_near_duplicate_plate_rows(df_in: pd.DataFrame, window_seconds: int) -> pd.DataFrame:
    df = df_in.copy()
    if df.empty:
        return df

    df["_DEDUP_FECHA"] = pd.to_datetime(df["FECHA"], errors="coerce").dt.normalize()
    df["_DEDUP_REF"] = build_duplicate_reference_time(df)
    df["_DEDUP_T1"] = df["LLEGADA COLA"].map(normalizar_hora)
    df["_DEDUP_T2"] = df["LLEGADA CASETA"].map(normalizar_hora)
    df["_DEDUP_T3"] = df["SALIDA CASETA"].map(normalizar_hora)
    df["_DEDUP_COMPLETITUD"] = df[["_DEDUP_T1", "_DEDUP_T2", "_DEDUP_T3"]].notna().sum(axis=1)

    group_cols = ["PEAJE", "SENTIDO", "_DEDUP_FECHA", "PLACA_FINAL_DECIDIDA"]
    valid_mask = (
        df["PLACA_EXCLUIR_ANALISIS"].eq(False)
        & df["PLACA_FINAL_DECIDIDA"].notna()
        & df["_DEDUP_FECHA"].notna()
        & df["_DEDUP_REF"].notna()
    )
    if not valid_mask.any():
        return df.drop(columns=[col for col in df.columns if col.startswith("_DEDUP_")])

    for _, sub_df in df[valid_mask].sort_values(group_cols + ["_DEDUP_REF", "_ORDEN_FILA"]).groupby(group_cols, dropna=False):
        if len(sub_df) < 2:
            continue

        clusters: list[list[int]] = []
        current_cluster: list[int] = []
        previous_ref = None
        for row_idx, row in sub_df.iterrows():
            current_ref = row["_DEDUP_REF"]
            if not current_cluster:
                current_cluster = [row_idx]
            elif previous_ref is not None and (current_ref - previous_ref).total_seconds() <= window_seconds:
                current_cluster.append(row_idx)
            else:
                if len(current_cluster) > 1:
                    clusters.append(current_cluster)
                current_cluster = [row_idx]
            previous_ref = current_ref
        if len(current_cluster) > 1:
            clusters.append(current_cluster)

        for cluster in clusters:
            cluster_df = df.loc[cluster].copy()
            cluster_df = cluster_df.sort_values(
                ["_DEDUP_COMPLETITUD", "_DEDUP_T2", "_DEDUP_T3", "_DEDUP_T1", "_ORDEN_FILA"],
                ascending=[False, False, False, False, True],
                na_position="last",
            )
            canonical_idx = cluster_df.index[0]
            source_indices = [idx for idx in cluster_df.index if idx != canonical_idx]
            if not source_indices:
                continue

            source_rows = df.loc[source_indices].copy()
            source_casetas = ", ".join(sorted(source_rows["CASETA"].astype(str).unique()))
            source_ordenes = ", ".join(str(int(value)) for value in source_rows["_ORDEN_FILA"].tolist())

            for source_idx in source_indices:
                for source_col, target_col in [
                    ("LLEGADA COLA", "_DEDUP_T1"),
                    ("LLEGADA CASETA", "_DEDUP_T2"),
                    ("SALIDA CASETA", "_DEDUP_T3"),
                    ("T. TEC", None),
                    ("T. CASETA", None),
                ]:
                    if pd.isna(df.at[canonical_idx, source_col]) and pd.notna(df.at[source_idx, source_col]):
                        df.at[canonical_idx, source_col] = df.at[source_idx, source_col]
                        if target_col:
                            df.at[canonical_idx, target_col] = df.at[source_idx, target_col]

            canonical_note = (
                "Consolidacion de duplicado cercano: misma placa final, mismo peaje/sentido/fecha, "
                f"fusionado con filas {_safe_text(source_ordenes)} y casetas {_safe_text(source_casetas)}."
            )
            df.at[canonical_idx, "PLACA_AJUSTE_MANUAL"] = append_note(df.at[canonical_idx, "PLACA_AJUSTE_MANUAL"], canonical_note)
            if str(df.at[canonical_idx, "PLACA_ACCION_FINAL"]) == "sin_cambio":
                df.at[canonical_idx, "PLACA_ACCION_FINAL"] = "consolidar_duplicado_cercano"

            for source_idx in source_indices:
                source_note = (
                    "Registro consolidado en duplicado cercano con la misma placa final; "
                    f"fila canonica {_safe_text(int(df.at[canonical_idx, '_ORDEN_FILA']))}, caseta canonica {_safe_text(df.at[canonical_idx, 'CASETA'])}."
                )
                df.at[source_idx, "PLACA_ACCION_FINAL"] = "excluir_duplicado_cercano"
                df.at[source_idx, "PLACA_EXCLUIR_ANALISIS"] = True
                df.at[source_idx, "PLACA_AJUSTE_MANUAL"] = append_note(df.at[source_idx, "PLACA_AJUSTE_MANUAL"], source_note)

    return df.drop(columns=[col for col in df.columns if col.startswith("_DEDUP_")], errors="ignore")


def consolidate_post_time_rows(df_in: pd.DataFrame, window_seconds: int) -> pd.DataFrame:
    df = df_in.copy()
    if df.empty:
        return df

    df["_POST_REF"] = df["T2_FINAL_5TA"].combine_first(df["T1_FINAL_5TA"]).combine_first(df["T3_FINAL_5TA"])
    df["_POST_ORIG_COMPLETITUD"] = df[["T1", "T2", "T3"]].notna().sum(axis=1)
    df["_POST_CASETA_SCORE"] = (
        df["T2"].notna().astype(int) * 3
        + df["T3"].notna().astype(int) * 3
        + df["T1"].notna().astype(int)
    )

    group_cols = ["PEAJE", "SENTIDO", "FECHA_DIA", "PLACA_FINAL"]
    valid_mask = (
        df["PLACA_FINAL"].notna()
        & df["FECHA_DIA"].notna()
        & df["_POST_REF"].notna()
    )
    if not valid_mask.any():
        return df.drop(columns=[col for col in df.columns if col.startswith("_POST_")], errors="ignore")

    rows_to_drop: list[int] = []
    for _, sub_df in df[valid_mask].sort_values(group_cols + ["_POST_REF", "_ORDEN_FILA"]).groupby(group_cols, dropna=False):
        if len(sub_df) < 2:
            continue

        clusters: list[list[int]] = []
        current_cluster: list[int] = []
        previous_ref = None
        for row_idx, row in sub_df.iterrows():
            current_ref = row["_POST_REF"]
            if not current_cluster:
                current_cluster = [row_idx]
            elif previous_ref is not None and (current_ref - previous_ref).total_seconds() <= window_seconds:
                current_cluster.append(row_idx)
            else:
                if len(current_cluster) > 1:
                    clusters.append(current_cluster)
                current_cluster = [row_idx]
            previous_ref = current_ref
        if len(current_cluster) > 1:
            clusters.append(current_cluster)

        for cluster in clusters:
            cluster_df = df.loc[cluster].copy()
            exact_duplicate_mask = cluster_df.duplicated(
                subset=["CASETA", "T1_FINAL_5TA", "T2_FINAL_5TA", "T3_FINAL_5TA"],
                keep=False,
            )
            should_merge = (~cluster_df["TIEMPOS_COMPLETOS"]).any() or exact_duplicate_mask.any()
            if not should_merge:
                continue

            cluster_df = cluster_df.sort_values(
                ["_POST_CASETA_SCORE", "_POST_ORIG_COMPLETITUD", "_POST_REF", "_ORDEN_FILA"],
                ascending=[False, False, True, True],
                na_position="last",
            )
            canonical_idx = cluster_df.index[0]
            source_indices = [idx for idx in cluster_df.index if idx != canonical_idx]
            if not source_indices:
                continue

            observed_t1 = cluster_df.loc[cluster_df["T1"].notna(), "T1"]
            if not observed_t1.empty:
                df.at[canonical_idx, "T1_FINAL_5TA"] = observed_t1.min()
                df.at[canonical_idx, "LLEGADA_COLA_FINAL"] = formatear_hora(df.at[canonical_idx, "T1_FINAL_5TA"])

            canonical_t2 = df.at[canonical_idx, "T2_FINAL_5TA"]
            canonical_t3 = df.at[canonical_idx, "T3_FINAL_5TA"]
            if pd.notna(df.at[canonical_idx, "T1_FINAL_5TA"]) and pd.notna(canonical_t2) and pd.notna(canonical_t3):
                df.at[canonical_idx, "TIEMPOS_COMPLETOS_CIERRE"] = True
                df.at[canonical_idx, "T_COLA_FINAL"] = canonical_t2 - df.at[canonical_idx, "T1_FINAL_5TA"]
                df.at[canonical_idx, "T_CASETA_FINAL"] = canonical_t3 - canonical_t2
                df.at[canonical_idx, "T_TEC_FINAL"] = canonical_t3 - df.at[canonical_idx, "T1_FINAL_5TA"]
                df.at[canonical_idx, "T_COLA_FINAL_TXT"] = formatear_hora(df.at[canonical_idx, "T_COLA_FINAL"])
                df.at[canonical_idx, "T_CASETA_FINAL_TXT"] = formatear_hora(df.at[canonical_idx, "T_CASETA_FINAL"])
                df.at[canonical_idx, "T_TEC_FINAL_TXT"] = formatear_hora(df.at[canonical_idx, "T_TEC_FINAL"])
                df.at[canonical_idx, "LLEGADA_CASETA_FINAL"] = formatear_hora(canonical_t2)
                df.at[canonical_idx, "SALIDA_CASETA_FINAL"] = formatear_hora(canonical_t3)

            source_rows = df.loc[source_indices].copy()
            source_ordenes = ", ".join(str(int(value)) for value in source_rows["_ORDEN_FILA"].tolist())
            source_casetas = ", ".join(sorted(source_rows["CASETA"].astype(str).unique()))
            merge_note = (
                "Consolidacion posterior a tiempos: misma placa final dentro de ventana corta, "
                f"fusionado con filas {_safe_text(source_ordenes)} y casetas {_safe_text(source_casetas)}."
            )
            current_action = str(df.at[canonical_idx, "TIEMPO_ACCION_CIERRE"] or "")
            if current_action == "sin_cambio":
                df.at[canonical_idx, "TIEMPO_ACCION_CIERRE"] = "consolidar_post_tiempo"
            df.at[canonical_idx, "TIEMPO_MOTIVO_CIERRE"] = append_note(df.at[canonical_idx, "TIEMPO_MOTIVO_CIERRE"], merge_note)
            rows_to_drop.extend(source_indices)

    if rows_to_drop:
        df = df.drop(index=sorted(set(rows_to_drop)))

    return df.drop(columns=[col for col in df.columns if col.startswith("_POST_")], errors="ignore")


def consolidate_fragmented_flow_rows(df_in: pd.DataFrame, max_delta_seconds: int = 180) -> pd.DataFrame:
    df = df_in.copy()
    if df.empty:
        return df

    ref_t1 = "T1_FINAL_5TA" if "T1_FINAL_5TA" in df.columns else "T1"
    ref_t2 = "T2_FINAL_5TA" if "T2_FINAL_5TA" in df.columns else "T2"
    ref_t3 = "T3_FINAL_5TA" if "T3_FINAL_5TA" in df.columns else "T3"
    if not {ref_t1, ref_t2, ref_t3}.issubset(df.columns):
        return df

    rows_to_drop: list[int] = []
    group_cols = ["PEAJE", "CASETA", "SENTIDO", "FECHA_DIA"]
    valid_mask = df["FECHA_DIA"].notna()
    if not valid_mask.any():
        return df

    for _, group in df[valid_mask].sort_values(group_cols + [ref_t1, ref_t2, ref_t3, "_ORDEN_FILA"], na_position="last").groupby(group_cols, dropna=False):
        group = group.reset_index()
        paired_rows: set[int] = set()
        complete_mask = group[[ref_t1, ref_t2, ref_t3]].notna().all(axis=1)
        for idx, row in group.iterrows():
            row_index = int(row["index"])
            if row_index in rows_to_drop or row_index in paired_rows or complete_mask.iloc[idx]:
                continue

            t1 = row[ref_t1]
            t2 = row[ref_t2]
            t3 = row[ref_t3]
            if pd.isna(t1) or pd.notna(t2) or pd.notna(t3):
                continue

            prev_complete = complete_mask.iloc[:idx].any()
            next_complete = complete_mask.iloc[idx + 1 :].any()
            if not (prev_complete and next_complete):
                continue

            for look_ahead in range(idx + 1, min(idx + 5, len(group))):
                other = group.iloc[look_ahead]
                other_index = int(other["index"])
                if other_index in rows_to_drop or other_index in paired_rows:
                    continue
                if pd.notna(other[ref_t1]) or pd.isna(other[ref_t2]) or pd.isna(other[ref_t3]):
                    continue
                delta_seg = (other[ref_t2] - t1).total_seconds()
                if delta_seg < 0 or delta_seg > max_delta_seconds:
                    continue
                is_match, confidence = classify_fragment_similarity(row.get("PLACA_FINAL"), other.get("PLACA_FINAL"), delta_seg)
                if not is_match or confidence != "alta":
                    continue

                canonical_idx = other_index
                source_idx = row_index
                df.at[canonical_idx, ref_t1] = t1
                df.at[canonical_idx, "TIEMPO_ACCION_CIERRE"] = "consolidar_fragmentacion_flujo"
                fragment_note = (
                    "Consolidacion por fragmentacion de flujo: una fila aporto cola y otra aporto caseta/salida; "
                    f"fila origen {_safe_text(int(df.at[source_idx, '_ORDEN_FILA']))}, delta {_safe_text(int(delta_seg))} s."
                )
                df.at[canonical_idx, "TIEMPO_MOTIVO_CIERRE"] = append_note(df.at[canonical_idx, "TIEMPO_MOTIVO_CIERRE"], fragment_note)
                rows_to_drop.append(source_idx)
                paired_rows.add(other_index)
                break

    if rows_to_drop:
        df = df.drop(index=sorted(set(rows_to_drop)))
    return refresh_final_time_outputs(df)


def refresh_final_time_outputs(df_in: pd.DataFrame) -> pd.DataFrame:
    df = df_in.copy()
    df["TIEMPOS_COMPLETOS_CIERRE"] = df[["T1_FINAL_5TA", "T2_FINAL_5TA", "T3_FINAL_5TA"]].notna().all(axis=1)
    df["LLEGADA_COLA_FINAL"] = df["T1_FINAL_5TA"].map(formatear_hora)
    df["LLEGADA_CASETA_FINAL"] = df["T2_FINAL_5TA"].map(formatear_hora)
    df["SALIDA_CASETA_FINAL"] = df["T3_FINAL_5TA"].map(formatear_hora)
    df["T_COLA_FINAL"] = df["T2_FINAL_5TA"] - df["T1_FINAL_5TA"]
    df["T_CASETA_FINAL"] = df["T3_FINAL_5TA"] - df["T2_FINAL_5TA"]
    df["T_TEC_FINAL"] = df["T3_FINAL_5TA"] - df["T1_FINAL_5TA"]
    df["T_COLA_FINAL_TXT"] = df["T_COLA_FINAL"].map(formatear_hora)
    df["T_CASETA_FINAL_TXT"] = df["T_CASETA_FINAL"].map(formatear_hora)
    df["T_TEC_FINAL_TXT"] = df["T_TEC_FINAL"].map(formatear_hora)
    return df


def apply_short_complete_time_swaps(df_in: pd.DataFrame, config: dict) -> pd.DataFrame:
    df = df_in.copy()
    if df.empty or not config.get("aplicar_swap_tiempos_completos_cortos", DEFAULT_CONFIG["aplicar_swap_tiempos_completos_cortos"]):
        return refresh_final_time_outputs(df)

    umbral_swap = pd.to_timedelta(UMBRAL_SWAP_TIEMPOS_NEGATIVOS_SEGUNDOS, unit="s")
    mascara_swap_t1_t2 = (
        df[["T1_FINAL_5TA", "T2_FINAL_5TA", "T3_FINAL_5TA"]].notna().all(axis=1)
        & (df["T1_FINAL_5TA"] > df["T2_FINAL_5TA"])
        & ((df["T1_FINAL_5TA"] - df["T2_FINAL_5TA"]) <= umbral_swap)
        & (df["T1_FINAL_5TA"] <= df["T3_FINAL_5TA"])
    )
    t1_swap_values = df.loc[mascara_swap_t1_t2, "T1_FINAL_5TA"].copy()
    df.loc[mascara_swap_t1_t2, "T1_FINAL_5TA"] = df.loc[mascara_swap_t1_t2, "T2_FINAL_5TA"]
    df.loc[mascara_swap_t1_t2, "T2_FINAL_5TA"] = t1_swap_values

    mascara_swap_t2_t3 = (
        df[["T1_FINAL_5TA", "T2_FINAL_5TA", "T3_FINAL_5TA"]].notna().all(axis=1)
        & (df["T2_FINAL_5TA"] > df["T3_FINAL_5TA"])
        & ((df["T2_FINAL_5TA"] - df["T3_FINAL_5TA"]) <= umbral_swap)
        & (df["T1_FINAL_5TA"] <= df["T3_FINAL_5TA"])
    )
    t2_swap_values = df.loc[mascara_swap_t2_t3, "T2_FINAL_5TA"].copy()
    df.loc[mascara_swap_t2_t3, "T2_FINAL_5TA"] = df.loc[mascara_swap_t2_t3, "T3_FINAL_5TA"]
    df.loc[mascara_swap_t2_t3, "T3_FINAL_5TA"] = t2_swap_values

    mascara_swap_doble = mascara_swap_t1_t2 & mascara_swap_t2_t3
    df.loc[mascara_swap_t1_t2 & ~mascara_swap_t2_t3, "TIEMPO_ACCION_CIERRE"] = "swap_t1_t2_caso_completo"
    df.loc[mascara_swap_t1_t2 & ~mascara_swap_t2_t3, "TIEMPO_MOTIVO_CIERRE"] = "t1_y_t2_intercambiados_por_inversion_corta_en_registro_completo"
    df.loc[mascara_swap_t2_t3 & ~mascara_swap_t1_t2, "TIEMPO_ACCION_CIERRE"] = "swap_t2_t3_caso_completo"
    df.loc[mascara_swap_t2_t3 & ~mascara_swap_t1_t2, "TIEMPO_MOTIVO_CIERRE"] = "t2_y_t3_intercambiados_por_inversion_corta_en_registro_completo"
    df.loc[mascara_swap_doble, "TIEMPO_ACCION_CIERRE"] = "swap_tiempos_caso_completo"
    df.loc[mascara_swap_doble, "TIEMPO_MOTIVO_CIERRE"] = "t1_t2_y_t2_t3_intercambiados_por_inversiones_cortas_en_registro_completo"
    return refresh_final_time_outputs(df)


def _safe_text(value) -> str:
    if pd.isna(value):
        return "s/d"
    return str(value)


def limpiar_placas_peaje(sub_df: pd.DataFrame) -> pd.DataFrame:
    result = sub_df.copy()
    original = result["PLACA"].fillna("").astype(str)
    original_upper = original.str.upper()
    trimmed_upper = original_upper.str.strip()
    normalized = trimmed_upper.str.replace(r"[^A-Z0-9]", "", regex=True)
    pattern = normalized.map(patron_alfanumerico)

    has_space_hyphen_or_symbol = original_upper.str.contains(r"[\s]|-|[^A-Z0-9\s-]", regex=True)
    invalid_length = ~normalized.str.len().isin(EXPECTED_PLATE_LENGTHS)
    uncommon_pattern = ~pattern.isin(COMMON_PATTERNS)
    placeholder_plate = normalized.isin(PLACEHOLDER_PLATES)
    suspicious_mask = has_space_hyphen_or_symbol | invalid_length | uncommon_pattern | placeholder_plate

    result["PLACA_NORMALIZADA"] = normalized
    result["PLACA_SUGERIDA"] = pd.Series(pd.NA, index=result.index, dtype="string")
    result["PLACA_FINAL"] = normalized
    result["PLACA_ESTADO"] = "ok"
    result["PLACA_MOTIVO"] = "sin observaciones"

    same_peaje_frequencies = normalized.value_counts()

    for idx in result.index:
        reasons = []
        plate = normalized.at[idx]
        plate_pattern = pattern.at[idx]
        current_state = "ok"
        final_plate = plate
        suggested_plate = pd.NA

        if has_space_hyphen_or_symbol.at[idx]:
            reasons.append("contiene espacios, guiones o simbolos")
        if invalid_length.at[idx]:
            reasons.append("longitud fuera de 6 caracteres")
        if uncommon_pattern.at[idx]:
            reasons.append(f"patron poco frecuente: {plate_pattern}")
        if placeholder_plate.at[idx]:
            reasons.append("placeholder o ejemplo")

        if suspicious_mask.at[idx]:
            current_state = "sospechosa"
            candidates = [
                candidate
                for candidate in generar_candidatas_confusion(plate)
                if candidate in same_peaje_frequencies.index
            ]
            if len(candidates) == 1:
                suggested_plate = candidates[0]
                reasons.append(f"sugerencia por confusion visual: {suggested_plate}")
                if same_peaje_frequencies.get(suggested_plate, 0) > same_peaje_frequencies.get(plate, 0):
                    current_state = "autocorregida_alta_confianza"
                    final_plate = suggested_plate
                    reasons.append(f"autocorreccion aplicada hacia {suggested_plate}")
        elif original.at[idx] != trimmed_upper.at[idx]:
            current_state = "normalizada"
            reasons = ["normalizacion segura de mayusculas o espacios externos"]

        result.at[idx, "PLACA_SUGERIDA"] = suggested_plate
        result.at[idx, "PLACA_FINAL"] = final_plate
        result.at[idx, "PLACA_ESTADO"] = current_state
        result.at[idx, "PLACA_MOTIVO"] = "; ".join(reasons) if reasons else "sin observaciones"

    return result


def distancia_levenshtein(a: str, b: str) -> int:
    if a == b:
        return 0
    prev = list(range(len(b) + 1))
    for i, char_a in enumerate(a, start=1):
        current = [i]
        for j, char_b in enumerate(b, start=1):
            current.append(
                min(
                    current[j - 1] + 1,
                    prev[j] + 1,
                    prev[j - 1] + (char_a != char_b),
                )
            )
        prev = current
    return prev[-1]


def tipo_operacion_lista(origen: str, candidata: str) -> str:
    if len(origen) == len(candidata):
        mismatches = [(a, b) for a, b in zip(origen, candidata) if a != b]
        if len(mismatches) == 1:
            old_char, new_char = mismatches[0]
            if new_char in CONFUSION_LOOKUP.get(old_char, set()) or old_char in CONFUSION_LOOKUP.get(new_char, set()):
                return "sustitucion_confusion"
            return "sustitucion_1_char"
    if len(origen) == len(candidata) + 1:
        for idx in range(len(origen)):
            if origen[:idx] + origen[idx + 1 :] == candidata:
                return "quitar_1_char"
    if len(origen) + 1 == len(candidata):
        for idx in range(len(candidata)):
            if candidata[:idx] + candidata[idx + 1 :] == origen:
                return "agregar_1_char"
    return "otro"


def buscar_coincidencias_lista(
    plate: str,
    placas_unicas_base: list[str],
    frecuencia_placas_base: dict[str, int],
) -> tuple[list[dict], list[dict]]:
    tipos_fuertes = {"sustitucion_confusion", "quitar_1_char", "agregar_1_char", "sustitucion_1_char"}
    prioridad_tipo = {
        "sustitucion_confusion": 0,
        "quitar_1_char": 1,
        "agregar_1_char": 1,
        "sustitucion_1_char": 2,
        "otro": 9,
    }
    coincidencias_fuertes = []
    coincidencias_debiles = []
    for candidata in placas_unicas_base:
        if candidata == plate or abs(len(candidata) - len(plate)) > 1:
            continue
        if distancia_levenshtein(plate, candidata) != 1:
            continue
        tipo_operacion = tipo_operacion_lista(plate, candidata)
        patron_candidata = patron_alfanumerico(candidata)
        item = {
            "candidata": candidata,
            "tipo_operacion": tipo_operacion,
            "frecuencia_candidata": int(frecuencia_placas_base.get(candidata, 0)),
            "patron_candidata": patron_candidata,
            "patron_comun": patron_candidata in COMMON_PATTERNS,
        }
        coincidencias_debiles.append(item)
        if item["patron_comun"] and tipo_operacion in tipos_fuertes:
            coincidencias_fuertes.append(item)
    sort_key = lambda item: (
        prioridad_tipo.get(item["tipo_operacion"], 9),
        -item["frecuencia_candidata"],
        item["candidata"],
    )
    coincidencias_fuertes = sorted(coincidencias_fuertes, key=sort_key)
    coincidencias_debiles = sorted(coincidencias_debiles, key=sort_key)
    return coincidencias_fuertes, coincidencias_debiles


def run_plate_cleaning(
    df_in: pd.DataFrame,
    config: dict,
    manual_rules_df: pd.DataFrame,
) -> dict[str, pd.DataFrame]:
    df_peajes = {
        peaje: limpiar_placas_peaje(sub_df)
        for peaje, sub_df in df_in.groupby("PEAJE", dropna=False)
    }
    df = pd.concat(df_peajes.values(), axis=0).sort_index()

    df["REVISION_GRUPO_BLOQUE"] = pd.NA
    df["REV_CANDIDATA_SIN_SUFIJO_X"] = pd.NA
    df["REV_LISTA_CANDIDATA_TOP"] = pd.NA
    df["PLACA_AJUSTE_MANUAL"] = pd.NA
    df["PLACA_ACCION_FINAL"] = "sin_cambio"
    df["PLACA_FINAL_DECIDIDA"] = df["PLACA_FINAL"]
    df["PLACA_EXCLUIR_ANALISIS"] = False

    df_revision_placas = df[df["PLACA_ESTADO"] == "sospechosa"].copy()
    if df_revision_placas.empty:
        if config["aplicar_reglas_manuales"]:
            df = apply_manual_rules_to_df(df, manual_rules_df)
        return {
            "df": df,
            "df_trabajo": df[~df["PLACA_EXCLUIR_ANALISIS"]].copy(),
            "df_eliminados": df[df["PLACA_EXCLUIR_ANALISIS"]].copy(),
            "df_revision_placas": df_revision_placas,
            "df_bloques_decision": pd.DataFrame(),
            "resumen_acciones_placa": df["PLACA_ACCION_FINAL"].value_counts(dropna=False).rename_axis("PLACA_ACCION_FINAL").to_frame("filas"),
        }

    df_revision_placas["REV_PATRON"] = df_revision_placas["PLACA_NORMALIZADA"].map(patron_alfanumerico)
    df_revision_placas["REV_LONGITUD"] = df_revision_placas["PLACA_NORMALIZADA"].str.len()
    df_revision_placas["REV_FLAG_SIMBOLOS"] = df_revision_placas["PLACA"].astype(str).str.upper().str.contains(r"[\s]|-|[^A-Z0-9\s-]", regex=True)
    df_revision_placas["REV_FLAG_LONGITUD"] = ~df_revision_placas["REV_LONGITUD"].isin(EXPECTED_PLATE_LENGTHS)
    df_revision_placas["REV_FLAG_PLACEHOLDER"] = df_revision_placas["PLACA_NORMALIZADA"].isin(PLACEHOLDER_PLATES)
    df_revision_placas["REV_FLAG_PATRON"] = ~df_revision_placas["REV_PATRON"].isin(COMMON_PATTERNS)

    frecuencia_global_placa = df["PLACA_NORMALIZADA"].value_counts()
    frecuencia_mismo_peaje = df.groupby(["PEAJE", "PLACA_NORMALIZADA"]).size()
    frecuencia_patron_base = df["PLACA_NORMALIZADA"].map(patron_alfanumerico).value_counts()

    df_revision_placas["REV_COINCIDENCIAS_MISMO_PEAJE"] = df_revision_placas.apply(
        lambda row: int(frecuencia_mismo_peaje.get((row["PEAJE"], row["PLACA_NORMALIZADA"]), 0)) - 1,
        axis=1,
    )
    df_revision_placas["REV_COINCIDENCIAS_BASE"] = df_revision_placas.apply(
        lambda row: int(frecuencia_global_placa.get(row["PLACA_NORMALIZADA"], 0))
        - int(frecuencia_mismo_peaje.get((row["PEAJE"], row["PLACA_NORMALIZADA"]), 0)),
        axis=1,
    )
    df_revision_placas["REV_FRECUENCIA_PATRON_BASE"] = df_revision_placas["REV_PATRON"].map(frecuencia_patron_base)

    candidatas_confusion = []
    candidatas_confusion_mismo_peaje = []
    candidatas_confusion_base = []
    for _, row in df_revision_placas.iterrows():
        candidates = []
        for candidate in generar_candidatas_confusion(row["PLACA_NORMALIZADA"]):
            global_count = int(frecuencia_global_placa.get(candidate, 0))
            if global_count:
                same_peaje_count = int(frecuencia_mismo_peaje.get((row["PEAJE"], candidate), 0))
                candidates.append((candidate, same_peaje_count, global_count))
        candidates = sorted(candidates, key=lambda item: (-item[1], -item[2], item[0]))
        if len(candidates) == 1:
            candidatas_confusion.append(candidates[0][0])
            candidatas_confusion_mismo_peaje.append(candidates[0][1])
            candidatas_confusion_base.append(candidates[0][2])
        else:
            candidatas_confusion.append(pd.NA)
            candidatas_confusion_mismo_peaje.append(pd.NA)
            candidatas_confusion_base.append(pd.NA)

    df_revision_placas["REV_CANDIDATA_CONFUSION"] = candidatas_confusion
    df_revision_placas["REV_CANDIDATA_CONF_MISMO_PEAJE"] = candidatas_confusion_mismo_peaje
    df_revision_placas["REV_CANDIDATA_CONF_BASE"] = candidatas_confusion_base
    df_revision_placas["REV_CANDIDATA_SIN_SUFIJO_X"] = df_revision_placas["PLACA_NORMALIZADA"].where(
        df_revision_placas["PLACA_NORMALIZADA"].astype(str).str.endswith("X")
        & (df_revision_placas["REV_LONGITUD"] > max(EXPECTED_PLATE_LENGTHS)),
        pd.NA,
    ).str[:-1]
    df_revision_placas["REV_SUFIJO_X_MISMO_PEAJE"] = df_revision_placas.apply(
        lambda row: int(frecuencia_mismo_peaje.get((row["PEAJE"], row["REV_CANDIDATA_SIN_SUFIJO_X"]), 0))
        if pd.notna(row["REV_CANDIDATA_SIN_SUFIJO_X"])
        else 0,
        axis=1,
    )
    df_revision_placas["REV_SUFIJO_X_BASE"] = df_revision_placas["REV_CANDIDATA_SIN_SUFIJO_X"].map(
        lambda plate: int(frecuencia_global_placa.get(plate, 0)) if pd.notna(plate) else 0
    )

    def clasificar_revision(row: pd.Series) -> pd.Series:
        normalized = row["PLACA_NORMALIZADA"]
        pattern = row["REV_PATRON"]
        pattern_frequency = int(row["REV_FRECUENCIA_PATRON_BASE"])
        length = int(row["REV_LONGITUD"])
        exact_same_peaje = int(row["REV_COINCIDENCIAS_MISMO_PEAJE"])
        exact_base = int(row["REV_COINCIDENCIAS_BASE"])
        candidate = row["REV_CANDIDATA_CONFUSION"]
        candidate_same_peaje = row["REV_CANDIDATA_CONF_MISMO_PEAJE"]
        candidate_base = row["REV_CANDIDATA_CONF_BASE"]
        suffix_x_candidate = row["REV_CANDIDATA_SIN_SUFIJO_X"]
        suffix_x_same_peaje = int(row["REV_SUFIJO_X_MISMO_PEAJE"])
        suffix_x_base = int(row["REV_SUFIJO_X_BASE"])
        has_symbols = bool(row["REV_FLAG_SIMBOLOS"])
        has_length_issue = bool(row["REV_FLAG_LONGITUD"])
        is_placeholder = bool(row["REV_FLAG_PLACEHOLDER"])
        has_suffix_x = pd.notna(suffix_x_candidate)
        suffix_x_pattern = patron_alfanumerico(str(suffix_x_candidate)) if has_suffix_x else ""

        if is_placeholder:
            return pd.Series({"REVISION_GRUPO": "placeholder_o_ejemplo", "REVISION_SUBGRUPO": normalized})
        if has_suffix_x and suffix_x_same_peaje > 0:
            return pd.Series({"REVISION_GRUPO": "sufijo_x_final_con_respaldo_local", "REVISION_SUBGRUPO": "recorte_confirmado_en_mismo_peaje"})
        if has_suffix_x and suffix_x_base > 0:
            return pd.Series({"REVISION_GRUPO": "sufijo_x_final_con_respaldo_global", "REVISION_SUBGRUPO": "recorte_confirmado_en_base"})
        if has_suffix_x and suffix_x_pattern in {"AAA999", "A9A999"}:
            return pd.Series({"REVISION_GRUPO": "sufijo_x_final_patron_peru", "REVISION_SUBGRUPO": suffix_x_pattern})
        if has_suffix_x:
            return pd.Series({"REVISION_GRUPO": "sufijo_x_final_sin_respaldo", "REVISION_SUBGRUPO": "recorte_sin_confirmacion_en_base"})
        if has_symbols and not has_length_issue and exact_same_peaje > 0:
            return pd.Series({"REVISION_GRUPO": "ruido_tipografico_con_respaldo_local", "REVISION_SUBGRUPO": "normalizacion_confirmada_en_mismo_peaje"})
        if has_symbols and not has_length_issue and exact_base > 0:
            return pd.Series({"REVISION_GRUPO": "ruido_tipografico_con_respaldo_global", "REVISION_SUBGRUPO": "normalizacion_confirmada_en_base"})
        if pd.notna(candidate):
            return pd.Series({"REVISION_GRUPO": "posible_confusion_visual", "REVISION_SUBGRUPO": pattern})
        if has_symbols and not has_length_issue:
            return pd.Series({"REVISION_GRUPO": "ruido_tipografico_sin_respaldo", "REVISION_SUBGRUPO": pattern})
        if has_length_issue:
            return pd.Series({"REVISION_GRUPO": "longitud_atipica", "REVISION_SUBGRUPO": f"longitud_{length}_patron_{pattern}"})
        if pattern_frequency >= 2:
            return pd.Series({"REVISION_GRUPO": "patron_atipico_recurrente", "REVISION_SUBGRUPO": pattern})
        return pd.Series({"REVISION_GRUPO": "patron_atipico_aislado", "REVISION_SUBGRUPO": pattern})

    df_revision_placas = pd.concat([df_revision_placas, df_revision_placas.apply(clasificar_revision, axis=1)], axis=1)

    target_blocks = {"longitud_atipica", "patron_atipico_aislado", "patron_atipico_recurrente"}
    placas_unicas_base = sorted(df["PLACA_NORMALIZADA"].dropna().astype(str).unique())
    frecuencia_placas_base = df["PLACA_NORMALIZADA"].value_counts().to_dict()

    estado_lista = []
    bloque_decision = []
    candidata_top = []
    for _, row in df_revision_placas.iterrows():
        group = row["REVISION_GRUPO"]
        plate = row["PLACA_NORMALIZADA"]
        if group not in target_blocks:
            estado_lista.append("no_aplica")
            bloque_decision.append(group)
            candidata_top.append(pd.NA)
            continue
        fuertes, debiles = buscar_coincidencias_lista(plate, placas_unicas_base, frecuencia_placas_base)
        if len(fuertes) == 1:
            estado = "coincidencia_unica"
            top = fuertes[0]
        elif len(fuertes) > 1:
            estado = "coincidencia_multiple"
            top = fuertes[0]
        elif debiles:
            estado = "sin_coincidencia_fuerte"
            top = debiles[0]
        else:
            estado = "sin_coincidencia"
            top = None
        estado_lista.append(estado)
        bloque_decision.append(f"{group}__{estado}")
        candidata_top.append(top["candidata"] if top else pd.NA)

    df_revision_placas["REV_LISTA_COINCIDENCIA_ESTADO"] = estado_lista
    df_revision_placas["REVISION_BLOQUE_DECISION"] = bloque_decision
    df_revision_placas["REV_LISTA_CANDIDATA_TOP"] = candidata_top

    conteo_peajes_por_bloque = (
        df_revision_placas.groupby(["REVISION_BLOQUE_DECISION", "PEAJE"]).size().rename("filas").reset_index()
    )
    resumen_peajes_por_bloque = (
        conteo_peajes_por_bloque.groupby("REVISION_BLOQUE_DECISION")
        .apply(lambda sub: " | ".join(f"{row['PEAJE']}: {row['filas']}" for _, row in sub.iterrows()), include_groups=False)
        .rename("peajes")
    )
    df_bloques_decision = (
        df_revision_placas.groupby("REVISION_BLOQUE_DECISION")
        .agg(
            grupo_base=("REVISION_GRUPO", "first"),
            filas=("PLACA", "size"),
            ejemplos=("PLACA", lambda s: ", ".join(s.astype(str).head(6))),
            candidatas=("REV_LISTA_CANDIDATA_TOP", lambda s: ", ".join(s.dropna().astype(str).head(6))),
        )
        .join(resumen_peajes_por_bloque)
        .reset_index()
    )

    decisiones_por_bloque = {bloque: "mantener_observada" for bloque in df_bloques_decision["REVISION_BLOQUE_DECISION"]}
    for bloque in decisiones_por_bloque:
        base = bloque.split("__", 1)[0]
        estado = bloque.split("__", 1)[1] if "__" in bloque else "base"
        if base in {"ruido_tipografico_con_respaldo_local", "ruido_tipografico_con_respaldo_global"} and config["aplicar_ruido_con_respaldo"]:
            decisiones_por_bloque[bloque] = "corregir_a_normalizada"
        elif base == "ruido_tipografico_sin_respaldo" and config["aplicar_ruido_sin_respaldo"]:
            decisiones_por_bloque[bloque] = "corregir_a_normalizada"
        elif base == "posible_confusion_visual" and config["aplicar_confusion_visual"]:
            decisiones_por_bloque[bloque] = "corregir_a_sugerida"
        elif base in {
            "sufijo_x_final_patron_peru",
            "sufijo_x_final_con_respaldo_local",
            "sufijo_x_final_con_respaldo_global",
            "sufijo_x_final_sin_respaldo",
        } and config["aplicar_recorte_sufijo_x"]:
            decisiones_por_bloque[bloque] = "corregir_a_recorte_sufijo_x"
        elif base == "placeholder_o_ejemplo" and config["aplicar_exclusion_placeholders"]:
            decisiones_por_bloque[bloque] = "excluir_analisis_placa"
        elif estado == "coincidencia_unica" and config["aplicar_coincidencia_unica_lista"]:
            decisiones_por_bloque[bloque] = "corregir_a_coincidencia_lista"

    df_bloques_decision["ACCION_ELEGIDA"] = df_bloques_decision["REVISION_BLOQUE_DECISION"].map(decisiones_por_bloque)
    df_revision_placas["PLACA_ACCION_BLOQUE"] = df_revision_placas["REVISION_BLOQUE_DECISION"].map(decisiones_por_bloque)

    df.loc[df_revision_placas.index, "REVISION_GRUPO_BLOQUE"] = df_revision_placas["REVISION_BLOQUE_DECISION"]
    df.loc[df_revision_placas.index, "REV_CANDIDATA_SIN_SUFIJO_X"] = df_revision_placas["REV_CANDIDATA_SIN_SUFIJO_X"]
    df.loc[df_revision_placas.index, "REV_LISTA_CANDIDATA_TOP"] = df_revision_placas["REV_LISTA_CANDIDATA_TOP"]

    for idx, row in df.iterrows():
        action = df_revision_placas["PLACA_ACCION_BLOQUE"].get(idx, "sin_cambio")
        normalized = str(row["PLACA_NORMALIZADA"])
        final_plate = row["PLACA_FINAL"]
        exclude_plate = False

        if action == "corregir_a_normalizada":
            final_plate = normalized
        elif action == "corregir_a_sugerida" and pd.notna(row["PLACA_SUGERIDA"]):
            final_plate = row["PLACA_SUGERIDA"]
        elif action == "corregir_a_coincidencia_lista" and pd.notna(row["REV_LISTA_CANDIDATA_TOP"]):
            final_plate = row["REV_LISTA_CANDIDATA_TOP"]
        elif action == "corregir_a_recorte_sufijo_x" and pd.notna(row["REV_CANDIDATA_SIN_SUFIJO_X"]):
            final_plate = row["REV_CANDIDATA_SIN_SUFIJO_X"]
        elif action == "excluir_analisis_placa":
            exclude_plate = True
        else:
            action = "sin_cambio" if idx not in df_revision_placas.index else "mantener_observada"

        df.at[idx, "PLACA_ACCION_FINAL"] = action
        df.at[idx, "PLACA_FINAL_DECIDIDA"] = final_plate
        df.at[idx, "PLACA_EXCLUIR_ANALISIS"] = exclude_plate

    if config["aplicar_reglas_manuales"]:
        df = apply_manual_rules_to_df(df, manual_rules_df)

    if config.get("aplicar_consolidacion_duplicados_cercanos", False):
        df = consolidate_near_duplicate_plate_rows(df, DEDUPLICACION_DUPLICADO_CERCANO_SEGUNDOS)

    df_eliminados = df[df["PLACA_EXCLUIR_ANALISIS"]].copy()
    df_trabajo = df[~df["PLACA_EXCLUIR_ANALISIS"]].copy()
    resumen_acciones_placa = (
        df["PLACA_ACCION_FINAL"].value_counts(dropna=False).rename_axis("PLACA_ACCION_FINAL").to_frame("filas")
    )

    return {
        "df": df,
        "df_trabajo": df_trabajo,
        "df_eliminados": df_eliminados,
        "df_revision_placas": df_revision_placas,
        "df_bloques_decision": df_bloques_decision,
        "resumen_acciones_placa": resumen_acciones_placa,
    }


def describir_faltantes_trio(row: pd.Series) -> str:
    faltantes = []
    if pd.isna(row["T1"]):
        faltantes.append("T1")
    if pd.isna(row["T2"]):
        faltantes.append("T2")
    if pd.isna(row["T3"]):
        faltantes.append("T3")
    return ", ".join(faltantes) if faltantes else "ninguno"


def evaluar_bordes_caseta(grupo: pd.DataFrame) -> pd.DataFrame:
    grupo = grupo.sort_values(
        ["TIEMPO_REFERENCIA", "T1", "T2", "T3", "_ORDEN_FILA"],
        na_position="last",
    ).copy()
    grupo["POSICION_FLUJO"] = range(1, len(grupo) + 1)
    grupo["VENTANA_COMPLETA_IDENTIFICADA"] = grupo["TIEMPOS_COMPLETOS"].any()
    grupo["BORDE_CASETA_ACCION"] = "mantener_para_tiempos"
    grupo["BORDE_CASETA_MOTIVO"] = "dentro_ventana_o_registro_completo"
    grupo["BORDE_CASETA_ELIMINAR"] = False
    grupo["BORDE_ES_INICIO"] = False
    grupo["BORDE_ES_CIERRE"] = False

    mascara_completa = grupo["TIEMPOS_COMPLETOS"]
    if not mascara_completa.any():
        grupo["BORDE_CASETA_ACCION"] = "sin_referencia_completa"
        grupo["BORDE_CASETA_MOTIVO"] = (
            "El flujo no tiene ningun vehiculo con T1, T2 y T3 completos; "
            "no se fija apertura ni cierre de manera automatica."
        )
        return grupo

    apertura_pos = int(grupo.loc[mascara_completa, "POSICION_FLUJO"].min())
    cierre_pos = int(grupo.loc[mascara_completa, "POSICION_FLUJO"].max())

    eliminar_inicio = (grupo["POSICION_FLUJO"] < apertura_pos) & (~grupo["TIEMPOS_COMPLETOS"])
    eliminar_cierre = (grupo["POSICION_FLUJO"] > cierre_pos) & (~grupo["TIEMPOS_COMPLETOS"])

    grupo.loc[eliminar_inicio, "BORDE_CASETA_ACCION"] = "eliminar_inicio_caseta"
    grupo.loc[eliminar_inicio, "BORDE_CASETA_MOTIVO"] = (
        "Registro incompleto antes del primer vehiculo con T1, T2 y T3 completos del flujo."
    )
    grupo.loc[eliminar_inicio, "BORDE_CASETA_ELIMINAR"] = True
    grupo.loc[eliminar_inicio, "BORDE_ES_INICIO"] = True

    grupo.loc[eliminar_cierre, "BORDE_CASETA_ACCION"] = "eliminar_cierre_caseta"
    grupo.loc[eliminar_cierre, "BORDE_CASETA_MOTIVO"] = (
        "Registro incompleto despues del ultimo vehiculo con T1, T2 y T3 completos del flujo."
    )
    grupo.loc[eliminar_cierre, "BORDE_CASETA_ELIMINAR"] = True
    grupo.loc[eliminar_cierre, "BORDE_ES_CIERRE"] = True

    return grupo


def patron_faltantes_tiempo(row: pd.Series, columnas: list[str]) -> str:
    faltantes = [col for col in columnas if pd.isna(row[col])]
    return "completo" if not faltantes else "+".join(faltantes)


def columnas_imputadas_tiempo(row: pd.Series, columnas_origen: list[str], columnas_finales: list[str]) -> str:
    imputadas = []
    for origen, final in zip(columnas_origen, columnas_finales):
        if pd.isna(row[origen]) and pd.notna(row[final]):
            imputadas.append(origen)
    return "ninguna" if not imputadas else "+".join(imputadas)


def interpolar_tiempos_grupo(grupo: pd.DataFrame) -> pd.DataFrame:
    grupo = grupo.sort_values("POSICION_FLUJO").copy()
    posicion = grupo["POSICION_FLUJO"].astype(float)
    for alias in TIME_ALIASES:
        segundos = grupo[alias].dt.total_seconds()
        observados = segundos.notna()
        prev_segundos = segundos.where(observados).ffill()
        next_segundos = segundos.where(observados).bfill()
        prev_posicion = posicion.where(observados).ffill()
        next_posicion = posicion.where(observados).bfill()
        denominador = next_posicion - prev_posicion
        proporcion = (posicion - prev_posicion) / denominador
        candidata_segundos = prev_segundos + (next_segundos - prev_segundos) * proporcion
        mascara_candidata = (
            grupo[alias].isna()
            & prev_segundos.notna()
            & next_segundos.notna()
            & prev_posicion.notna()
            & next_posicion.notna()
            & (denominador > 0)
        )
        grupo[f"{alias}_INTERP"] = candidata_segundos.where(mascara_candidata).map(segundos_a_timedelta)
    return grupo


def run_time_cleaning(df_trabajo: pd.DataFrame, config: dict) -> dict[str, pd.DataFrame]:
    df_tiempos_base = df_trabajo.copy()
    df_tiempos_base["PLACA_FINAL"] = df_tiempos_base["PLACA_FINAL_DECIDIDA"]
    df_tiempos_base = df_tiempos_base[
        [
            "PEAJE",
            "CASETA",
            "SENTIDO",
            "FECHA",
            "VEHICULO",
            "PLACA_FINAL",
            "PLACA_ACCION_FINAL",
            "LLEGADA COLA",
            "LLEGADA CASETA",
            "SALIDA CASETA",
            "T. TEC",
            "T. CASETA",
            "_ORDEN_FILA",
        ]
    ].copy()
    df_tiempos_base["FECHA_DIA"] = pd.to_datetime(df_tiempos_base["FECHA"], errors="coerce").dt.normalize()
    for columna_origen, alias in TIME_COLUMNS.items():
        df_tiempos_base[alias] = df_tiempos_base[columna_origen].map(normalizar_hora)
        df_tiempos_base[f"{alias}_TXT"] = df_tiempos_base[alias].map(formatear_hora)
    df_tiempos_base["TIEMPOS_COMPLETOS"] = df_tiempos_base[["T1", "T2", "T3"]].notna().all(axis=1)
    df_tiempos_base["TIEMPO_REFERENCIA"] = (
        df_tiempos_base["T1"].combine_first(df_tiempos_base["T2"]).combine_first(df_tiempos_base["T3"])
    )

    if config["eliminar_bordes_caseta"]:
        grupos_evaluados = [
            evaluar_bordes_caseta(sub_df)
            for _, sub_df in df_tiempos_base.groupby(TIME_GROUP_KEYS, sort=True, dropna=False)
        ]
        df_tiempos_bordes = pd.concat(grupos_evaluados, axis=0).sort_index()
    else:
        df_tiempos_bordes = df_tiempos_base.copy()
        df_tiempos_bordes["POSICION_FLUJO"] = df_tiempos_bordes.groupby(TIME_GROUP_KEYS, dropna=False).cumcount() + 1
        df_tiempos_bordes["VENTANA_COMPLETA_IDENTIFICADA"] = df_tiempos_bordes.groupby(TIME_GROUP_KEYS, dropna=False)["TIEMPOS_COMPLETOS"].transform("any")
        df_tiempos_bordes["BORDE_CASETA_ACCION"] = "sin_cambio"
        df_tiempos_bordes["BORDE_CASETA_MOTIVO"] = "etapa_bordes_desactivada"
        df_tiempos_bordes["BORDE_CASETA_ELIMINAR"] = False
        df_tiempos_bordes["BORDE_ES_INICIO"] = False
        df_tiempos_bordes["BORDE_ES_CIERRE"] = False

    # Los registros de borde se mantienen temporalmente en juego para darles una ultima oportunidad
    # de recuperacion antes de excluirlos definitivamente.
    df_tiempos_bordes["PENDIENTE_TIEMPOS_INTERNO"] = ~df_tiempos_bordes["TIEMPOS_COMPLETOS"]
    df_tiempos_trabajo = df_tiempos_bordes.copy()

    grupos_interpolados = [
        interpolar_tiempos_grupo(sub_df)
        for _, sub_df in df_tiempos_trabajo.groupby(TIME_GROUP_KEYS, sort=True, dropna=False)
    ]
    df_inter = pd.concat(grupos_interpolados, axis=0).sort_index()
    df_inter["PATRON_FALTANTES_ORIGINAL"] = df_inter.apply(lambda row: patron_faltantes_tiempo(row, TIME_ALIASES), axis=1)
    df_inter["ELEGIBLE_INTERPOLACION"] = (
        df_inter["PENDIENTE_TIEMPOS_INTERNO"]
        & df_inter["VENTANA_COMPLETA_IDENTIFICADA"]
        & df_inter["TIEMPO_REFERENCIA"].notna()
    )

    mascara_candidata_completa = pd.Series(True, index=df_inter.index)
    for alias in TIME_ALIASES:
        mascara_candidata_completa &= df_inter[alias].notna() | df_inter[f"{alias}_INTERP"].notna()
        df_inter[f"{alias}_PROPUESTA"] = df_inter[alias].combine_first(df_inter[f"{alias}_INTERP"])
    df_inter["INTERPOLACION_COMPLETA"] = mascara_candidata_completa
    mascara_propuesta_completa = df_inter[[f"{alias}_PROPUESTA" for alias in TIME_ALIASES]].notna().all(axis=1)
    df_inter["INTERPOLACION_ORDEN_VALIDO"] = (
        mascara_propuesta_completa
        & (df_inter["T1_PROPUESTA"] <= df_inter["T2_PROPUESTA"])
        & (df_inter["T2_PROPUESTA"] <= df_inter["T3_PROPUESTA"])
    )
    mascara_incompleta = ~df_inter["TIEMPOS_COMPLETOS"]
    mascara_imputar = (
        config["aplicar_interpolacion"]
        & mascara_incompleta
        & df_inter["ELEGIBLE_INTERPOLACION"]
        & df_inter["INTERPOLACION_COMPLETA"]
        & df_inter["INTERPOLACION_ORDEN_VALIDO"]
    )
    for alias in TIME_ALIASES:
        df_inter[f"{alias}_FINAL"] = df_inter[alias]
        df_inter.loc[mascara_imputar, f"{alias}_FINAL"] = df_inter.loc[mascara_imputar, f"{alias}_PROPUESTA"]
    df_inter["TIEMPO_ACCION_INTERPOLACION"] = "sin_cambio"
    df_inter["TIEMPO_MOTIVO_INTERPOLACION"] = "registro_ya_completo"
    df_inter.loc[mascara_incompleta, "TIEMPO_ACCION_INTERPOLACION"] = "mantener_pendiente"
    df_inter.loc[mascara_incompleta, "TIEMPO_MOTIVO_INTERPOLACION"] = "pendiente_por_evaluar"
    df_inter.loc[
        mascara_incompleta & ~df_inter["VENTANA_COMPLETA_IDENTIFICADA"],
        "TIEMPO_MOTIVO_INTERPOLACION",
    ] = "flujo_sin_referencia_completa"
    df_inter.loc[
        mascara_incompleta & df_inter["VENTANA_COMPLETA_IDENTIFICADA"] & df_inter["TIEMPO_REFERENCIA"].isna(),
        "TIEMPO_MOTIVO_INTERPOLACION",
    ] = "fila_sin_tiempo_referencia"
    df_inter.loc[
        mascara_incompleta & df_inter["ELEGIBLE_INTERPOLACION"] & ~df_inter["INTERPOLACION_COMPLETA"],
        "TIEMPO_MOTIVO_INTERPOLACION",
    ] = "faltan_anclas_para_interpolar"
    df_inter.loc[
        mascara_incompleta
        & df_inter["ELEGIBLE_INTERPOLACION"]
        & df_inter["INTERPOLACION_COMPLETA"]
        & ~df_inter["INTERPOLACION_ORDEN_VALIDO"],
        "TIEMPO_MOTIVO_INTERPOLACION",
    ] = "interpolacion_rompe_orden"
    df_inter.loc[mascara_imputar, "TIEMPO_ACCION_INTERPOLACION"] = "imputar_interpolacion"
    df_inter.loc[mascara_imputar, "TIEMPO_MOTIVO_INTERPOLACION"] = "interpolacion_interna_con_ancla_previa_y_posterior"

    base_medianas_locales = df_inter[df_inter[[f"{alias}_FINAL" for alias in TIME_ALIASES]].notna().all(axis=1)].copy()
    base_medianas_locales["T_COLA_FINAL"] = base_medianas_locales["T2_FINAL"] - base_medianas_locales["T1_FINAL"]
    base_medianas_locales["T_CASETA_FINAL"] = base_medianas_locales["T3_FINAL"] - base_medianas_locales["T2_FINAL"]
    base_medianas_locales = base_medianas_locales[
        base_medianas_locales["T_COLA_FINAL"].notna()
        & base_medianas_locales["T_CASETA_FINAL"].notna()
        & (base_medianas_locales["T_COLA_FINAL"] >= pd.Timedelta(0))
        & (base_medianas_locales["T_CASETA_FINAL"] >= pd.Timedelta(0))
    ].copy()
    medianas_locales_flujo = (
        base_medianas_locales.groupby(TIME_GROUP_KEYS, dropna=False)
        .agg(
            MEDIANA_COLA_FLUJO=("T_COLA_FINAL", "median"),
            MEDIANA_CASETA_FLUJO=("T_CASETA_FINAL", "median"),
        )
        .reset_index()
    )
    df_2da = df_inter.merge(medianas_locales_flujo, on=TIME_GROUP_KEYS, how="left")
    objetivo_segunda_pasada = df_2da["TIEMPO_MOTIVO_INTERPOLACION"] == "interpolacion_rompe_orden"

    def proponer_segunda_pasada(row: pd.Series) -> pd.Series:
        resultado = {
            "SEGUNDA_T1": row["T1_FINAL"],
            "SEGUNDA_T2": row["T2_FINAL"],
            "SEGUNDA_T3": row["T3_FINAL"],
            "SEGUNDA_PASADA_MOTIVO": pd.NA,
        }
        if row["TIEMPO_MOTIVO_INTERPOLACION"] != "interpolacion_rompe_orden":
            resultado["SEGUNDA_PASADA_MOTIVO"] = "no_aplica"
            return pd.Series(resultado)
        mediana_cola = row["MEDIANA_COLA_FLUJO"]
        mediana_caseta = row["MEDIANA_CASETA_FLUJO"]
        if pd.isna(mediana_cola) or pd.isna(mediana_caseta):
            resultado["SEGUNDA_PASADA_MOTIVO"] = "medianas_locales_insuficientes"
            return pd.Series(resultado)
        t1, t2, t3 = row["T1_FINAL"], row["T2_FINAL"], row["T3_FINAL"]
        for _ in range(3):
            cambio = False
            if pd.isna(t2):
                if pd.notna(t1):
                    t2 = t1 + mediana_cola
                    cambio = True
                elif pd.notna(t3):
                    t2 = t3 - mediana_caseta
                    cambio = True
            if pd.isna(t3) and pd.notna(t2):
                t3 = t2 + mediana_caseta
                cambio = True
            if pd.isna(t1) and pd.notna(t2):
                t1 = t2 - mediana_cola
                cambio = True
            if not cambio:
                break
        resultado["SEGUNDA_T1"] = t1
        resultado["SEGUNDA_T2"] = t2
        resultado["SEGUNDA_T3"] = t3
        if pd.isna(t1) or pd.isna(t2) or pd.isna(t3):
            resultado["SEGUNDA_PASADA_MOTIVO"] = "segunda_pasada_incompleta"
        elif not (t1 <= t2 <= t3):
            resultado["SEGUNDA_PASADA_MOTIVO"] = "segunda_pasada_rompe_orden"
        else:
            resultado["SEGUNDA_PASADA_MOTIVO"] = "segunda_pasada_mediana_local"
        return pd.Series(resultado)

    propuestas_segunda = df_2da.apply(proponer_segunda_pasada, axis=1)
    df_2da = pd.concat([df_2da, propuestas_segunda], axis=1)
    mascara_segunda_recuperada = (
        config["aplicar_mediana_local"]
        & objetivo_segunda_pasada
        & df_2da["SEGUNDA_PASADA_MOTIVO"].eq("segunda_pasada_mediana_local")
    )
    for alias in TIME_ALIASES:
        df_2da[f"{alias}_FINAL_2DA"] = df_2da[f"{alias}_FINAL"]
        df_2da.loc[mascara_segunda_recuperada, f"{alias}_FINAL_2DA"] = df_2da.loc[
            mascara_segunda_recuperada,
            f"SEGUNDA_{alias}",
        ]
    df_2da["TIEMPO_ACCION_FINAL"] = df_2da["TIEMPO_ACCION_INTERPOLACION"]
    df_2da["TIEMPO_MOTIVO_FINAL"] = df_2da["TIEMPO_MOTIVO_INTERPOLACION"]
    df_2da.loc[mascara_segunda_recuperada, "TIEMPO_ACCION_FINAL"] = "imputar_mediana_local"
    df_2da.loc[mascara_segunda_recuperada, "TIEMPO_MOTIVO_FINAL"] = "segunda_pasada_mediana_local"
    df_2da.loc[
        objetivo_segunda_pasada & ~mascara_segunda_recuperada,
        "TIEMPO_MOTIVO_FINAL",
    ] = df_2da.loc[objetivo_segunda_pasada & ~mascara_segunda_recuperada, "SEGUNDA_PASADA_MOTIVO"]

    base_donantes = df_2da[df_2da[[f"{alias}_FINAL_2DA" for alias in TIME_ALIASES]].notna().all(axis=1)].copy()
    base_donantes["COLA_DONANTE"] = base_donantes["T2_FINAL_2DA"] - base_donantes["T1_FINAL_2DA"]
    base_donantes["CASETA_DONANTE"] = base_donantes["T3_FINAL_2DA"] - base_donantes["T2_FINAL_2DA"]
    base_donantes = base_donantes[
        base_donantes["COLA_DONANTE"].notna()
        & base_donantes["CASETA_DONANTE"].notna()
        & (base_donantes["COLA_DONANTE"] >= pd.Timedelta(0))
        & (base_donantes["CASETA_DONANTE"] >= pd.Timedelta(0))
    ].copy()
    donantes_caseta = (
        base_donantes.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
        .agg(
            MEDIANA_COLA_DONANTE_CASETA=("COLA_DONANTE", "median"),
            MEDIANA_CASETA_DONANTE_CASETA=("CASETA_DONANTE", "median"),
            REFERENCIAS_DONANTE_CASETA=("PLACA_FINAL", "size"),
        )
        .reset_index()
    )
    donantes_dia = (
        base_donantes.groupby(["PEAJE", "FECHA_DIA", "SENTIDO"], dropna=False)
        .agg(
            MEDIANA_COLA_DONANTE_DIA=("COLA_DONANTE", "median"),
            MEDIANA_CASETA_DONANTE_DIA=("CASETA_DONANTE", "median"),
            REFERENCIAS_DONANTE_DIA=("PLACA_FINAL", "size"),
        )
        .reset_index()
    )
    donantes_peaje_sentido = (
        base_donantes.groupby(["PEAJE", "SENTIDO"], dropna=False)
        .agg(
            MEDIANA_COLA_DONANTE_PEAJE_SENTIDO=("COLA_DONANTE", "median"),
            MEDIANA_CASETA_DONANTE_PEAJE_SENTIDO=("CASETA_DONANTE", "median"),
            REFERENCIAS_DONANTE_PEAJE_SENTIDO=("PLACA_FINAL", "size"),
        )
        .reset_index()
    )
    df_3ra = df_2da.merge(donantes_caseta, on=["PEAJE", "CASETA", "SENTIDO"], how="left")
    df_3ra = df_3ra.merge(donantes_dia, on=["PEAJE", "FECHA_DIA", "SENTIDO"], how="left")
    df_3ra = df_3ra.merge(donantes_peaje_sentido, on=["PEAJE", "SENTIDO"], how="left")

    def elegir_donante_cola(row: pd.Series) -> pd.Series:
        if pd.notna(row["MEDIANA_COLA_DONANTE_CASETA"]):
            return pd.Series({
                "COLA_DONANTE_ELEGIDA": row["MEDIANA_COLA_DONANTE_CASETA"],
                "CASETA_DONANTE_ELEGIDA": row["MEDIANA_CASETA_DONANTE_CASETA"],
                "FUENTE_DONANTE": "misma_caseta_sentido",
            })
        if pd.notna(row["MEDIANA_COLA_DONANTE_DIA"]):
            return pd.Series({
                "COLA_DONANTE_ELEGIDA": row["MEDIANA_COLA_DONANTE_DIA"],
                "CASETA_DONANTE_ELEGIDA": row["MEDIANA_CASETA_DONANTE_DIA"],
                "FUENTE_DONANTE": "mismo_dia_mismo_sentido",
            })
        if pd.notna(row["MEDIANA_COLA_DONANTE_PEAJE_SENTIDO"]):
            return pd.Series({
                "COLA_DONANTE_ELEGIDA": row["MEDIANA_COLA_DONANTE_PEAJE_SENTIDO"],
                "CASETA_DONANTE_ELEGIDA": row["MEDIANA_CASETA_DONANTE_PEAJE_SENTIDO"],
                "FUENTE_DONANTE": "mismo_peaje_mismo_sentido",
            })
        return pd.Series({"COLA_DONANTE_ELEGIDA": pd.NaT, "CASETA_DONANTE_ELEGIDA": pd.NaT, "FUENTE_DONANTE": pd.NA})

    df_3ra = pd.concat([df_3ra, df_3ra.apply(elegir_donante_cola, axis=1)], axis=1)
    objetivo_t1_donante = (
        config["aplicar_donantes"]
        & df_3ra["T1_FINAL_2DA"].isna()
        & df_3ra["T2_FINAL_2DA"].notna()
        & df_3ra["T3_FINAL_2DA"].notna()
    )
    objetivo_t2_t3_donante = (
        config["aplicar_donantes"]
        & df_3ra["T1_FINAL_2DA"].notna()
        & df_3ra["T2_FINAL_2DA"].isna()
        & df_3ra["T3_FINAL_2DA"].isna()
    )
    df_3ra["T1_FINAL_3RA"] = df_3ra["T1_FINAL_2DA"]
    df_3ra["T2_FINAL_3RA"] = df_3ra["T2_FINAL_2DA"]
    df_3ra["T3_FINAL_3RA"] = df_3ra["T3_FINAL_2DA"]
    mascara_t1_recuperada = objetivo_t1_donante & df_3ra["COLA_DONANTE_ELEGIDA"].notna()
    df_3ra.loc[mascara_t1_recuperada, "T1_FINAL_3RA"] = (
        df_3ra.loc[mascara_t1_recuperada, "T2_FINAL_2DA"] - df_3ra.loc[mascara_t1_recuperada, "COLA_DONANTE_ELEGIDA"]
    )

    def proponer_t2_t3_con_donantes(row: pd.Series) -> pd.Series:
        resultado = {"T2_DONANTE_PROPUESTO": pd.NaT, "T3_DONANTE_PROPUESTO": pd.NaT, "RECUPERABLE_T2_T3": False}
        if not (pd.notna(row["T1_FINAL_2DA"]) and pd.isna(row["T2_FINAL_2DA"]) and pd.isna(row["T3_FINAL_2DA"])):
            return pd.Series(resultado)
        if pd.notna(row["COLA_DONANTE_ELEGIDA"]) and pd.notna(row["CASETA_DONANTE_ELEGIDA"]):
            t2 = row["T1_FINAL_2DA"] + row["COLA_DONANTE_ELEGIDA"]
            t3 = t2 + row["CASETA_DONANTE_ELEGIDA"]
            if row["T1_FINAL_2DA"] <= t2 <= t3:
                resultado["T2_DONANTE_PROPUESTO"] = t2
                resultado["T3_DONANTE_PROPUESTO"] = t3
                resultado["RECUPERABLE_T2_T3"] = True
        return pd.Series(resultado)

    df_3ra = pd.concat([df_3ra, df_3ra.apply(proponer_t2_t3_con_donantes, axis=1)], axis=1)
    mascara_t2_t3_recuperada = objetivo_t2_t3_donante & df_3ra["RECUPERABLE_T2_T3"]
    df_3ra.loc[mascara_t2_t3_recuperada, "T2_FINAL_3RA"] = df_3ra.loc[mascara_t2_t3_recuperada, "T2_DONANTE_PROPUESTO"]
    df_3ra.loc[mascara_t2_t3_recuperada, "T3_FINAL_3RA"] = df_3ra.loc[mascara_t2_t3_recuperada, "T3_DONANTE_PROPUESTO"]
    df_3ra["TIEMPO_ACCION_GLOBAL"] = df_3ra["TIEMPO_ACCION_FINAL"]
    df_3ra["TIEMPO_MOTIVO_GLOBAL"] = df_3ra["TIEMPO_MOTIVO_FINAL"]
    df_3ra.loc[mascara_t1_recuperada, "TIEMPO_ACCION_GLOBAL"] = "imputar_donante_cola"
    df_3ra.loc[mascara_t1_recuperada, "TIEMPO_MOTIVO_GLOBAL"] = "t1_recuperado_con_mediana_cola_donante"
    df_3ra.loc[mascara_t2_t3_recuperada, "TIEMPO_ACCION_GLOBAL"] = "imputar_donante_t2_t3"
    df_3ra.loc[mascara_t2_t3_recuperada, "TIEMPO_MOTIVO_GLOBAL"] = "t2_t3_recuperados_con_donantes"

    df_final = df_3ra.copy()
    df_final["T1_FINAL_4TA"] = df_final["T1_FINAL_3RA"]
    df_final["T2_FINAL_4TA"] = df_final["T2_FINAL_3RA"]
    df_final["T3_FINAL_4TA"] = df_final["T3_FINAL_3RA"]
    objetivo_ajuste_final = (
        config["aplicar_swap_final_t2_t3"]
        & (df_final["TIEMPO_MOTIVO_GLOBAL"] == "segunda_pasada_rompe_orden")
        & df_final["T1_FINAL_3RA"].isna()
        & df_final["T2_FINAL_3RA"].notna()
        & df_final["T3_FINAL_3RA"].notna()
        & (df_final["T2_FINAL_3RA"] > df_final["T3_FINAL_3RA"])
    )
    df_final.loc[objetivo_ajuste_final, "T2_FINAL_4TA"] = df_final.loc[objetivo_ajuste_final, "T3_FINAL_3RA"]
    df_final.loc[objetivo_ajuste_final, "T3_FINAL_4TA"] = df_final.loc[objetivo_ajuste_final, "T2_FINAL_3RA"]
    mascara_con_donante_final = objetivo_ajuste_final & df_final["COLA_DONANTE_ELEGIDA"].notna()
    df_final.loc[mascara_con_donante_final, "T1_FINAL_4TA"] = (
        df_final.loc[mascara_con_donante_final, "T2_FINAL_4TA"] - df_final.loc[mascara_con_donante_final, "COLA_DONANTE_ELEGIDA"]
    )
    mascara_recuperada_final = (
        objetivo_ajuste_final
        & df_final["T1_FINAL_4TA"].notna()
        & df_final["T2_FINAL_4TA"].notna()
        & df_final["T3_FINAL_4TA"].notna()
        & (df_final["T1_FINAL_4TA"] <= df_final["T2_FINAL_4TA"])
        & (df_final["T2_FINAL_4TA"] <= df_final["T3_FINAL_4TA"])
    )
    df_final["TIEMPO_ACCION_CIERRE"] = df_final["TIEMPO_ACCION_GLOBAL"]
    df_final["TIEMPO_MOTIVO_CIERRE"] = df_final["TIEMPO_MOTIVO_GLOBAL"]
    df_final.loc[mascara_recuperada_final, "TIEMPO_ACCION_CIERRE"] = "swap_t2_t3_y_reconstruir_t1"
    df_final.loc[mascara_recuperada_final, "TIEMPO_MOTIVO_CIERRE"] = "t2_t3_intercambiados_y_t1_reconstruido_con_cola_donante"

    df_final["T1_FINAL_5TA"] = df_final["T1_FINAL_4TA"]
    df_final["T2_FINAL_5TA"] = df_final["T2_FINAL_4TA"]
    df_final["T3_FINAL_5TA"] = df_final["T3_FINAL_4TA"]
    df_final = apply_short_complete_time_swaps(df_final, config)

    df_final = consolidate_post_time_rows(df_final, DEDUPLICACION_DUPLICADO_CERCANO_SEGUNDOS)
    df_final = apply_short_complete_time_swaps(df_final, config)
    df_final = consolidate_fragmented_flow_rows(df_final)
    df_final = apply_short_complete_time_swaps(df_final, config)

    df_tiempos_analisis_flujo = df_final.copy()

    mascara_eliminacion_borde_final = df_final["BORDE_CASETA_ELIMINAR"] & ~df_final["TIEMPOS_COMPLETOS_CIERRE"]
    df_tiempos_eliminados_borde = df_final[mascara_eliminacion_borde_final].copy()
    df_final = df_final[~mascara_eliminacion_borde_final].copy()
    df_pendientes = df_final[~df_final["TIEMPOS_COMPLETOS_CIERRE"]].copy()
    resumen_tiempos = (
        df_final["TIEMPO_ACCION_CIERRE"].value_counts(dropna=False).rename_axis("TIEMPO_ACCION_CIERRE").to_frame("filas")
    )
    return {
        "df_tiempos_base": df_tiempos_base,
        "df_tiempos_bordes": df_tiempos_bordes,
        "df_tiempos_analisis_flujo": df_tiempos_analisis_flujo,
        "df_tiempos_eliminados_borde": df_tiempos_eliminados_borde,
        "df_tiempos_final": df_final,
        "df_tiempos_pendientes": df_pendientes,
        "resumen_tiempos": resumen_tiempos,
    }


def build_export_tables(
    df_original: pd.DataFrame,
    plate_result: dict[str, pd.DataFrame],
    time_result: dict[str, pd.DataFrame],
    config: dict,
    manual_rules_df: pd.DataFrame,
) -> dict[str, pd.DataFrame]:
    df = plate_result["df"]
    df_trabajo = plate_result["df_trabajo"]
    df_eliminados_placa = plate_result["df_eliminados"]
    df_tiempos_final = time_result["df_tiempos_final"]
    df_eliminados_borde = time_result["df_tiempos_eliminados_borde"]
    df_tiempos_pendientes = time_result["df_tiempos_pendientes"]
    flow_findings = detect_flow_fuga_candidates({"df_tiempos_bordes": time_result["df_tiempos_bordes"]})

    placa_meta_export = df_trabajo[
        ["_ORDEN_FILA", "PLACA", "PLACA_FINAL_DECIDIDA", "PLACA_ACCION_FINAL", "PLACA_AJUSTE_MANUAL"]
    ].copy()
    df_export_base = df_tiempos_final.merge(
        placa_meta_export,
        on="_ORDEN_FILA",
        how="left",
        validate="one_to_one",
        suffixes=("", "_PLACA"),
    )
    df_export_base["PLACA_FINAL"] = df_export_base["PLACA_FINAL_DECIDIDA"]
    df_export_base["TIEMPO_ACCION_FINAL"] = df_export_base["TIEMPO_ACCION_CIERRE"]
    df_export_base["TIEMPO_MOTIVO_FINAL"] = df_export_base["TIEMPO_MOTIVO_CIERRE"]
    df_export_base["ACCION_REALIZADA"] = df_export_base.apply(
        lambda row: combine_action(row["PLACA_ACCION_FINAL"], row["TIEMPO_ACCION_FINAL"]),
        axis=1,
    )
    df_export_base["T_COLA_FINAL_SEGUNDOS"] = df_export_base["T_COLA_FINAL"].map(timedelta_a_segundos)
    df_export_base["T_COLA_FINAL_MINUTOS"] = df_export_base["T_COLA_FINAL"].map(timedelta_a_minutos)
    df_export_base["T_CASETA_FINAL_SEGUNDOS"] = df_export_base["T_CASETA_FINAL"].map(timedelta_a_segundos)
    df_export_base["T_CASETA_FINAL_MINUTOS"] = df_export_base["T_CASETA_FINAL"].map(timedelta_a_minutos)
    df_export_base["T_TEC_FINAL_SEGUNDOS"] = df_export_base["T_TEC_FINAL"].map(timedelta_a_segundos)
    df_export_base["T_TEC_FINAL_MINUTOS"] = df_export_base["T_TEC_FINAL"].map(timedelta_a_minutos)
    df_export_base["REGISTRO_AJUSTADO"] = (
        df_export_base["PLACA_ACCION_FINAL"].isin(ACCIONES_PLACA_AJUSTE)
        | df_export_base["TIEMPO_ACCION_FINAL"].ne("sin_cambio")
    )
    df_export_base["PLACA_PENDIENTE_REVISION"] = df_export_base["PLACA_ACCION_FINAL"].eq("mantener_observada")

    if not flow_findings["fugas_probables"].empty:
        fuga_detail = flow_findings["fugas_probables"].copy()
        fuga_detail["_FUGA_PRIORIDAD"] = fuga_detail["TIPO_FUGA"].astype(str).map(
            lambda value: 0
            if value.startswith("fuga_fuerte_")
            else 1
            if value.startswith("fuga_probable_")
            else 2
            if value.startswith("incompleto_no_concluyente")
            else 9
        )
        fuga_detail["_FUGA_FECHA_KEY"] = pd.to_datetime(fuga_detail["FECHA"], errors="coerce").dt.normalize()
        fuga_detail["_FUGA_PEAJE_KEY"] = fuga_detail["PEAJE"].map(normalize_text_key)
        fuga_detail["_FUGA_CASETA_KEY"] = fuga_detail["CASETA"].map(normalize_text_key)
        fuga_detail["_FUGA_SENTIDO_KEY"] = fuga_detail["SENTIDO"].map(normalize_text_key)
        fuga_detail["_FUGA_PLACA_KEY"] = fuga_detail["PLACA_FINAL"].map(normalize_text_key)
        fuga_detail = fuga_detail.sort_values(
            ["_FUGA_PRIORIDAD", "SCORE_FUGA", "PEAJE", "CASETA", "SENTIDO", "_FUGA_FECHA_KEY", "PLACA_FINAL"],
            ascending=[True, False, True, True, True, True, True],
        )
        fuga_detail = fuga_detail.drop_duplicates(
            subset=["_FUGA_PEAJE_KEY", "_FUGA_CASETA_KEY", "_FUGA_SENTIDO_KEY", "_FUGA_FECHA_KEY", "_FUGA_PLACA_KEY"],
            keep="first",
        )

        df_export_base["_FUGA_FECHA_KEY"] = pd.to_datetime(df_export_base["FECHA"], errors="coerce").dt.normalize()
        df_export_base["_FUGA_PEAJE_KEY"] = df_export_base["PEAJE"].map(normalize_text_key)
        df_export_base["_FUGA_CASETA_KEY"] = df_export_base["CASETA"].map(normalize_text_key)
        df_export_base["_FUGA_SENTIDO_KEY"] = df_export_base["SENTIDO"].map(normalize_text_key)
        df_export_base["_FUGA_PLACA_KEY"] = df_export_base["PLACA_FINAL"].map(normalize_text_key)
        df_export_base = df_export_base.merge(
            fuga_detail[
                [
                    "_FUGA_PEAJE_KEY",
                    "_FUGA_CASETA_KEY",
                    "_FUGA_SENTIDO_KEY",
                    "_FUGA_FECHA_KEY",
                    "_FUGA_PLACA_KEY",
                    "TIPO_FUGA",
                    "NIVEL_CONFIANZA",
                    "SCORE_FUGA",
                    "DETALLE",
                    "ES_FUGA_FUERTE",
                    "ES_FUGA_PROBABLE",
                    "ES_INCOMPLETO_NO_CONCLUYENTE",
                ]
            ],
            on=["_FUGA_PEAJE_KEY", "_FUGA_CASETA_KEY", "_FUGA_SENTIDO_KEY", "_FUGA_FECHA_KEY", "_FUGA_PLACA_KEY"],
            how="left",
        )
        df_export_base = df_export_base.drop(
            columns=["_FUGA_FECHA_KEY", "_FUGA_PEAJE_KEY", "_FUGA_CASETA_KEY", "_FUGA_SENTIDO_KEY", "_FUGA_PLACA_KEY"]
        )
    else:
        df_export_base["TIPO_FUGA"] = pd.NA
        df_export_base["NIVEL_CONFIANZA"] = pd.NA
        df_export_base["SCORE_FUGA"] = pd.NA
        df_export_base["DETALLE"] = pd.NA
        df_export_base["ES_FUGA_FUERTE"] = False
        df_export_base["ES_FUGA_PROBABLE"] = False
        df_export_base["ES_INCOMPLETO_NO_CONCLUYENTE"] = False

    df_export_base["CLASIFICACION_FUGA_FLUJO"] = df_export_base["TIPO_FUGA"].fillna("sin_hallazgo")
    df_export_base["NIVEL_FUGA_FLUJO"] = df_export_base["NIVEL_CONFIANZA"].fillna("sin_hallazgo")
    df_export_base["FUGA_FLUJO_IDENTIFICADA"] = df_export_base["TIPO_FUGA"].notna()
    df_export_base = df_export_base.rename(columns={"DETALLE": "DETALLE_FUGA_FLUJO"})

    casos_excluidos_contraste = pd.DataFrame()
    if config.get("modo_contraste_estricto", False):
        mascara_exclusion_contraste = df_export_base["TIEMPOS_COMPLETOS_CIERRE"] & df_export_base["TIEMPO_ACCION_FINAL"].isin(
            ["imputar_mediana_local", "imputar_donante_cola"]
        )
        if mascara_exclusion_contraste.any():
            casos_excluidos_contraste = df_export_base.loc[mascara_exclusion_contraste].copy()
            casos_excluidos_contraste = casos_excluidos_contraste[
                [
                    "PEAJE",
                    "CASETA",
                    "SENTIDO",
                    "FECHA",
                    "VEHICULO",
                    "PLACA",
                    "PLACA_FINAL",
                    "ACCION_REALIZADA",
                    "TIEMPO_MOTIVO_FINAL",
                    "LLEGADA_COLA_FINAL",
                    "LLEGADA_CASETA_FINAL",
                    "SALIDA_CASETA_FINAL",
                    "T_TEC_FINAL_TXT",
                    "T_CASETA_FINAL_TXT",
                ]
            ].rename(
                columns={
                    "TIEMPO_MOTIVO_FINAL": "MOTIVO_ELIMINACION",
                    "LLEGADA_COLA_FINAL": "LLEGADA COLA",
                    "LLEGADA_CASETA_FINAL": "LLEGADA CASETA",
                    "SALIDA_CASETA_FINAL": "SALIDA CASETA",
                    "T_TEC_FINAL_TXT": "T. TEC",
                    "T_CASETA_FINAL_TXT": "T. CASETA",
                }
            )
            casos_excluidos_contraste["ETAPA_ELIMINACION"] = "contraste_estricto"
            casos_excluidos_contraste["ACCION_REALIZADA"] = (
                "excluido_contraste_estricto:" + casos_excluidos_contraste["ACCION_REALIZADA"].astype(str)
            )
            df_export_base = df_export_base.loc[~mascara_exclusion_contraste].copy()

    base_limpia = df_export_base[df_export_base["TIEMPOS_COMPLETOS_CIERRE"]].copy()
    base_limpia = base_limpia[
        [
            "PEAJE",
            "CASETA",
            "SENTIDO",
            "FECHA",
            "VEHICULO",
            "PLACA_FINAL",
            "LLEGADA_COLA_FINAL",
            "LLEGADA_CASETA_FINAL",
            "SALIDA_CASETA_FINAL",
            "T_COLA_FINAL_TXT",
            "T_COLA_FINAL_SEGUNDOS",
            "T_COLA_FINAL_MINUTOS",
            "T_CASETA_FINAL_TXT",
            "T_CASETA_FINAL_SEGUNDOS",
            "T_CASETA_FINAL_MINUTOS",
            "T_TEC_FINAL_TXT",
            "T_TEC_FINAL_SEGUNDOS",
            "T_TEC_FINAL_MINUTOS",
            "FUGA_FLUJO_IDENTIFICADA",
            "CLASIFICACION_FUGA_FLUJO",
            "NIVEL_FUGA_FLUJO",
            "SCORE_FUGA",
            "DETALLE_FUGA_FLUJO",
            "ES_FUGA_FUERTE",
            "ES_FUGA_PROBABLE",
            "ES_INCOMPLETO_NO_CONCLUYENTE",
            "ACCION_REALIZADA",
        ]
    ].rename(
        columns={
            "PLACA_FINAL": "PLACA",
            "LLEGADA_COLA_FINAL": "LLEGADA COLA",
            "LLEGADA_CASETA_FINAL": "LLEGADA CASETA",
            "SALIDA_CASETA_FINAL": "SALIDA CASETA",
            "T_COLA_FINAL_TXT": "T. COLA",
            "T_COLA_FINAL_SEGUNDOS": "T. COLA_SEGUNDOS",
            "T_COLA_FINAL_MINUTOS": "T. COLA_MINUTOS",
            "T_CASETA_FINAL_TXT": "T. CASETA",
            "T_CASETA_FINAL_SEGUNDOS": "T. CASETA_SEGUNDOS",
            "T_CASETA_FINAL_MINUTOS": "T. CASETA_MINUTOS",
            "T_TEC_FINAL_TXT": "T. TEC",
            "T_TEC_FINAL_SEGUNDOS": "T. TEC_SEGUNDOS",
            "T_TEC_FINAL_MINUTOS": "T. TEC_MINUTOS",
            "SCORE_FUGA": "SCORE_FUGA_FLUJO",
        }
    )

    eliminados_placa = pd.DataFrame()
    if not df_eliminados_placa.empty:
        eliminados_placa = df_eliminados_placa[
            EXPECTED_COLUMNS + ["PLACA_FINAL_DECIDIDA", "PLACA_ACCION_FINAL", "PLACA_AJUSTE_MANUAL", "REVISION_GRUPO_BLOQUE", "PLACA_MOTIVO"]
        ].copy()
        eliminados_placa["PLACA_FINAL"] = eliminados_placa["PLACA_FINAL_DECIDIDA"]
        eliminados_placa["ETAPA_ELIMINACION"] = "placa"
        eliminados_placa["ACCION_REALIZADA"] = "eliminado:" + eliminados_placa["PLACA_ACCION_FINAL"].astype(str)
        eliminados_placa["MOTIVO_ELIMINACION"] = (
            eliminados_placa["PLACA_AJUSTE_MANUAL"]
            .combine_first(eliminados_placa["REVISION_GRUPO_BLOQUE"])
            .combine_first(eliminados_placa["PLACA_MOTIVO"])
        )

    eliminados_tiempo = pd.DataFrame()
    if not df_eliminados_borde.empty:
        eliminados_tiempo = df_eliminados_borde[
            [
                "PEAJE",
                "CASETA",
                "SENTIDO",
                "FECHA",
                "VEHICULO",
                "LLEGADA COLA",
                "LLEGADA CASETA",
                "SALIDA CASETA",
                "T. TEC",
                "T. CASETA",
                "_ORDEN_FILA",
                "BORDE_CASETA_ACCION",
                "BORDE_CASETA_MOTIVO",
            ]
        ].copy()
        eliminados_tiempo = eliminados_tiempo.merge(
            placa_meta_export[["_ORDEN_FILA", "PLACA", "PLACA_FINAL_DECIDIDA"]],
            on="_ORDEN_FILA",
            how="left",
            validate="many_to_one",
        )
        eliminados_tiempo["PLACA_FINAL"] = eliminados_tiempo["PLACA_FINAL_DECIDIDA"]
        eliminados_tiempo["ETAPA_ELIMINACION"] = "tiempo_borde_caseta"
        eliminados_tiempo["ACCION_REALIZADA"] = "eliminado:" + eliminados_tiempo["BORDE_CASETA_ACCION"].astype(str)
        eliminados_tiempo["MOTIVO_ELIMINACION"] = eliminados_tiempo["BORDE_CASETA_MOTIVO"]

    columnas_eliminados = [
        "PEAJE",
        "CASETA",
        "SENTIDO",
        "FECHA",
        "VEHICULO",
        "PLACA",
        "PLACA_FINAL",
        "ETAPA_ELIMINACION",
        "ACCION_REALIZADA",
        "MOTIVO_ELIMINACION",
        "LLEGADA COLA",
        "LLEGADA CASETA",
        "SALIDA CASETA",
        "T. TEC",
        "T. CASETA",
    ]
    frames_eliminados = [
        eliminados_placa.reindex(columns=columnas_eliminados),
        eliminados_tiempo.reindex(columns=columnas_eliminados),
        casos_excluidos_contraste.reindex(columns=columnas_eliminados),
    ]
    frames_eliminados = [frame for frame in frames_eliminados if not frame.empty]
    if frames_eliminados:
        casos_eliminados = pd.concat(frames_eliminados, ignore_index=True)
    else:
        casos_eliminados = pd.DataFrame(columns=columnas_eliminados)

    casos_pendientes = pd.DataFrame()
    if not df_tiempos_pendientes.empty:
        casos_pendientes = df_tiempos_pendientes[
            [
                "PEAJE",
                "CASETA",
                "SENTIDO",
                "FECHA",
                "VEHICULO",
                "PLACA_FINAL",
                "LLEGADA COLA",
                "LLEGADA CASETA",
                "SALIDA CASETA",
                "TIEMPO_ACCION_CIERRE",
                "TIEMPO_MOTIVO_CIERRE",
            ]
        ].copy()
        casos_pendientes = casos_pendientes.rename(
            columns={
                "PLACA_FINAL": "PLACA",
                "TIEMPO_ACCION_CIERRE": "ACCION_REALIZADA",
                "TIEMPO_MOTIVO_CIERRE": "MOTIVO",
            }
        )

    resumen_rows = []
    resumen_rows.extend(
        [
            {"seccion": "general", "indicador": "filas_entrada", "valor": len(df_original)},
            {"seccion": "general", "indicador": "filas_base_limpia", "valor": len(base_limpia)},
            {"seccion": "general", "indicador": "filas_eliminadas", "valor": len(casos_eliminados)},
            {"seccion": "general", "indicador": "filas_pendientes", "valor": len(casos_pendientes)},
            {"seccion": "general", "indicador": "filas_excluidas_contraste_estricto", "valor": len(casos_excluidos_contraste)},
        ]
    )
    for idx, row in plate_result["resumen_acciones_placa"].reset_index().iterrows():
        resumen_rows.append({"seccion": "placa", "indicador": row.iloc[0], "valor": row.iloc[1]})
    for idx, row in time_result["resumen_tiempos"].reset_index().iterrows():
        resumen_rows.append({"seccion": "tiempo", "indicador": row.iloc[0], "valor": row.iloc[1]})
    for clave, valor in config.items():
        resumen_rows.append({"seccion": "config", "indicador": clave, "valor": valor})
    reporte_resumen = pd.DataFrame(resumen_rows)

    config_df = pd.DataFrame(
        [{"clave": clave, "valor": valor} for clave, valor in config.items()]
        + [{"clave": "manual_rules_rows", "valor": len(manual_rules_df)}]
    )

    return {
        "base_limpia": base_limpia,
        "casos_eliminados": casos_eliminados,
        "casos_pendientes": casos_pendientes,
        "fugas_flujo": flow_findings["fugas_probables"],
        "fragmentaciones_probables": flow_findings["fragmentaciones_probables"],
        "reporte_resumen": reporte_resumen,
        "revision_placas": plate_result["df_revision_placas"],
        "bloques_decision": plate_result["df_bloques_decision"],
        "config_usada": config_df,
        "export_base_detalle": df_export_base,
    }


def build_exact_export_package(
    df_original: pd.DataFrame,
    export_tables: dict[str, pd.DataFrame],
) -> dict[str, pd.DataFrame]:
    base_limpia = export_tables["base_limpia"].copy()
    casos_eliminados = export_tables["casos_eliminados"].copy()
    df_export_base = export_tables["export_base_detalle"].copy()

    resumen_general = pd.Series(
        {
            "filas_base_original": len(df_original),
            "filas_base_limpia": len(base_limpia),
            "filas_ajustadas": int(df_export_base["REGISTRO_AJUSTADO"].sum()),
            "filas_con_placa_pendiente_revision": int(df_export_base["PLACA_PENDIENTE_REVISION"].sum()),
            "filas_eliminadas_placa": int(casos_eliminados["ETAPA_ELIMINACION"].eq("placa").sum()),
            "filas_eliminadas_tiempo": int(casos_eliminados["ETAPA_ELIMINACION"].eq("tiempo_borde_caseta").sum()),
            "filas_eliminadas_total": len(casos_eliminados),
        }
    ).rename_axis("indicador").reset_index(name="valor")

    resumen_accion_realizada = (
        base_limpia["ACCION_REALIZADA"]
        .value_counts(dropna=False)
        .rename_axis("ACCION_REALIZADA")
        .reset_index(name="filas")
    )

    resumen_acciones_placa = (
        df_export_base["PLACA_ACCION_FINAL"]
        .value_counts(dropna=False)
        .rename_axis("PLACA_ACCION_FINAL")
        .reset_index(name="filas")
    )

    resumen_acciones_tiempo = (
        df_export_base["TIEMPO_ACCION_FINAL"]
        .value_counts(dropna=False)
        .rename_axis("TIEMPO_ACCION_FINAL")
        .reset_index(name="filas")
    )

    resumen_eliminados = (
        casos_eliminados.groupby(["ETAPA_ELIMINACION", "ACCION_REALIZADA"], dropna=False)
        .size()
        .rename("filas")
        .reset_index()
        .sort_values(["ETAPA_ELIMINACION", "filas"], ascending=[True, False])
    )

    return {
        "base_limpia": base_limpia,
        "casos_eliminados": casos_eliminados,
        "casos_pendientes": export_tables["casos_pendientes"].copy(),
        "revision_placas": export_tables["revision_placas"].copy(),
        "bloques_decision": export_tables["bloques_decision"].copy(),
        "fugas_flujo": export_tables["fugas_flujo"].copy(),
        "fragmentaciones_probables": export_tables["fragmentaciones_probables"].copy(),
        "resumen_general": resumen_general,
        "resumen_accion_realizada": resumen_accion_realizada,
        "resumen_acciones_placa": resumen_acciones_placa,
        "resumen_acciones_tiempo": resumen_acciones_tiempo,
        "resumen_eliminados": resumen_eliminados,
    }


def derive_output_filenames(uploaded_name: str) -> dict[str, str]:
    stem = Path(uploaded_name).stem.strip() or "resultado"
    if stem.lower().startswith("data "):
        label = stem[5:].strip() or stem
        return {
            "report_label": label,
            "clean_excel": f"{stem}_limpio.xlsx",
            "report_excel": f"Resultados {label}.xlsx",
            "report_docx": f"Anexo tablas {label}.docx",
            "report_docx_model": f"Informe final {label} modelo 2022.docx",
            "extra_excel": f"Resultados {label} complementarios.xlsx",
        }
    return {
        "report_label": stem,
        "clean_excel": f"{stem}_limpio.xlsx",
        "report_excel": f"{stem}_resultados.xlsx",
        "report_docx": f"{stem}_anexo_tablas.docx",
        "report_docx_model": f"{stem}_informe_final_modelo_2022.docx",
        "extra_excel": f"{stem}_resultados_complementarios.xlsx",
    }


def build_processing_signature(
    uploaded_name: str,
    selected_sheet: str | None,
    file_bytes: bytes,
    mapping: dict[str, str],
    config: dict[str, object],
    manual_rules_df: pd.DataFrame,
) -> str:
    manual_rules_payload = (
        manual_rules_df.fillna("")
        .sort_index(axis=1)
        .to_dict(orient="records")
    )
    signature_payload = {
        "version": PROCESSING_SIGNATURE_VERSION,
        "uploaded_name": uploaded_name,
        "selected_sheet": selected_sheet or "",
        "file_sha1": hashlib.sha1(file_bytes).hexdigest(),
        "mapping": {key: mapping.get(key, "") for key in sorted(mapping)},
        "config": {key: config.get(key) for key in sorted(config)},
        "manual_rules": manual_rules_payload,
    }
    serialized = json.dumps(signature_payload, sort_keys=True, ensure_ascii=True, default=str)
    return hashlib.sha1(serialized.encode("utf-8")).hexdigest()


def build_processing_artifacts(
    uploaded_name: str,
    selected_sheet: str | None,
    processing_signature: str,
    result: dict[str, object],
) -> dict[str, object]:
    export_tables = result["export_tables"]
    dashboard = build_processing_dashboard(result["input_df"], result)
    output_filenames = derive_output_filenames(uploaded_name)
    report_label = output_filenames["report_label"]
    exact_export_bytes = to_exact_excel_bytes(result["exact_export"])
    return {
        "input_signature": processing_signature,
        "source_name": uploaded_name,
        "selected_sheet": selected_sheet,
        "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "export_tables": export_tables,
        "output_filenames": output_filenames,
        "base_limpia_bytes": exact_export_bytes,
        "base_limpia_only_bytes": to_excel_bytes({"base_limpia": export_tables["base_limpia"]}),
        "exact_export_bytes": exact_export_bytes,
        "complementary_excel_bytes": to_excel_bytes(result["complementary_package"]["excel_sheets"]),
        "report_docx_bytes": to_docx_bytes(report_label, result["informe_package"]),
        "report_docx_model_bytes": to_templated_docx_bytes(report_label, result["informe_package"]),
        "dashboard": dashboard,
        "fugas_report_bytes": to_excel_bytes(build_fugas_report_sheets(dashboard)),
    }


def render_processing_dashboard(dashboard: dict[str, object], processed_payload: dict[str, object], source_name: str) -> None:
    overview = dashboard["overview"]
    raw_tables = dashboard["raw_tables"]
    clean_tables = dashboard["clean_tables"]
    queue_theory = dashboard["queue_theory"]
    fugas_report_bytes = processed_payload["fugas_report_bytes"]
    output_filenames = processed_payload["output_filenames"]
    fugas_patron_detalle = dashboard["fugas_patron_detalle"]
    fugas_probables_detalle = dashboard["fugas_probables_detalle"]
    fragmentaciones_detalle = dashboard["fragmentaciones_detalle"]
    fragmentaciones_confianza = dashboard["fragmentaciones_confianza"]
    plate_actions = dashboard["plate_actions"]
    time_actions = dashboard["time_actions"]
    caseta_changes = dashboard["caseta_changes"]

    st.markdown('<div class="section-heading">Lectura ejecutiva del procesamiento</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-copy">El tablero resume la base original, las fugas detectadas, el trabajo de limpieza realizado y el estado final de la base limpia.</div>',
        unsafe_allow_html=True,
    )
    st.markdown(build_processing_flow_diagram_html(), unsafe_allow_html=True)

    st.markdown('<div class="section-heading">1. Estado de la base original</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-copy">Composicion inicial del archivo antes del pipeline, con foco en volumen por peaje, sentido y caseta.</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        build_dashboard_metric_items(
            [
                (overview["raw_rows"], "registros de entrada"),
                (overview["raw_peajes"], "peajes presentes"),
                (overview["raw_casetas"], "casetas presentes"),
                (overview["raw_sentidos"], "sentidos presentes"),
                (overview["fugas_rows"], "filas con patron X / longitud atipica"),
                (overview["fugas_unique"], "placas unicas con ese patron"),
            ]
        ),
        unsafe_allow_html=True,
    )
    raw_col1, raw_col2 = st.columns([1.08, 1.12])
    with raw_col1:
        render_dashboard_figure(plot_volume_by_sentido(raw_tables["por_peaje_sentido"], "Composicion original por peaje y sentido"))
    with raw_col2:
        render_dashboard_figure(plot_top_labels(raw_tables["por_caseta"][ ["ETIQUETA", "REGISTROS"]], "Casetas con mayor volumen en la base original", "#77a8ff"))

    st.markdown('<div class="section-heading">2. Hallazgos y valor del procesamiento</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-copy">Aqui se concentra la evidencia del trabajo del pipeline: deteccion de fugas, correcciones de placa, cambios de caseta y recuperaciones de tiempo. La lectura operativa sigue una jerarquia: fuga fuerte por flujo, fuga probable por flujo e incompleto no concluyente.</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="section-copy">El reporte consolida fugas por patron, fugas por flujo con nivel de confianza, fragmentaciones con confianza y placas con cambio de caseta. Fuga fuerte implica mejor soporte contextual; fuga probable sigue siendo alerta analitica; incompleto no concluyente no debe interpretarse como fuga confirmada.</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        build_dashboard_metric_items(
            [
                (overview["flagged_plates"], "filas identificadas para revision de placa"),
                (overview["corrected_plate_rows"], "filas con correccion de placa"),
                (overview["fugas_fuertes"], "fugas fuertes por flujo"),
                (overview["fugas_probables"], "fugas probables por flujo"),
                (overview["fugas_no_concluyentes"], "incompletos no concluyentes"),
                (overview["fragmentaciones_probables"], "fragmentaciones probables"),
                (overview["fragmentaciones_alta_confianza"], "filas potencialmente unificables"),
                (overview["caseta_change_groups"], "placas finales con cambio de caseta"),
                (overview["recovered_time_rows"], "tiempos recuperados"),
                (overview["final_time_adjustments"], "ajustes finales de tiempo"),
                (overview["deleted_rows"], "filas excluidas del resultado"),
                (overview["pending_rows"], "pendientes finales"),
            ]
        ),
        unsafe_allow_html=True,
    )
    central_col1, central_col2 = st.columns(2)
    with central_col1:
        render_dashboard_figure(
            plot_top_labels(
                plate_actions.rename(columns={"ACCION_UI": "ETIQUETA", "FILAS": "REGISTROS"})[["ETIQUETA", "REGISTROS"]],
                "Acciones de placa ejecutadas",
                "#163d72",
            )
        )
    with central_col2:
        render_dashboard_figure(
            plot_top_labels(
                time_actions.rename(columns={"ACCION_UI": "ETIQUETA", "FILAS": "REGISTROS"})[["ETIQUETA", "REGISTROS"]],
                "Acciones de tiempo ejecutadas",
                "#2f6ddc",
            )
        )
    detail_col1, detail_col2 = st.columns(2)
    with detail_col1:
        st.markdown('<div class="section-copy">Fugas por patron de placa: termina en X y supera 6 caracteres normalizados.</div>', unsafe_allow_html=True)
        st.dataframe(fugas_patron_detalle.head(20), use_container_width=True, hide_index=True)
    with detail_col2:
        st.markdown('<div class="section-copy">Placas finales que aparecen en mas de una caseta para el mismo peaje, sentido y fecha.</div>', unsafe_allow_html=True)
        st.dataframe(caseta_changes.head(20), use_container_width=True, hide_index=True)
    fuga_col1, fuga_col2 = st.columns(2)
    with fuga_col1:
        st.markdown(
            '<div class="section-copy">Jerarquia de fugas por flujo: fuga fuerte cuando hay mejor soporte contextual y recurrencia util; fuga probable cuando la evidencia es intermedia; incompleto no concluyente cuando el caso no debe leerse como fuga confirmada. En esta tabla conviven las tres etiquetas con su nivel de confianza y score.</div>',
            unsafe_allow_html=True,
        )
        st.dataframe(fugas_probables_detalle.head(20), use_container_width=True, hide_index=True)
    with fuga_col2:
        st.markdown('<div class="section-copy">Fragmentaciones probables: una fila trae cola y otra cercana trae caseta/salida con placa muy similar. Esta categoria compite contra fuga y ayuda a no sobredeclarar evasiones cuando la mejor explicacion es un registro partido.</div>', unsafe_allow_html=True)
        st.dataframe(fragmentaciones_detalle.head(20), use_container_width=True, hide_index=True)
    summary_col = st.columns(1)[0]
    with summary_col:
        st.markdown('<div class="section-copy">Lectura operativa consolidada: primero se revisan fragmentaciones, luego fugas fuertes, despues fugas probables y finalmente incompletos no concluyentes.</div>', unsafe_allow_html=True)
        st.dataframe(
            pd.DataFrame(
                [
                    {"hallazgo": "fragmentaciones probables", "filas": overview["fragmentaciones_probables"]},
                    {"hallazgo": "alta confianza para unificar", "filas": overview["fragmentaciones_alta_confianza"]},
                    {"hallazgo": "fugas fuertes por flujo", "filas": overview["fugas_fuertes"]},
                    {"hallazgo": "fugas probables por flujo", "filas": overview["fugas_probables"]},
                    {"hallazgo": "incompletos no concluyentes", "filas": overview["fugas_no_concluyentes"]},
                ]
            ),
            use_container_width=True,
            hide_index=True,
        )

    st.markdown('<div class="section-heading">3. Estado de la base limpia</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-copy">Distribucion final despues del pipeline, con porcentaje de retencion y concentracion por peaje, sentido y caseta.</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        build_dashboard_metric_items(
            [
                (overview["clean_rows"], "registros en base limpia"),
                (f"{overview['retention_pct']}%", "retencion sobre entrada"),
                (overview["clean_peajes"], "peajes en salida"),
                (overview["clean_casetas"], "casetas en salida"),
                (overview["clean_sentidos"], "sentidos en salida"),
                (overview["pending_rows"], "pendientes finales"),
            ]
        ),
        unsafe_allow_html=True,
    )
    clean_col1, clean_col2 = st.columns([1.02, 1.24])
    with clean_col1:
        render_dashboard_figure(plot_volume_by_peaje(clean_tables["por_peaje"], "Registros finales por peaje", "#0f8b6d"))
    with clean_col2:
        render_dashboard_figure(plot_volume_by_sentido(clean_tables["por_peaje_sentido"], "Composicion final por peaje y sentido"))
    clean_col3, clean_col4 = st.columns([1.08, 1.12])
    with clean_col3:
        render_dashboard_figure(plot_top_labels(clean_tables["por_caseta"][["ETIQUETA", "REGISTROS"]], "Casetas con mayor volumen en la base limpia", "#63c29b"))
    with clean_col4:
        st.markdown('<div class="section-copy">Top operativo de casetas en la base limpia</div>', unsafe_allow_html=True)
        st.dataframe(clean_tables["por_caseta"][["PEAJE", "CASETA", "SENTIDO", "REGISTROS", "PARTICIPACION_%"]].head(12), use_container_width=True, hide_index=True)

    st.markdown('<div class="section-heading">4. Resultados del analisis de teoria de colas</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-copy">Este bloque resume lo que dice la base limpia sobre tiempos de espera y tamano de cola real. Se calcula solo con registros que terminan con tiempos completos, para comparar de forma estable el comportamiento global, por peaje y por caseta.</div>',
        unsafe_allow_html=True,
    )
    if queue_theory["general_insights"].empty:
        st.info("No hay suficientes registros con tiempos completos para resumir teoria de colas en el dashboard.")
    else:
        st.markdown(build_dashboard_metric_items(queue_theory["general_cards"]), unsafe_allow_html=True)
        st.markdown('<div class="section-copy">Lectura general de teoria de colas</div>', unsafe_allow_html=True)
        st.dataframe(queue_theory["general_insights"], use_container_width=True, hide_index=True)
        st.markdown('<div class="section-copy">Resultados consolidados por peaje</div>', unsafe_allow_html=True)
        st.dataframe(queue_theory["by_peaje"], use_container_width=True, hide_index=True)
        st.markdown('<div class="section-copy">Resultados consolidados por caseta</div>', unsafe_allow_html=True)
        st.dataframe(queue_theory["by_caseta"], use_container_width=True, hide_index=True)

    st.markdown('<div class="section-heading">5. Descargas disponibles</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-copy">Descarga desde aqui el libro general, la hoja sola de base limpia, los resultados complementarios y los documentos Word del procesamiento.</div>',
        unsafe_allow_html=True,
    )
    download_col1, download_col2, download_col3, download_col4, download_col5 = st.columns(5)
    with download_col1:
        st.download_button(
            "Descargar Base limpia",
            data=processed_payload["base_limpia_bytes"],
            file_name=output_filenames["report_excel"],
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="download_base_limpia",
        )
    with download_col2:
        st.download_button(
            "Descargar solo hoja base limpia",
            data=processed_payload["base_limpia_only_bytes"],
            file_name=output_filenames["clean_excel"],
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="download_hoja_base_limpia",
        )
    with download_col3:
        st.download_button(
            "Descargar resultados complementarios",
            data=processed_payload["complementary_excel_bytes"],
            file_name=output_filenames["extra_excel"],
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="download_resultados_complementarios",
        )
    with download_col4:
        st.download_button(
            "Descargar anexo DOCX",
            data=processed_payload["report_docx_bytes"],
            file_name=output_filenames["report_docx"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_anexo_docx",
        )
    with download_col5:
        st.download_button(
            "Descargar informe modelo",
            data=processed_payload["report_docx_model_bytes"],
            file_name=output_filenames["report_docx_model"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_informe_modelo_docx",
        )
    st.download_button(
        "Descargar reporte de fugas",
        data=fugas_report_bytes,
        file_name=f"reporte_fugas_{Path(source_name).stem}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key="download_fugas_report",
    )


def render_processing_outputs(processed_payload: dict[str, object], storage_backend, can_view_history: bool) -> None:
    export_tables = processed_payload["export_tables"]
    dashboard = processed_payload["dashboard"]

    info_col, clear_col = st.columns([4, 1])
    info_col.caption(
        "Resultados disponibles para "
        f"{processed_payload['source_name']}"
        + (
            f" | Hoja: {processed_payload['selected_sheet']}"
            if processed_payload["selected_sheet"]
            else ""
        )
        + f" | Procesado: {processed_payload['processed_at']}"
    )
    if clear_col.button("Limpiar resultado", use_container_width=True, key="clear_processing_result"):
        st.session_state.pop(TEC_RESULT_STATE_KEY, None)
        st.rerun()

    tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs(
        ["Dashboard", "Resumen", "Base limpia", "Eliminados", "Pendientes", "Revision placas", "Bloques", "Fugas flujo", "Fragmentaciones"]
    )
    with tab0:
        render_processing_dashboard(dashboard, processed_payload, processed_payload["source_name"])
    with tab1:
        st.dataframe(export_tables["reporte_resumen"], use_container_width=True)
    with tab2:
        st.dataframe(export_tables["base_limpia"], use_container_width=True)
    with tab3:
        st.dataframe(export_tables["casos_eliminados"], use_container_width=True)
    with tab4:
        st.dataframe(export_tables["casos_pendientes"], use_container_width=True)
    with tab5:
        st.dataframe(export_tables["revision_placas"], use_container_width=True)
    with tab6:
        st.dataframe(export_tables["bloques_decision"], use_container_width=True)
    with tab7:
        st.dataframe(export_tables["fugas_flujo"], use_container_width=True)
    with tab8:
        st.dataframe(export_tables["fragmentaciones_probables"], use_container_width=True)

def calcular_cola_espera_real(grupo: pd.DataFrame) -> pd.DataFrame:
    grupo = grupo.sort_values(
        [
            "LLEGADA_COLA_FINAL_TD",
            "LLEGADA_CASETA_FINAL_TD",
            "SALIDA_CASETA_FINAL_TD",
            "_ORDEN_FILA",
        ],
        na_position="last",
    ).copy()

    activos = []
    colas = []

    for row in grupo.itertuples(index=False):
        t1 = getattr(row, "LLEGADA_COLA_FINAL_TD")
        activos = [veh for veh in activos if veh["t3"] > t1]
        cola_espera = sum(1 for veh in activos if veh["t2"] > t1)
        colas.append(cola_espera)
        activos.append(
            {
                "t2": getattr(row, "LLEGADA_CASETA_FINAL_TD"),
                "t3": getattr(row, "SALIDA_CASETA_FINAL_TD"),
            }
        )

    grupo["COLA_ESPERA_USUARIOS"] = colas
    return grupo


def format_zero_blank(value):
    if pd.isna(value):
        return ""
    if isinstance(value, str):
        return value
    try:
        return "" if float(value) == 0 else value
    except (TypeError, ValueError):
        return value


def formatear_tabla_frecuencia(tabla: pd.DataFrame) -> pd.DataFrame:
    tabla = tabla.copy()
    columnas_valor = tabla.columns[1:]
    for columna in columnas_valor:
        tabla[columna] = tabla[columna].map(format_zero_blank)
    return tabla


def agregar_tabla_docx(doc: Document, titulo: str, tabla: pd.DataFrame) -> None:
    doc.add_paragraph(titulo, style="Heading 2")
    docx_table = doc.add_table(rows=1, cols=len(tabla.columns))
    docx_table.style = "Table Grid"
    docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, columna in enumerate(tabla.columns):
        docx_table.rows[0].cells[idx].text = str(columna)

    for fila in tabla.itertuples(index=False):
        celdas = docx_table.add_row().cells
        for idx, valor in enumerate(fila):
            celdas[idx].text = "" if pd.isna(valor) else str(valor)


def resolve_informe_template_path() -> Path | None:
    for candidate in INFORME_TEMPLATE_CANDIDATES:
        if candidate.exists():
            return candidate
    return None


def write_dataframe_to_existing_docx_table(table, tabla: pd.DataFrame) -> None:
    while len(table.columns) < len(tabla.columns):
        table.add_column(Inches(1.0))

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    if not table.rows:
        table.add_row()

    header_cells = table.rows[0].cells
    for idx, cell in enumerate(header_cells):
        cell.text = str(tabla.columns[idx]) if idx < len(tabla.columns) else ""

    for fila in tabla.itertuples(index=False):
        celdas = table.add_row().cells
        for idx, cell in enumerate(celdas):
            if idx >= len(fila):
                cell.text = ""
                continue
            valor = fila[idx]
            cell.text = "" if pd.isna(valor) else str(valor)


def normalize_text_key(value: object) -> str:
    text = "" if value is None else str(value)
    normalized = unicodedata.normalize("NFKD", text)
    ascii_only = normalized.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"\s+", " ", ascii_only).strip().lower()


def replace_docx_paragraph_text(doc: Document, search_text: str, replacement: str) -> bool:
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            paragraph.text = replacement
            return True
    return False


def replace_docx_paragraph_contains(doc: Document, search_snippet: str, replacement: str) -> bool:
    snippet_key = normalize_text_key(search_snippet)
    for paragraph in doc.paragraphs:
        if snippet_key in normalize_text_key(paragraph.text):
            paragraph.text = replacement
            return True
    return False


def replace_docx_paragraph_at_index(doc: Document, paragraph_idx: int, replacement: str) -> bool:
    if 0 <= paragraph_idx < len(doc.paragraphs):
        doc.paragraphs[paragraph_idx].text = replacement
        return True
    return False


def pick_report_variant(seed: str, key: str, options: list[str]) -> str:
    if not options:
        return ""
    digest = hashlib.sha256(f"{seed}|{key}".encode("utf-8")).hexdigest()
    return options[int(digest[:8], 16) % len(options)]


def build_report_variation_seed(report_label: str, informe_package: dict[str, object]) -> str:
    rows = len(informe_package.get("df_resultados", pd.DataFrame()))
    now_key = datetime.now().strftime("%Y%m%d%H%M%S%f")
    return hashlib.sha256(f"{report_label}|{rows}|{now_key}".encode("utf-8")).hexdigest()


def summarize_informe_facts(informe_package: dict[str, object]) -> dict[str, object]:
    tabla_programacion = informe_package.get("tabla_programacion", pd.DataFrame())
    tabla_personal = informe_package.get("tabla_personal", pd.DataFrame())
    tabla_tec_caseta = informe_package.get("tabla_tec_caseta", pd.DataFrame())
    tabla_tec_peaje = informe_package.get("tabla_tec_peaje", pd.DataFrame())
    tabla_cola = informe_package.get("tabla_cola_maxima", pd.DataFrame())
    df_resultados = informe_package.get("df_resultados", pd.DataFrame())

    peajes = []
    if not tabla_tec_peaje.empty and "Peaje" in tabla_tec_peaje.columns:
        peajes = [str(value) for value in pd.Series(tabla_tec_peaje["Peaje"]).dropna().astype(str).unique().tolist()]
    peajes_text = ", ".join(peajes[:-1]) + (" y " + peajes[-1] if len(peajes) > 1 else (peajes[0] if peajes else "las estaciones evaluadas"))

    total_casetas = int(tabla_tec_caseta[["Peaje", "Caseta Controlada", "Sentido de Circulacion"]].drop_duplicates().shape[0]) if not tabla_tec_caseta.empty else 0
    total_sentidos = int(tabla_tec_peaje["Sentido de Circulacion"].astype(str).nunique()) if not tabla_tec_peaje.empty else 0
    total_peajes = int(pd.Series(peajes).nunique()) if peajes else 0
    total_programaciones = int(tabla_programacion.shape[0]) if not tabla_programacion.empty else 0
    total_personal_rows = int(tabla_personal.shape[0]) if not tabla_personal.empty else 0

    top_tec_peaje = None
    if not tabla_tec_peaje.empty:
        top_tec_peaje = tabla_tec_peaje.sort_values("Tiempo de Espera en Cola - TEC", ascending=False).iloc[0]

    top_tec_caseta = None
    if not tabla_tec_caseta.empty:
        top_tec_caseta = tabla_tec_caseta.sort_values("Tiempo de Espera en Cola - TEC", ascending=False).iloc[0]

    top_cola = None
    if not tabla_cola.empty:
        top_cola = tabla_cola.sort_values("Cola maxima real", ascending=False).iloc[0]

    tec_peaje_values = pd.to_numeric(tabla_tec_peaje.get("Tiempo de Espera en Cola - TEC"), errors="coerce") if not tabla_tec_peaje.empty else pd.Series(dtype=float)
    tec_caseta_values = pd.to_numeric(tabla_tec_caseta.get("Tiempo de Espera en Cola - TEC"), errors="coerce") if not tabla_tec_caseta.empty else pd.Series(dtype=float)
    cola_values = pd.to_numeric(tabla_cola.get("Cola maxima real"), errors="coerce") if not tabla_cola.empty else pd.Series(dtype=float)

    tec_threshold = 3.0
    tec_near_threshold = 2.5
    peaje_over_threshold = int((tec_peaje_values > tec_threshold).sum()) if not tec_peaje_values.empty else 0
    caseta_over_threshold = int((tec_caseta_values > tec_threshold).sum()) if not tec_caseta_values.empty else 0
    peaje_near_threshold = int(((tec_peaje_values >= tec_near_threshold) & (tec_peaje_values <= tec_threshold)).sum()) if not tec_peaje_values.empty else 0
    caseta_near_threshold = int(((tec_caseta_values >= tec_near_threshold) & (tec_caseta_values <= tec_threshold)).sum()) if not tec_caseta_values.empty else 0

    return {
        "peajes_text": peajes_text,
        "total_peajes": total_peajes,
        "total_casetas": total_casetas,
        "total_sentidos": total_sentidos,
        "total_programaciones": total_programaciones,
        "total_personal_rows": total_personal_rows,
        "date_source_text": build_report_date_range_text(informe_package.get("df_resultados", pd.DataFrame())).replace("Fuente: ", ""),
        "top_tec_peaje": top_tec_peaje,
        "top_tec_caseta": top_tec_caseta,
        "top_cola": top_cola,
        "max_tec_peaje_value": float(tec_peaje_values.max()) if not tec_peaje_values.empty else None,
        "avg_tec_peaje_value": float(tec_peaje_values.mean()) if not tec_peaje_values.empty else None,
        "max_tec_caseta_value": float(tec_caseta_values.max()) if not tec_caseta_values.empty else None,
        "avg_tec_caseta_value": float(tec_caseta_values.mean()) if not tec_caseta_values.empty else None,
        "max_cola_value": int(cola_values.max()) if not cola_values.empty else None,
        "avg_cola_value": float(cola_values.mean()) if not cola_values.empty else None,
        "peaje_over_threshold": peaje_over_threshold,
        "caseta_over_threshold": caseta_over_threshold,
        "peaje_near_threshold": peaje_near_threshold,
        "caseta_near_threshold": caseta_near_threshold,
        "total_peaje_rows": int(tec_peaje_values.notna().sum()) if not tec_peaje_values.empty else 0,
        "total_caseta_rows": int(tec_caseta_values.notna().sum()) if not tec_caseta_values.empty else 0,
        "total_vehiculos_evaluados": int(df_resultados["PLACA_FINAL"].notna().sum()) if not df_resultados.empty and "PLACA_FINAL" in df_resultados.columns else 0,
    }


def format_minutes_reference(value: float | None) -> str:
    return f"{value:.2f} minutos" if value is not None and pd.notna(value) else "sin referencia suficiente"


def classify_tec_status(max_value: float | None) -> str:
    if max_value is None or pd.isna(max_value):
        return "sin evidencia suficiente para emitir una lectura de cumplimiento"
    if max_value > 3:
        return "con superacion del umbral contractual de 3.00 minutos"
    if max_value >= 2.5:
        return "dentro del umbral contractual, aunque cercana al limite de 3.00 minutos"
    return "holgadamente por debajo del umbral contractual de 3.00 minutos"


def classify_queue_status(max_cola: int | None) -> str:
    if max_cola is None:
        return "sin evidencia suficiente para caracterizar la cola real"
    if max_cola >= 10:
        return "episodios de acumulacion relevantes que merecen seguimiento operativo"
    if max_cola >= 5:
        return "picos de cola puntuales, aunque contenidos dentro de la jornada observada"
    return "colas acotadas y de baja acumulacion durante la muestra"


def build_dynamic_report_paragraphs(report_label: str, informe_package: dict[str, object]) -> dict[int, str]:
    seed = build_report_variation_seed(report_label, informe_package)
    facts = summarize_informe_facts(informe_package)
    top_tec_peaje = facts["top_tec_peaje"]
    top_tec_caseta = facts["top_tec_caseta"]
    top_cola = facts["top_cola"]

    top_tec_peaje_text = (
        f"{top_tec_peaje['Peaje']} - {top_tec_peaje['Sentido de Circulacion']} con {float(top_tec_peaje['Tiempo de Espera en Cola - TEC']):.2f} minutos"
        if top_tec_peaje is not None
        else "los frentes evaluados"
    )
    top_tec_caseta_text = (
        f"{top_tec_caseta['Peaje']} - caseta {top_tec_caseta['Caseta Controlada']} - {top_tec_caseta['Sentido de Circulacion']} con {float(top_tec_caseta['Tiempo de Espera en Cola - TEC']):.2f} minutos"
        if top_tec_caseta is not None
        else "las casetas observadas"
    )
    top_cola_text = (
        f"{int(top_cola['Cola maxima real'])} usuarios en {top_cola['Peaje']} - caseta {top_cola['Caseta Controlada']} - {top_cola['Sentido de Circulacion']}"
        if top_cola is not None
        else "la mayor cola observada durante la campana"
    )
    compliance_status_text = classify_tec_status(facts["max_tec_peaje_value"])
    caseta_status_text = classify_tec_status(facts["max_tec_caseta_value"])
    queue_status_text = classify_queue_status(facts["max_cola_value"])
    peaje_threshold_summary = (
        f"{facts['total_peaje_rows'] - facts['peaje_over_threshold']} de {facts['total_peaje_rows']} frentes peaje-sentido quedaron en o por debajo de 3.00 minutos"
        if facts["total_peaje_rows"]
        else "no se cuenta con frentes peaje-sentido suficientes para evaluar cumplimiento"
    )
    caseta_threshold_summary = (
        f"{facts['total_caseta_rows'] - facts['caseta_over_threshold']} de {facts['total_caseta_rows']} casetas-sentido quedaron en o por debajo de 3.00 minutos"
        if facts["total_caseta_rows"]
        else "no se cuenta con casetas-sentido suficientes para evaluar cumplimiento"
    )
    recommendation_focus_text = (
        f"priorizar la revision operativa de {top_tec_peaje['Peaje']} - {top_tec_peaje['Sentido de Circulacion']}"
        if top_tec_peaje is not None and facts["max_tec_peaje_value"] is not None and facts["max_tec_peaje_value"] > 3
        else f"mantener seguimiento preventivo sobre {top_tec_peaje['Peaje']} - {top_tec_peaje['Sentido de Circulacion']}"
        if top_tec_peaje is not None
        else "mantener seguimiento sobre los frentes evaluados"
    )

    variants = {
        24: [
            "En el marco del seguimiento a los indicadores de servicialidad de la Red Vial N° 5, el presente informe sintetiza la evaluacion del Tiempo de Espera en Cola desarrollada en {peajes_text}, considerando los registros procesados con la aplicacion TEC.",
            "Como parte del control de niveles de servicio de la Red Vial N° 5, este informe consolida la medicion del Tiempo de Espera en Cola efectuada en {peajes_text}, a partir de la base depurada y trazable generada por el sistema.",
            "El presente documento expone los resultados de la medicion del Tiempo de Espera en Cola para {peajes_text}, dentro del esquema de supervision de servicialidad aplicable a la Red Vial N° 5.",
        ],
        25: [
            "Para esta evaluacion, el TEC se interpreta como el tiempo promedio ponderado que transcurre desde que el usuario se incorpora a la cola hasta que concluye la atencion en la caseta, criterio que permite describir con consistencia el desempeno operativo observado.",
            "Bajo el enfoque utilizado en campo y gabinete, el TEC corresponde al intervalo promedio ponderado entre el inicio de la espera en cola y la culminacion del servicio, lo que permite caracterizar el comportamiento real de la atencion vehicular.",
            "El indicador TEC se calcula como el tiempo promedio ponderado de espera entre la formacion de la cola y el cierre de la atencion en caseta, de modo que refleje la experiencia efectiva del usuario durante la operacion.",
        ],
        31: [
            "El objetivo del estudio es medir y documentar el Tiempo de Espera en Cola en {total_peajes} frentes operativos, con un total de {total_casetas} casetas-sentido evaluadas, a fin de verificar el comportamiento del servicio durante las jornadas observadas.",
            "La evaluacion tiene por finalidad cuantificar el Tiempo de Espera en Cola en las estaciones consideradas, cubriendo {total_casetas} combinaciones caseta-sentido, para contar con evidencia objetiva sobre el desempeno operativo registrado.",
            "Este estudio busca estimar el Tiempo de Espera en Cola en los puntos de cobro analizados, organizando la informacion por caseta y sentido para sustentar la revision tecnica de la operacion observada.",
        ],
        39: [
            "La evaluacion se desarrollo bajo criterios de cobertura operativa, continuidad de observacion y consistencia en el registro de placas y tiempos, respetando la programacion definida para la medicion.",
            "Para asegurar la comparabilidad de resultados, la medicion se ejecuto con criterios uniformes de cobertura, ventanas de observacion y registro de eventos operativos relevantes.",
            "Las consideraciones generales del estudio priorizaron la trazabilidad del dato, la continuidad de la captura y la cobertura integral de las casetas operativas incluidas en la jornada.",
        ],
        41: [
            "Las labores abarcaron las estaciones y sentidos operativos comprendidos en la programacion de campo, de manera que el analisis represente el funcionamiento efectivo del sistema durante la muestra.",
            "Se relevaron los puntos de cobro que estuvieron operativos dentro de la ventana de observacion, procurando reflejar el comportamiento real de la atencion en cada frente evaluado.",
            "La muestra cubrio las estaciones y casetas habilitadas durante la jornada, con el proposito de capturar una lectura representativa del desempeno del sistema de cobro.",
        ],
        43: [
            "El tiempo de medicion se organizo sobre jornadas de control continuas y suficientes para capturar la dinamica de llegada, formacion de cola y atencion en caseta dentro de cada frente evaluado.",
            "La ventana de observacion se definio para recoger un volumen de datos adecuado y permitir una lectura estable del comportamiento de cola y servicio por caseta.",
            "Las franjas de control consideradas permitieron registrar la secuencia completa de llegada, espera y atencion de los vehiculos en las casetas observadas.",
        ],
        48: [
            "La medicion se ejecuto sobre las casetas y sentidos que permanecieron operativos durante la jornada programada, utilizando como referencia los horarios efectivamente observados en campo.",
            "Las ventanas de evaluacion se aplicaron a las casetas activas en cada estacion, considerando los periodos reales de funcionamiento detectados durante la medicion.",
            "Para el relevamiento se consideraron las casetas operativas dentro de los horarios efectivos de observacion, con el fin de mantener consistencia entre cobertura y resultados.",
        ],
        50: [
            "El levantamiento de informacion se sustento en el metodo de placas de rodaje, por tratarse de una tecnica adecuada para enlazar los hitos de llegada, atencion y salida de cada vehiculo.",
            "Se utilizo el metodo de placas de rodaje como procedimiento central de captura, ya que permite seguir la trayectoria temporal del vehiculo a lo largo del sistema de cola y atencion.",
            "La recoleccion de datos se realizo mediante el metodo de placas de rodaje, herramienta que facilita reconstruir con precision los tiempos asociados al proceso de cobro.",
        ],
        52: [
            "La labor de campo quedo respaldada con evidencias visuales y registros complementarios que permiten contextualizar los eventos operativos observados durante la muestra.",
            "El trabajo desarrollado en campo se complemento con soportes visuales y anotaciones de control, a fin de reforzar la trazabilidad de la campana de medicion.",
            "La evaluacion cuenta con evidencia de campo suficiente para documentar tanto la operacion observada como las condiciones bajo las cuales se obtuvo la muestra.",
        ],
        57: [
            "Para fines de esta evaluacion, el TEC se entiende como el intervalo transcurrido entre la incorporacion del vehiculo a la cola y la finalizacion de su atencion en la caseta correspondiente.",
            "En esta medicion, el TEC representa el tiempo que media entre el momento en que el usuario entra en espera y el cierre de la atencion en cabina de cobro.",
            "El criterio adoptado define el TEC como el tiempo total de espera desde la formacion de cola hasta la salida del vehiculo luego de ser atendido en la caseta.",
        ],
        59: [
            "Bajo el metodo de placas, el registro de campo recoge hitos temporales y operativos suficientes para reconstruir el proceso de espera y servicio de cada vehiculo observado.",
            "El metodo aplicado requiere levantar referencias temporales y de identificacion vehicular que permitan estimar el TEC por caseta y sentido de circulacion.",
            "La logica del metodo de placas consiste en registrar los hitos necesarios para seguir el desplazamiento del vehiculo dentro del sistema de cola y cobro.",
        ],
        60: [
            "Se registra la hora en la que el vehiculo se detiene para integrarse a la cola o para ingresar directamente al sistema de atencion.",
            "Como primer hito, se consigna el momento en que el vehiculo se incorpora a la espera o inicia su aproximacion a la caseta.",
            "El levantamiento considera la hora en que el usuario entra al proceso de espera, aun cuando no se forme una cola visible.",
        ],
        62: [
            "Asimismo, se toma el instante de salida de la caseta de cobro, por lo que cada cabina es tratada como una unidad independiente de servicio.",
            "Tambien se releva la hora en que culmina la atencion en la caseta, tratando cada cabina como un sistema de cola individual.",
            "El cierre del servicio se registra al momento de abandonar la caseta, asumiendo para el analisis que cada cabina opera como un sistema de cola independiente.",
        ],
        63: [
            "Con esta informacion, el metodo permite estimar el tiempo promedio ponderado por vehiculo conforme a los lineamientos contractuales y a la estructura operativa observada.",
            "A partir de los hitos levantados, se obtiene un promedio ponderado por vehiculo coherente con la definicion operacional del TEC contemplada para la evaluacion.",
            "La metodologia permite calcular el promedio ponderado de espera por vehiculo siguiendo los criterios tecnicos aplicables al indicador TEC.",
        ],
        64: [
            "Para cada registro muestreado, el TEC se determina con la diferencia entre la hora de finalizacion del servicio y la hora de inicio de la espera en cola.",
            "El calculo por vehiculo se efectua restando el momento en que concluye la atencion al instante en que el usuario inicia su espera dentro del sistema.",
            "En terminos operativos, el TEC individual resulta de comparar el hito de cierre de servicio con el punto de inicio de la espera del vehiculo.",
        ],
        65: [
            "En la practica, esta definicion permite evaluar el desempeno de cada sentido de cobro con base en la experiencia efectiva del usuario durante la jornada.",
            "Aplicado en campo, este criterio brinda una lectura directa del servicio percibido por el usuario en cada frente de cobro observado.",
            "Desde la perspectiva operativa, la metodologia facilita comparar el servicio ofrecido por cada sentido de circulacion y cada caseta evaluada.",
        ],
        69: [
            "El equipo de trabajo incluyo una jefatura de proyecto responsable de consolidar resultados, coordinar la campana y asegurar la integracion tecnica del informe final.",
            "La estructura operativa considero una jefatura de proyecto encargada de conducir la coordinacion general y articular la elaboracion del informe tecnico.",
            "La campana conto con una jefatura de proyecto que asumio la coordinacion global, el seguimiento de actividades y la preparacion del informe final.",
        ],
        72: [
            "Asimismo, se dispuso de un especialista en medicion de TEC para planificar la operacion de campo y conducir tecnicamente la toma de informacion.",
            "La labor se complemento con un especialista en TEC responsable de la organizacion metodologica y de la supervision tecnica de la muestra.",
            "El trabajo de campo fue dirigido por un especialista en TEC que estructuro la campana y acompanio tecnicamente la ejecucion de la medicion.",
        ],
        74: [
            "Entre sus funciones principales estuvo organizar y planificar las actividades previas al inicio de la captura de muestra.",
            "Dentro de sus responsabilidades se incluyo la preparacion y programacion de las tareas previas al levantamiento de informacion.",
            "Como parte de su rol, se encargo de estructurar y ordenar las labores necesarias antes del inicio de la medicion en campo.",
        ],
        76: [
            "Tambien correspondio capacitar a aforadores y supervisores para asegurar criterios uniformes de registro durante la campana.",
            "Otra funcion relevante fue entrenar al personal de campo para homogeneizar la captura de placas, tiempos y eventos de interes.",
            "El especialista tuvo ademas la tarea de instruir a los equipos de campo, con el fin de estandarizar el levantamiento de datos.",
        ],
        78: [
            "Durante la ejecucion, se brindo asistencia tecnica continua a supervisores y aforadores para resolver incidencias y mantener la calidad del dato.",
            "En campo se proporciono acompanamiento tecnico permanente a los equipos operativos para sostener la consistencia del relevamiento.",
            "La asistencia tecnica al personal desplegado permitio atender consultas y resguardar la calidad de la captura durante la muestra.",
        ],
        80: [
            "Finalmente, se superviso que las labores se desarrollaran con continuidad y normalidad, en linea con los fines del estudio.",
            "Se verifico asimismo que la campana se mantuviera continua y estable, conforme a los requerimientos tecnicos de la evaluacion.",
            "El control de continuidad de las labores fue parte central del seguimiento, para asegurar que la muestra respondiera a los objetivos del informe.",
        ],
        82: [
            "La operacion en campo se apoyo en supervisores y aforadores asignados por estacion, responsables del seguimiento directo a la captura de placas y tiempos por caseta.",
            "El despliegue de campo considero supervisores y aforadores por estacion, encargados de vigilar la toma de informacion y la cobertura de las casetas observadas.",
            "Para la ejecucion de la medicion se dispuso personal de supervision y aforo en cada estacion, con funciones directas sobre el registro operativo de la muestra.",
        ],
        85: [
            "Durante la muestra tambien se registraron eventos y observaciones complementarias que ayudan a interpretar la operacion del sistema y sus variaciones.",
            "Adicionalmente, se tomo nota de incidencias y elementos contextuales utiles para explicar el comportamiento operativo detectado en la jornada.",
            "El relevamiento considero hechos y datos adicionales que permiten contextualizar la operacion observada durante la medicion.",
        ],
        88: [
            "El procesamiento y depuracion de la informacion estuvo a cargo del componente estadistico del equipo, responsable de estructurar la base y preparar los resultados del informe.",
            "La fase de procesamiento se sostuvo en un apoyo estadistico encargado de ordenar, depurar y consolidar los registros provenientes de campo.",
            "El equipo incorporo un frente estadistico para transformar la informacion relevada en resultados trazables y comparables dentro del informe final.",
        ],
        98: [
            "Para la obtencion de la muestra se utilizo material de campo orientado a registrar de manera ordenada la placa, los hitos horarios y las incidencias asociadas al paso de cada vehiculo.",
            "El trabajo de levantamiento empleo formatos y soportes de campo disenados para capturar de forma estructurada los parametros necesarios de cada vehiculo observado.",
            "La medicion se apoyo en instrumentos de campo preparados para documentar, con criterios uniformes, la informacion operativa de cada registro vehicular.",
        ],
        99: [
            "El formato empleado permitio relevar la secuencia completa del vehiculo dentro del sistema, desde su llegada a la cola hasta la finalizacion del servicio.",
            "La planilla de campo se diseno para seguir el recorrido temporal del vehiculo y dejar trazados los momentos de espera y atencion en caseta.",
            "El instrumento de registro hizo posible capturar la trayectoria operativa del vehiculo, incluyendo ingreso al sistema, espera y cierre de atencion.",
        ],
        101: [
            "En terminos operativos, el proceso de medicion consistio en observar la llegada del vehiculo al sistema, registrar su espera y documentar la culminacion de la atencion en cabina.",
            "La medicion en campo se desarrollo siguiendo la secuencia natural del servicio: ingreso del vehiculo, permanencia en cola y salida luego de ser atendido.",
            "El proceso aplicado en campo registro la cronologia del servicio desde la incorporacion del vehiculo a la espera hasta su salida efectiva de la caseta.",
        ],
        108: [
            "En la Tabla N° 3 se presentan los promedios de TEC por caseta y sentido de circulacion. La lectura consolidada muestra como referencia mas exigente a {top_tec_caseta_text}.",
            "La Tabla N° 3 resume el Tiempo de Espera en Cola promedio por caseta y sentido. Dentro de los resultados obtenidos, resalta {top_tec_caseta_text}.",
            "Los valores consignados en la Tabla N° 3 sintetizan el TEC promedio por caseta y sentido de cobro, destacando especialmente {top_tec_caseta_text}.",
        ],
        112: [
            "Lectura ejecutiva: {caseta_threshold_summary}; el promedio maximo por caseta-sentido fue de {max_tec_caseta_value_text}, por lo que el bloque caseta se interpreta como {caseta_status_text}.",
            "Como lectura gerencial del detalle por caseta, {caseta_threshold_summary} y el punto mas exigente alcanzo {max_tec_caseta_value_text}; en conjunto, el resultado se aprecia {caseta_status_text}.",
            "En terminos ejecutivos, el analisis por caseta muestra que {caseta_threshold_summary}; el valor mas alto, {max_tec_caseta_value_text}, ubica a este bloque como {caseta_status_text}.",
        ],
        113: [
            "La Tabla N° 4 consolida el comportamiento promedio por peaje y sentido; en el conjunto evaluado, el mayor valor corresponde a {top_tec_peaje_text}.",
            "En la Tabla N° 4 se resume el TEC promedio por peaje y sentido, observandose como resultado mas alto {top_tec_peaje_text}.",
            "La lectura agregada de la Tabla N° 4 permite identificar el valor promedio mas elevado en {top_tec_peaje_text}.",
        ],
        117: [
            "Desde la perspectiva de cumplimiento, {peaje_threshold_summary}; el promedio agregado del bloque fue de {avg_tec_peaje_value_text} y la lectura global se considera {compliance_status_text}.",
            "La lectura ejecutiva por peaje-sentido indica que {peaje_threshold_summary}; con un promedio general de {avg_tec_peaje_value_text}, el resultado agregado se ubica {compliance_status_text}.",
            "En clave de gestion, {peaje_threshold_summary}; al contrastarlo con el promedio consolidado de {avg_tec_peaje_value_text}, el comportamiento agregado permanece {compliance_status_text}.",
        ],
        119: [
            "La Tabla N° 5 y las tablas siguientes muestran la distribucion de frecuencias de cola por estacion y caseta, registrandose como mayor cola observada {top_cola_text}.",
            "En las Tablas N° 5 a N° 8 se presenta la frecuencia de usuarios segun tamano de cola; dentro de la campana, la mayor referencia corresponde a {top_cola_text}.",
            "La evaluacion de cola maxima real se resume en las Tablas N° 5, 6, 7 y 8, con un valor maximo reportado de {top_cola_text}.",
        ],
        120: [
            "La interpretacion operativa de las colas indica {queue_status_text}; el maximo registrado fue de {max_cola_value_text} y el promedio de las colas maximas por caseta se ubico en {avg_cola_value_text}.",
            "Como lectura ejecutiva del bloque de colas, se aprecia {queue_status_text}; la referencia maxima fue {max_cola_value_text}, mientras que la media de cola maxima por caseta alcanzo {avg_cola_value_text}.",
            "Operativamente, el patron de cola se resume como {queue_status_text}; el valor maximo observado fue {max_cola_value_text} y la media de cola maxima por caseta fue de {avg_cola_value_text}.",
        ],
        137: [
            "Las Imagenes N° 1, 2 y 3 permiten apreciar la distribucion porcentual de los tamanos de cola observados por estacion, complementando la lectura tabular de la evaluacion.",
            "Las Imagenes N° 1, 2 y 3 ilustran la frecuencia relativa de las colas registradas en cada estacion evaluada y refuerzan la interpretacion de los resultados.",
            "La lectura grafica contenida en las Imagenes N° 1, 2 y 3 resume la estructura porcentual de las colas observadas durante la campana de medicion.",
        ],
        152: [
            "Como conclusion general, el bloque evaluado se comporto {compliance_status_text}; {peaje_threshold_summary} y el frente de mayor exigencia fue {top_tec_peaje_text}.",
            "En sintesis, la evaluacion muestra un desempeno {compliance_status_text}; {peaje_threshold_summary}, con su referencia mas alta en {top_tec_peaje_text}.",
            "A nivel global, los resultados del informe ubican al conjunto evaluado {compliance_status_text}; {peaje_threshold_summary}, siendo {top_tec_peaje_text} el punto mas demandante.",
        ],
        153: [
            "En un segundo nivel de lectura, el detalle por caseta confirma que {caseta_threshold_summary}; la mayor presion operativa se concentro en {top_tec_caseta_text}, mientras que la cola maxima observada fue {top_cola_text}.",
            "Desde una mirada operativa mas fina, el comportamiento por caseta indica que {caseta_threshold_summary}; ademas, la referencia critica del detalle fue {top_tec_caseta_text} y la mayor cola correspondio a {top_cola_text}.",
            "Profundizando en el analisis, el nivel caseta-sentido muestra que {caseta_threshold_summary}; el punto mas exigente fue {top_tec_caseta_text}, en paralelo con una cola maxima de {top_cola_text}.",
        ],
        154: [
            "En consecuencia, corresponde {recommendation_focus_text}, especialmente si se busca sostener el indicador con margen frente al limite contractual y anticipar episodios de mayor demanda.",
            "Por ello, la recomendacion inmediata es {recommendation_focus_text}, manteniendo seguimiento sobre la evolucion del TEC y de la cola maxima en las proximas mediciones.",
            "Bajo esta evidencia, resulta pertinente {recommendation_focus_text} y conservar trazabilidad especifica sobre los frentes con mayor presion de cola y espera.",
        ],
    }

    replacements = {}
    for paragraph_idx, options in variants.items():
        replacements[paragraph_idx] = pick_report_variant(seed, f"p{paragraph_idx}", options).format(
            peajes_text=facts["peajes_text"],
            total_peajes=facts["total_peajes"],
            total_casetas=facts["total_casetas"],
            total_sentidos=facts["total_sentidos"],
            total_programaciones=facts["total_programaciones"],
            total_personal_rows=facts["total_personal_rows"],
            date_source_text=facts["date_source_text"],
            top_tec_peaje_text=top_tec_peaje_text,
            top_tec_caseta_text=top_tec_caseta_text,
            top_cola_text=top_cola_text,
            max_tec_peaje_value_text=format_minutes_reference(facts["max_tec_peaje_value"]),
            avg_tec_peaje_value_text=format_minutes_reference(facts["avg_tec_peaje_value"]),
            max_tec_caseta_value_text=format_minutes_reference(facts["max_tec_caseta_value"]),
            max_cola_value_text=f"{facts['max_cola_value']} usuarios" if facts["max_cola_value"] is not None else "sin referencia suficiente",
            avg_cola_value_text=f"{facts['avg_cola_value']:.2f} usuarios" if facts["avg_cola_value"] is not None and pd.notna(facts["avg_cola_value"]) else "sin referencia suficiente",
            compliance_status_text=compliance_status_text,
            caseta_status_text=caseta_status_text,
            queue_status_text=queue_status_text,
            peaje_threshold_summary=peaje_threshold_summary,
            caseta_threshold_summary=caseta_threshold_summary,
            recommendation_focus_text=recommendation_focus_text,
        )
    return replacements


def discover_cover_logo_paths() -> list[Path | None]:
    available_images = []
    for directory in CLIENT_LOGO_DIR_CANDIDATES:
        if directory.exists() and directory.is_dir():
            available_images.extend(
                sorted(
                    [path for path in directory.iterdir() if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg"}],
                    key=lambda path: path.name.lower(),
                )
            )

    def find_logo(keywords: tuple[str, ...]) -> Path | None:
        for image_path in available_images:
            stem_key = normalize_text_key(image_path.stem)
            if any(keyword in stem_key for keyword in keywords):
                return image_path
        return None

    left_logo = find_logo(("norvial", "concesionario", "cliente1", "izq", "left"))
    center_logo = find_logo(("ositran", "regulador", "cliente2", "centro", "center"))
    right_logo = find_logo(("cidatt", "consult", "cliente3", "der", "right"))
    if right_logo is None and CONTRACTOR_LOGO_PATH.exists():
        right_logo = CONTRACTOR_LOGO_PATH
    return [left_logo, center_logo, right_logo]


def update_cover_logos(doc: Document) -> None:
    if not doc.paragraphs:
        return
    cover_paragraph = doc.paragraphs[0]
    rel_ids = []
    for blip in cover_paragraph._p.xpath('.//a:blip'):
        rel_id = blip.get(qn('r:embed'))
        if rel_id:
            rel_ids.append(rel_id)
    if not rel_ids:
        return

    for rel_id, logo_path in zip(rel_ids, discover_cover_logo_paths()):
        if logo_path is None or not logo_path.exists():
            continue
        replace_related_image(doc, rel_id, logo_path.read_bytes())


def apply_dynamic_template_narrative(doc: Document, report_label: str, informe_package: dict[str, object]) -> None:
    for paragraph_idx, replacement in build_dynamic_report_paragraphs(report_label, informe_package).items():
        replace_docx_paragraph_at_index(doc, paragraph_idx, replacement)


def classify_peaje_bucket(peaje: object) -> str:
    peaje_key = normalize_text_key(peaje)
    if "paraiso" in peaje_key:
        return "paraiso"
    if "variante" in peaje_key:
        return "variante"
    if "serpentin" in peaje_key and "pesaje" in peaje_key:
        return "serpentin_pesaje"
    if "serpentin" in peaje_key:
        return "serpentin"
    return peaje_key or "otros"


def build_frequency_table_for_bucket(df_resultados: pd.DataFrame, bucket: str) -> pd.DataFrame:
    column_map = {
        "paraiso": lambda caseta: f"PARAISO {int(caseta)}",
        "variante": lambda caseta: f"VARIANTE {int(caseta)}",
        "serpentin": lambda caseta: f"Serpentin Caseta {int(caseta)}",
        "serpentin_pesaje": lambda caseta: f"Serpentin Pesaje Caseta {int(caseta)}",
    }
    formatter = column_map.get(bucket, lambda caseta: f"Caseta {int(caseta)}")

    sub_df = df_resultados[df_resultados["PEAJE_BUCKET"] == bucket].copy()
    if sub_df.empty:
        return pd.DataFrame(columns=["Cantidad de usuarios en la cola", "Total"])

    tabla = pd.pivot_table(
        sub_df,
        index="COLA_ESPERA_USUARIOS",
        columns="CASETA",
        values="PLACA_FINAL",
        aggfunc="size",
        fill_value=0,
    ).sort_index()
    tabla.columns = [formatter(columna) for columna in tabla.columns]
    tabla["Total"] = tabla.sum(axis=1)
    tabla.loc["Total general"] = tabla.sum(axis=0)
    return tabla.reset_index().rename(columns={"COLA_ESPERA_USUARIOS": "Cantidad de usuarios en la cola"})


def format_template_peaje_name(bucket: str) -> str:
    names = {
        "serpentin": "Serpentin de Pasamayo",
        "variante": "Variante de Pasamayo",
        "paraiso": "Paraiso (Huacho)",
        "serpentin_pesaje": "Serpentin de Pasamayo (Pesaje)",
    }
    return names.get(bucket, bucket.upper())


def format_template_sentido_name(value: object) -> str:
    key = normalize_text_key(value)
    if "asc" in key:
        return "Ascendente"
    if "desc" in key:
        return "Descendente"
    return str(value).title() if value is not None else ""


def format_template_date(value: pd.Timestamp) -> str:
    return value.strftime("%d.%m.%Y")


def format_template_day_name(value: pd.Timestamp) -> str:
    dias = {
        0: "Lunes",
        1: "Martes",
        2: "Miercoles",
        3: "Jueves",
        4: "Viernes",
        5: "Sabado",
        6: "Domingo",
    }
    return dias[value.weekday()]


def format_template_timerange(start_value: pd.Timedelta, end_value: pd.Timedelta) -> str:
    if pd.isna(start_value) or pd.isna(end_value):
        return ""

    def fmt(td: pd.Timedelta) -> str:
        total_seconds = int(td.total_seconds()) % (24 * 3600)
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        return f"{hours:02d}:{minutes:02d}"

    return f"{fmt(start_value)} a {fmt(end_value)}"


def blank_repeated_first_column(tabla: pd.DataFrame, column_name: str) -> pd.DataFrame:
    tabla = tabla.copy()
    previous = None
    values = []
    for value in tabla[column_name]:
        if previous == value:
            values.append("")
        else:
            values.append(value)
            previous = value
    tabla[column_name] = values
    return tabla


def build_medicion_programada_table(df_resultados: pd.DataFrame) -> pd.DataFrame:
    columnas = ["Peaje", "Fecha", "Dia", "Rango Horario", "Sentido"]
    if df_resultados.empty:
        return pd.DataFrame(columns=columnas)

    base = df_resultados.copy()
    if "LLEGADA_COLA_FINAL_TD" not in base.columns:
        base["LLEGADA_COLA_FINAL_TD"] = pd.to_timedelta(base["LLEGADA_COLA_FINAL"].astype(str), errors="coerce")
    if "SALIDA_CASETA_FINAL_TD" not in base.columns:
        base["SALIDA_CASETA_FINAL_TD"] = pd.to_timedelta(base["SALIDA_CASETA_FINAL"].astype(str), errors="coerce")
    base["FECHA_DT"] = pd.to_datetime(base["FECHA"], errors="coerce")

    tabla = (
        base.groupby(["PEAJE_BUCKET", "FECHA_DT", "SENTIDO"], dropna=False)
        .agg(hora_inicio=("LLEGADA_COLA_FINAL_TD", "min"), hora_fin=("SALIDA_CASETA_FINAL_TD", "max"))
        .reset_index()
    )
    tabla = tabla[tabla["FECHA_DT"].notna()].copy()
    if tabla.empty:
        return pd.DataFrame(columns=columnas)

    bucket_order = {"serpentin": 0, "variante": 1, "paraiso": 2, "serpentin_pesaje": 3}
    tabla["_bucket_order"] = tabla["PEAJE_BUCKET"].map(lambda value: bucket_order.get(value, 99))
    tabla["_sentido_order"] = tabla["SENTIDO"].map(lambda value: 0 if "desc" in normalize_text_key(value) else 1)
    tabla = tabla.sort_values(["_bucket_order", "FECHA_DT", "_sentido_order"], ascending=[True, False, True])

    tabla_out = pd.DataFrame(
        {
            "Peaje": tabla["PEAJE_BUCKET"].map(format_template_peaje_name),
            "Fecha": tabla["FECHA_DT"].map(format_template_date),
            "Dia": tabla["FECHA_DT"].map(format_template_day_name),
            "Rango Horario": [format_template_timerange(start, end) for start, end in zip(tabla["hora_inicio"], tabla["hora_fin"])],
            "Sentido": tabla["SENTIDO"].map(format_template_sentido_name),
        }
    )
    return blank_repeated_first_column(tabla_out, "Peaje")


def build_personal_asignado_table(df_resultados: pd.DataFrame) -> pd.DataFrame:
    columnas = ["Peaje", "Cantidad de Casetas", "Supervisores", "Aforadores", "Sentido"]
    if df_resultados.empty:
        return pd.DataFrame(columns=columnas)

    tabla = (
        df_resultados.groupby(["PEAJE_BUCKET", "SENTIDO"], dropna=False)
        .agg(casetas=("CASETA", pd.Series.nunique))
        .reset_index()
    )
    if tabla.empty:
        return pd.DataFrame(columns=columnas)

    bucket_order = {"serpentin": 0, "variante": 1, "paraiso": 2, "serpentin_pesaje": 3}
    tabla["_bucket_order"] = tabla["PEAJE_BUCKET"].map(lambda value: bucket_order.get(value, 99))
    tabla["_sentido_order"] = tabla["SENTIDO"].map(lambda value: 1 if "desc" in normalize_text_key(value) else 0)
    tabla = tabla.sort_values(["_bucket_order", "_sentido_order"])

    tabla_out = pd.DataFrame(
        {
            "Peaje": tabla["PEAJE_BUCKET"].map(format_template_peaje_name),
            "Cantidad de Casetas": tabla["casetas"].astype(int).astype(str),
            "Supervisores": "1",
            "Aforadores": (tabla["casetas"].astype(int) * 2).astype(str),
            "Sentido": tabla["SENTIDO"].map(format_template_sentido_name),
        }
    )
    return blank_repeated_first_column(tabla_out, "Peaje")


def build_frequency_percentage_series(df_resultados: pd.DataFrame, buckets: list[str]) -> pd.DataFrame:
    sub_df = df_resultados[df_resultados["PEAJE_BUCKET"].isin(buckets)].copy()
    if sub_df.empty:
        return pd.DataFrame(columns=["cola", "pct"])

    serie = (
        sub_df.groupby("COLA_ESPERA_USUARIOS", dropna=False)
        .size()
        .rename("total")
        .reset_index()
        .sort_values("COLA_ESPERA_USUARIOS")
    )
    total = int(serie["total"].sum())
    if total <= 0:
        return pd.DataFrame(columns=["cola", "pct"])

    serie["pct"] = 100 * serie["total"] / total
    serie["cola"] = serie["COLA_ESPERA_USUARIOS"].astype(int)
    return serie[["cola", "pct"]]


def render_frequency_chart_bytes(title: str, percentage_series: pd.DataFrame, accent: str) -> bytes | None:
    if percentage_series.empty:
        return None

    fig, ax = plt.subplots(figsize=(8.2, 4.6), dpi=180)
    ax.bar(
        percentage_series["cola"].astype(str),
        percentage_series["pct"],
        color=accent,
        edgecolor="#16315f",
        linewidth=0.6,
    )
    ax.set_title(title, fontsize=12, fontweight="bold", color="#16315f", pad=14)
    ax.set_xlabel("Cantidad de usuarios en cola", fontsize=9)
    ax.set_ylabel("Frecuencia (%)", fontsize=9)
    ax.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.35)
    ax.set_axisbelow(True)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(axis="x", labelsize=8)
    ax.tick_params(axis="y", labelsize=8)

    for idx, value in enumerate(percentage_series["pct"]):
        ax.text(idx, value + 0.4, f"{value:.1f}%", ha="center", va="bottom", fontsize=7, color="#16315f")

    fig.tight_layout()
    output = BytesIO()
    fig.savefig(output, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    output.seek(0)
    return output.getvalue()


def build_report_date_range_text(df_resultados: pd.DataFrame) -> str:
    if df_resultados is None or getattr(df_resultados, "empty", True) or "FECHA" not in df_resultados.columns:
        return "Fuente: Base procesada en la aplicacion TEC."

    fechas = pd.to_datetime(df_resultados.get("FECHA"), errors="coerce").dropna().sort_values().dt.normalize().unique()
    if len(fechas) == 0:
        return "Fuente: Base procesada en la aplicacion TEC."

    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
        7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }

    def fmt(fecha):
        return f"{fecha.day:02d} de {meses[fecha.month]} de {fecha.year}"

    if len(fechas) == 1:
        return f"Fuente: Labores de campo realizadas el {fmt(pd.Timestamp(fechas[0]))}"

    return (
        "Fuente: Labores de campo realizadas entre "
        f"{fmt(pd.Timestamp(fechas[0]))} y {fmt(pd.Timestamp(fechas[-1]))}"
    )


def build_template_table3_text(informe_package: dict[str, object]) -> str:
    tabla_tec_caseta = informe_package["tabla_tec_caseta"]
    if tabla_tec_caseta.empty:
        return (
            "En la Tabla N° 3 no se identificaron resultados suficientes para resumir el Tiempo de Espera en Cola por caseta y sentido de circulacion."
        )

    fila = tabla_tec_caseta.sort_values("Tiempo de Espera en Cola - TEC", ascending=False).iloc[0]
    return (
        "En la Tabla N° 3, se presentan los resultados obtenidos a partir de los promedios de Tiempo de Espera en Cola "
        "por vehiculo ponderado por el numero de vehiculos atendidos. Los resultados se presentan por caseta y sentido de circulacion, "
        f"observandose el mayor promedio en {fila['Peaje']} - caseta {fila['Caseta Controlada']} - {fila['Sentido de Circulacion']}, "
        f"con {float(fila['Tiempo de Espera en Cola - TEC']):.2f} minutos."
    )


def build_template_queue_text(informe_package: dict[str, object]) -> str:
    tabla_cola = informe_package["tabla_cola_maxima"]
    if tabla_cola.empty:
        return (
            "En la Tabla N° 5 se muestra la frecuencia de usuarios segun el tamano de cola por peaje y por casetas durante los periodos de evaluacion."
        )

    fila = tabla_cola.sort_values("Cola maxima real", ascending=False).iloc[0]
    return (
        "En la Tabla N° 5 se muestra la frecuencia de usuarios segun el tamano de cola por peaje y por casetas, durante los periodos de evaluacion. "
        "La cola de espera real para cada caseta de peaje se presenta en las Tablas N° 5, 6, 7 y 8; "
        f"destacando como mayor valor observado {int(fila['Cola maxima real'])} usuarios en {fila['Peaje']} - caseta {fila['Caseta Controlada']} - {fila['Sentido de Circulacion']}."
    )


def replace_result_section_sources(doc: Document, source_text: str) -> None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx >= 105 and "Fuente:" in paragraph.text:
            paragraph.text = source_text


def replace_related_image(doc: Document, rel_id: str, image_bytes: bytes) -> None:
    image_part = doc.part.related_parts[rel_id]
    image_part._blob = image_bytes


def update_result_section_graphs(doc: Document, informe_package: dict[str, object]) -> None:
    df_resultados = informe_package.get("df_resultados")
    if df_resultados is None or getattr(df_resultados, "empty", True):
        return

    chart_specs = {
        139: {
            "buckets": ["paraiso"],
            "title": "Imagen N° 1. Frecuencias de cola en vehiculos observadas en la estacion de peaje Paraiso",
            "accent": "#245cc6",
        },
        143: {
            "buckets": ["serpentin", "serpentin_pesaje"],
            "title": "Imagen N° 2. Frecuencias de cola en vehiculos observadas en la estacion de peaje y pesaje Serpentin",
            "accent": "#1849a9",
        },
        147: {
            "buckets": ["variante"],
            "title": "Imagen N° 3. Frecuencias de cola en vehiculos observadas en la estacion de peaje Variante",
            "accent": "#3b7be6",
        },
    }

    for paragraph_idx, spec in chart_specs.items():
        if paragraph_idx >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[paragraph_idx]
        rel_ids = []
        for blip in paragraph._p.xpath('.//a:blip'):
            rel_id = blip.get(qn('r:embed'))
            if rel_id:
                rel_ids.append(rel_id)
        if not rel_ids:
            continue
        series = build_frequency_percentage_series(df_resultados, spec["buckets"])
        chart_bytes = render_frequency_chart_bytes(spec["title"], series, spec["accent"])
        if chart_bytes is None:
            continue
        replace_related_image(doc, rel_ids[0], chart_bytes)


def build_template_compliance_text(informe_package: dict[str, object]) -> str:
    tabla_tec_peaje = informe_package["tabla_tec_peaje"]
    if tabla_tec_peaje.empty:
        return (
            "En la Tabla N° 4 no se cuenta con resultados suficientes para evaluar el cumplimiento del TEC "
            "por peaje/pesaje segun sentido."
        )

    fila_critica = tabla_tec_peaje.sort_values("Tiempo de Espera en Cola - TEC", ascending=False).iloc[0]
    valor_critico = float(fila_critica["Tiempo de Espera en Cola - TEC"])
    if valor_critico <= 3:
        return (
            "En la Tabla N° 4, observamos que los resultados promedios por cada peaje/pesaje "
            "segun sentido no superan los 3 minutos promedio para la atencion de los usuarios."
        )

    return (
        "En la Tabla N° 4, observamos que los resultados promedios por cada peaje/pesaje segun sentido "
        f"muestran que {fila_critica['Peaje']} - {fila_critica['Sentido de Circulacion']} alcanza {valor_critico:.2f} minutos promedio, "
        "superando los 3 minutos promedio para la atencion de los usuarios."
    )


def build_template_conclusion_text(informe_package: dict[str, object]) -> str:
    tabla_tec_peaje = informe_package["tabla_tec_peaje"]

    if tabla_tec_peaje.empty:
        return (
            "No se cuenta con resultados suficientes en la base procesada para emitir conclusiones con el modelo de informe."
        )

    fila_critica = tabla_tec_peaje.sort_values("Tiempo de Espera en Cola - TEC", ascending=False).iloc[0]
    valor_critico = float(fila_critica["Tiempo de Espera en Cola - TEC"])
    if valor_critico <= 3:
        return (
            "De la tabla N° 4 del presente Informe, se evidencia que los valores obtenidos por los niveles de servicio "
            "que el Concesionario tiene actualmente se encuentran dentro de los niveles que deben ser cumplidos, debido a que, "
            "durante el periodo muestreado en cada unidad de peaje/pesaje el promedio de la muestra no supero los 3 minutos; "
            "lo cual no amerita por ahora la modificacion del sistema de atencion."
        )

    return (
        "De la tabla N° 4 del presente Informe, se evidencia que los valores obtenidos por los niveles de servicio "
        f"superan el umbral esperado en {fila_critica['Peaje']} - {fila_critica['Sentido de Circulacion']}, donde el promedio "
        f"alcanzado fue de {valor_critico:.2f} minutos; por lo que corresponde revisar las condiciones operativas del sistema de atencion."
    )


def to_templated_docx_bytes(report_label: str, informe_package: dict[str, object]) -> bytes:
    template_path = resolve_informe_template_path()
    if template_path is None:
        return to_docx_bytes(report_label, informe_package)

    doc = Document(template_path)
    update_cover_logos(doc)
    table_mapping = {
        0: informe_package["tabla_programacion"],
        1: informe_package["tabla_personal"],
        3: informe_package["tabla_tec_caseta"],
        4: informe_package["tabla_tec_peaje"],
        5: informe_package["tabla_frecuencia_paraiso"],
        6: informe_package["tabla_frecuencia_serpentin"],
        7: informe_package["tabla_frecuencia_serpentin_pesaje"],
        8: informe_package["tabla_frecuencia_variante"],
    }

    for table_idx, tabla in table_mapping.items():
        if table_idx < len(doc.tables):
            write_dataframe_to_existing_docx_table(doc.tables[table_idx], tabla)

    source_text = build_report_date_range_text(informe_package.get("df_resultados", pd.DataFrame()))
    replace_result_section_sources(doc, source_text)

    replace_docx_paragraph_contains(
        doc,
        "en la tabla n 3, se presenta los resultados obtenidos",
        build_template_table3_text(informe_package),
    )
    replace_docx_paragraph_contains(
        doc,
        "los resultados promedios por cada peaje/pesaje segun sentido",
        build_template_compliance_text(informe_package),
    )
    replace_docx_paragraph_contains(
        doc,
        "en la tabla n 5 se muestra la frecuencia de usuarios",
        build_template_queue_text(informe_package),
    )
    replace_docx_paragraph_contains(
        doc,
        "de la tabla n 4 del presente informe",
        build_template_conclusion_text(informe_package),
    )
    apply_dynamic_template_narrative(doc, report_label, informe_package)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def q85(serie: pd.Series) -> float:
    return float(serie.quantile(0.85))


def q95(serie: pd.Series) -> float:
    serie_numerica = pd.to_numeric(serie, errors="coerce").dropna()
    if serie_numerica.empty:
        return float("nan")
    return float(serie_numerica.quantile(0.95))


def coerce_numeric_result_columns(df_in: pd.DataFrame) -> pd.DataFrame:
    df = df_in.copy()
    numeric_columns = [
        "T_COLA_FINAL_SEGUNDOS",
        "T_COLA_FINAL_MINUTOS",
        "T_CASETA_FINAL_SEGUNDOS",
        "T_CASETA_FINAL_MINUTOS",
        "T_TEC_FINAL_SEGUNDOS",
        "T_TEC_FINAL_MINUTOS",
        "COLA_ESPERA_USUARIOS",
    ]
    for column in numeric_columns:
        if column in df.columns:
            df[column] = pd.to_numeric(df[column], errors="coerce")
    return df


def resumir_metricas(df_in: pd.DataFrame, group_cols: list[str], value_col: str, label: str) -> pd.DataFrame:
    df_metricas = df_in.copy()
    df_metricas[value_col] = pd.to_numeric(df_metricas[value_col], errors="coerce")
    tabla = (
        df_metricas.groupby(group_cols, dropna=False)[value_col]
        .agg(
            casos="size",
            promedio="mean",
            mediana="median",
            p85=q85,
            p95=q95,
            maximo="max",
        )
        .reset_index()
    )
    tabla.insert(len(group_cols), "indicador", label)
    tabla[["promedio", "mediana", "p85", "p95", "maximo"]] = tabla[
        ["promedio", "mediana", "p85", "p95", "maximo"]
    ].round(2)
    return tabla


def build_resultados_dataframe(df_export_base: pd.DataFrame) -> pd.DataFrame:
    df_resultados_base = df_export_base[df_export_base["TIEMPOS_COMPLETOS_CIERRE"]].copy()
    if df_resultados_base.empty:
        df_resultados_base["COLA_ESPERA_USUARIOS"] = pd.Series(dtype="int64")
        return coerce_numeric_result_columns(df_resultados_base)

    for columna in ["LLEGADA_COLA_FINAL", "LLEGADA_CASETA_FINAL", "SALIDA_CASETA_FINAL"]:
        df_resultados_base[f"{columna}_TD"] = pd.to_timedelta(df_resultados_base[columna].astype(str), errors="coerce")

    grupos_resultados = [
        calcular_cola_espera_real(sub_df)
        for _, sub_df in df_resultados_base.groupby(["PEAJE", "CASETA", "SENTIDO", "FECHA"], dropna=False)
    ]
    if not grupos_resultados:
        df_resultados_base["COLA_ESPERA_USUARIOS"] = pd.Series(dtype="int64")
        return coerce_numeric_result_columns(df_resultados_base)

    df_resultados = pd.concat(grupos_resultados, axis=0).sort_values(
        ["PEAJE", "SENTIDO", "CASETA", "FECHA", "LLEGADA_COLA_FINAL_TD", "_ORDEN_FILA"]
    ).reset_index(drop=True)
    return coerce_numeric_result_columns(df_resultados)


def build_informe_package(df_export_base: pd.DataFrame) -> dict[str, object]:
    df_resultados = build_resultados_dataframe(df_export_base)

    tabla_programacion_informe = pd.DataFrame(columns=["Peaje", "Fecha", "Dia", "Rango Horario", "Sentido"])
    tabla_personal_informe = pd.DataFrame(columns=["Peaje", "Cantidad de Casetas", "Supervisores", "Aforadores", "Sentido"])

    tabla_tec_caseta_informe = pd.DataFrame(
        columns=["Peaje", "Caseta Controlada", "Sentido de Circulacion", "Tiempo de Espera en Cola - TEC", "Unidades de tiempo"]
    )
    tabla_tec_peaje_informe = pd.DataFrame(
        columns=["Peaje", "Sentido de Circulacion", "Tiempo de Espera en Cola - TEC", "Unidades de tiempo"]
    )
    tabla_cola_maxima_informe = pd.DataFrame(
        columns=["Peaje", "Caseta Controlada", "Sentido de Circulacion", "Cola maxima real", "Vehiculos evaluados"]
    )
    tabla_frecuencia_paraiso_informe = pd.DataFrame(columns=["Cantidad de usuarios en la cola", "Total"])
    tabla_frecuencia_serpentin_informe = pd.DataFrame(columns=["Cantidad de usuarios en la cola", "Total"])
    tabla_frecuencia_serpentin_pesaje_informe = pd.DataFrame(columns=["Cantidad de usuarios en la cola", "Total"])
    tabla_frecuencia_variante_informe = pd.DataFrame(columns=["Cantidad de usuarios en la cola", "Total"])
    resumen_narrativo = pd.DataFrame({"Texto sugerido para informe": ["No hay datos suficientes para generar resultados de informe."]})

    if not df_resultados.empty:
        df_resultados = coerce_numeric_result_columns(df_resultados)
        df_resultados["PEAJE_BUCKET"] = df_resultados["PEAJE"].map(classify_peaje_bucket)
        tabla_programacion_informe = build_medicion_programada_table(df_resultados)
        tabla_personal_informe = build_personal_asignado_table(df_resultados)

        tabla_tec_caseta = (
            df_resultados.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
            .agg(
                TEC_MINUTOS=("T_TEC_FINAL_MINUTOS", "mean"),
                VEHICULOS=("PLACA_FINAL", "size"),
            )
            .reset_index()
            .sort_values(["PEAJE", "SENTIDO", "CASETA"])
        )
        tabla_tec_caseta["TEC_MINUTOS"] = tabla_tec_caseta["TEC_MINUTOS"].round(2)
        tabla_tec_caseta_informe = tabla_tec_caseta.rename(
            columns={
                "PEAJE": "Peaje",
                "CASETA": "Caseta Controlada",
                "SENTIDO": "Sentido de Circulacion",
                "TEC_MINUTOS": "Tiempo de Espera en Cola - TEC",
            }
        )[
            ["Peaje", "Caseta Controlada", "Sentido de Circulacion", "Tiempo de Espera en Cola - TEC"]
        ].copy()
        tabla_tec_caseta_informe["Unidades de tiempo"] = "Minutos"

        tabla_tec_peaje = (
            df_resultados.groupby(["PEAJE", "SENTIDO"], dropna=False)
            .agg(
                TEC_MINUTOS=("T_TEC_FINAL_MINUTOS", "mean"),
                VEHICULOS=("PLACA_FINAL", "size"),
            )
            .reset_index()
            .sort_values(["PEAJE", "SENTIDO"])
        )
        tabla_tec_peaje["TEC_MINUTOS"] = tabla_tec_peaje["TEC_MINUTOS"].round(2)
        tabla_tec_peaje_informe = tabla_tec_peaje.rename(
            columns={
                "PEAJE": "Peaje",
                "SENTIDO": "Sentido de Circulacion",
                "TEC_MINUTOS": "Tiempo de Espera en Cola - TEC",
            }
        )[
            ["Peaje", "Sentido de Circulacion", "Tiempo de Espera en Cola - TEC"]
        ].copy()
        tabla_tec_peaje_informe["Unidades de tiempo"] = "Minutos"

        tabla_cola_maxima = (
            df_resultados.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
            .agg(
                COLA_MAXIMA_REAL=("COLA_ESPERA_USUARIOS", "max"),
                VEHICULOS=("PLACA_FINAL", "size"),
            )
            .reset_index()
            .sort_values(["PEAJE", "SENTIDO", "CASETA"])
        )
        tabla_cola_maxima_informe = tabla_cola_maxima.rename(
            columns={
                "PEAJE": "Peaje",
                "CASETA": "Caseta Controlada",
                "SENTIDO": "Sentido de Circulacion",
                "COLA_MAXIMA_REAL": "Cola maxima real",
                "VEHICULOS": "Vehiculos evaluados",
            }
        )

        tabla_frecuencia_paraiso_informe = formatear_tabla_frecuencia(build_frequency_table_for_bucket(df_resultados, "paraiso"))
        tabla_frecuencia_serpentin_informe = formatear_tabla_frecuencia(build_frequency_table_for_bucket(df_resultados, "serpentin"))
        tabla_frecuencia_serpentin_pesaje_informe = formatear_tabla_frecuencia(
            build_frequency_table_for_bucket(df_resultados, "serpentin_pesaje")
        )
        tabla_frecuencia_variante_informe = formatear_tabla_frecuencia(build_frequency_table_for_bucket(df_resultados, "variante"))

        fila_mayor_tec_caseta = tabla_tec_caseta.sort_values("TEC_MINUTOS", ascending=False).iloc[0]
        fila_mayor_tec_peaje = tabla_tec_peaje.sort_values("TEC_MINUTOS", ascending=False).iloc[0]
        fila_mayor_cola = tabla_cola_maxima.sort_values("COLA_MAXIMA_REAL", ascending=False).iloc[0]
        resumen_narrativo = pd.DataFrame(
            {
                "Texto sugerido para informe": [
                    (
                        "El mayor TEC promedio por peaje y sentido se observo en "
                        f"{fila_mayor_tec_peaje['PEAJE']} - {fila_mayor_tec_peaje['SENTIDO']}, "
                        f"con {fila_mayor_tec_peaje['TEC_MINUTOS']:.2f} minutos."
                    ),
                    (
                        "La caseta con mayor TEC promedio fue "
                        f"{fila_mayor_tec_caseta['PEAJE']} - caseta {int(fila_mayor_tec_caseta['CASETA'])} - "
                        f"{fila_mayor_tec_caseta['SENTIDO']}, con {fila_mayor_tec_caseta['TEC_MINUTOS']:.2f} minutos."
                    ),
                    (
                        "La mayor cola de espera real se observo en "
                        f"{fila_mayor_cola['PEAJE']} - caseta {int(fila_mayor_cola['CASETA'])} - "
                        f"{fila_mayor_cola['SENTIDO']}, con {int(fila_mayor_cola['COLA_MAXIMA_REAL'])} usuarios en cola."
                    ),
                    (
                        "Las tablas de frecuencia por tamano de cola se calcularon contando solo a los usuarios "
                        "que aun no llegan a la caseta al momento en que arriba un nuevo vehiculo."
                    ),
                ]
            }
        )

    excel_sheets = {
        "programacion_medicion": tabla_programacion_informe,
        "personal_asignado": tabla_personal_informe,
        "tec_caseta": tabla_tec_caseta_informe,
        "tec_peaje_sentido": tabla_tec_peaje_informe,
        "cola_maxima_real": tabla_cola_maxima_informe,
        "cola_paraiso": tabla_frecuencia_paraiso_informe,
        "cola_serpentin": tabla_frecuencia_serpentin_informe,
        "cola_serpentin_pesaje": tabla_frecuencia_serpentin_pesaje_informe,
        "cola_variante": tabla_frecuencia_variante_informe,
        "texto_informe": resumen_narrativo,
    }
    return {
        "df_resultados": df_resultados,
        "excel_sheets": excel_sheets,
        "tabla_programacion": tabla_programacion_informe,
        "tabla_personal": tabla_personal_informe,
        "tabla_tec_caseta": tabla_tec_caseta_informe,
        "tabla_tec_peaje": tabla_tec_peaje_informe,
        "tabla_cola_maxima": tabla_cola_maxima_informe,
        "tabla_frecuencia_paraiso": tabla_frecuencia_paraiso_informe,
        "tabla_frecuencia_serpentin": tabla_frecuencia_serpentin_informe,
        "tabla_frecuencia_serpentin_pesaje": tabla_frecuencia_serpentin_pesaje_informe,
        "tabla_frecuencia_variante": tabla_frecuencia_variante_informe,
        "texto_informe": resumen_narrativo,
    }


def build_complementary_package(
    df_resultados: pd.DataFrame,
    df_export_base: pd.DataFrame,
    df_original: pd.DataFrame,
    fugas_flujo: pd.DataFrame,
    fragmentaciones_probables: pd.DataFrame,
) -> dict[str, object]:
    resumen_complementario = pd.DataFrame(columns=["indicador", "valor"])
    descriptivos_peaje_sentido = pd.DataFrame(columns=["PEAJE", "SENTIDO", "indicador", "casos", "promedio", "mediana", "p85", "p95", "maximo"])
    descriptivos_caseta = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "indicador", "casos", "promedio", "mediana", "p85", "p95", "maximo"])
    cumplimiento_3min_peaje = pd.DataFrame(columns=["PEAJE", "SENTIDO", "vehiculos", "dentro_3_min", "fuera_3_min", "tec_promedio_min", "tec_p95_min", "tec_max_min", "pct_dentro_3_min"])
    cumplimiento_3min_caseta = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "vehiculos", "dentro_3_min", "fuera_3_min", "tec_promedio_min", "tec_p95_min", "tec_max_min", "pct_dentro_3_min"])
    fugas_por_peaje = pd.DataFrame(columns=["PEAJE", "REGISTROS_DATA", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"])
    fugas_por_caseta = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "REGISTROS_DATA", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"])
    top_casetas_tec = pd.DataFrame(columns=descriptivos_caseta.columns)
    top_bloques_30min = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "BLOQUE_30MIN", "vehiculos", "tec_promedio_min", "tec_p95_min", "cola_promedio_veh", "cola_maxima_veh"])
    acciones_por_peaje = pd.DataFrame(columns=["PEAJE", "ACCION_REALIZADA", "filas"])
    texto_sugerido_extra = pd.DataFrame({"Texto sugerido": ["No hay datos suficientes para generar resultados complementarios."]})
    fuga_rate_tables = build_fuga_rate_tables(df_original, fugas_flujo, fragmentaciones_probables)
    fugas_por_peaje = fuga_rate_tables["by_peaje"].copy()
    fugas_por_caseta = fuga_rate_tables["by_caseta"].copy()

    if not df_resultados.empty:
        df_resultados_extra = df_resultados.copy()
        df_resultados_extra["FECHA_DT"] = pd.to_datetime(df_resultados_extra["FECHA"], errors="coerce")
        df_resultados_extra["LLEGADA_COLA_TIMESTAMP"] = (
            df_resultados_extra["FECHA_DT"] + df_resultados_extra["LLEGADA_COLA_FINAL_TD"]
        )
        df_resultados_extra["BLOQUE_30MIN"] = (
            df_resultados_extra["LLEGADA_COLA_TIMESTAMP"].dt.floor("30min").dt.strftime("%Y-%m-%d %H:%M")
        )

        resumen_complementario = pd.DataFrame(
            {
                "indicador": [
                    "vehiculos_base_limpia",
                    "registros_ajustados",
                    "porcentaje_ajustados",
                    "tec_promedio_global_min",
                    "tec_mediana_global_min",
                    "tec_p95_global_min",
                    "cola_promedio_global_veh",
                    "cola_p95_global_veh",
                    "cola_maxima_global_veh",
                ],
                "valor": [
                    len(df_resultados_extra),
                    int(df_export_base["REGISTRO_AJUSTADO"].sum()),
                    round(100 * df_export_base["REGISTRO_AJUSTADO"].mean(), 2),
                    round(df_resultados_extra["T_TEC_FINAL_MINUTOS"].mean(), 2),
                    round(df_resultados_extra["T_TEC_FINAL_MINUTOS"].median(), 2),
                    round(df_resultados_extra["T_TEC_FINAL_MINUTOS"].quantile(0.95), 2),
                    round(df_resultados_extra["COLA_ESPERA_USUARIOS"].mean(), 2),
                    round(df_resultados_extra["COLA_ESPERA_USUARIOS"].quantile(0.95), 2),
                    int(df_resultados_extra["COLA_ESPERA_USUARIOS"].max()),
                ],
            }
        )

        descriptivos_peaje_sentido = pd.concat(
            [
                resumir_metricas(df_resultados_extra, ["PEAJE", "SENTIDO"], "T_TEC_FINAL_MINUTOS", "T_TEC_MINUTOS"),
                resumir_metricas(df_resultados_extra, ["PEAJE", "SENTIDO"], "T_COLA_FINAL_MINUTOS", "T_COLA_MINUTOS"),
                resumir_metricas(df_resultados_extra, ["PEAJE", "SENTIDO"], "T_CASETA_FINAL_MINUTOS", "T_CASETA_MINUTOS"),
            ],
            ignore_index=True,
        ).sort_values(["PEAJE", "SENTIDO", "indicador"])

        descriptivos_caseta = resumir_metricas(
            df_resultados_extra,
            ["PEAJE", "CASETA", "SENTIDO"],
            "T_TEC_FINAL_MINUTOS",
            "T_TEC_MINUTOS",
        ).sort_values(["promedio", "p95", "maximo"], ascending=[False, False, False])

        cumplimiento_3min_peaje = (
            df_resultados_extra.groupby(["PEAJE", "SENTIDO"], dropna=False)
            .agg(
                vehiculos=("PLACA_FINAL", "size"),
                dentro_3_min=("T_TEC_FINAL_MINUTOS", lambda s: int((s <= 3).sum())),
                fuera_3_min=("T_TEC_FINAL_MINUTOS", lambda s: int((s > 3).sum())),
                tec_promedio_min=("T_TEC_FINAL_MINUTOS", "mean"),
                tec_p95_min=("T_TEC_FINAL_MINUTOS", q95),
                tec_max_min=("T_TEC_FINAL_MINUTOS", "max"),
            )
            .reset_index()
        )
        cumplimiento_3min_peaje["pct_dentro_3_min"] = (
            100 * cumplimiento_3min_peaje["dentro_3_min"] / cumplimiento_3min_peaje["vehiculos"]
        ).round(2)
        cumplimiento_3min_peaje[["tec_promedio_min", "tec_p95_min", "tec_max_min"]] = cumplimiento_3min_peaje[
            ["tec_promedio_min", "tec_p95_min", "tec_max_min"]
        ].round(2)
        cumplimiento_3min_peaje["PEAJE"] = cumplimiento_3min_peaje["PEAJE"].map(format_dashboard_dimension)
        cumplimiento_3min_peaje["SENTIDO"] = cumplimiento_3min_peaje["SENTIDO"].map(format_dashboard_dimension)
        cumplimiento_3min_peaje = cumplimiento_3min_peaje.merge(
            fugas_por_peaje[["PEAJE", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]],
            on=["PEAJE"],
            how="left",
        )
        cumplimiento_3min_peaje[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]] = cumplimiento_3min_peaje[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]].fillna(0)
        cumplimiento_3min_peaje["FUGAS_FLUJO"] = cumplimiento_3min_peaje["FUGAS_FLUJO"].astype(int)
        cumplimiento_3min_peaje["FRAGMENTACIONES_FLUJO"] = cumplimiento_3min_peaje["FRAGMENTACIONES_FLUJO"].astype(int)

        cumplimiento_3min_caseta = (
            df_resultados_extra.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
            .agg(
                vehiculos=("PLACA_FINAL", "size"),
                dentro_3_min=("T_TEC_FINAL_MINUTOS", lambda s: int((s <= 3).sum())),
                fuera_3_min=("T_TEC_FINAL_MINUTOS", lambda s: int((s > 3).sum())),
                tec_promedio_min=("T_TEC_FINAL_MINUTOS", "mean"),
                tec_p95_min=("T_TEC_FINAL_MINUTOS", q95),
                tec_max_min=("T_TEC_FINAL_MINUTOS", "max"),
            )
            .reset_index()
        )
        cumplimiento_3min_caseta["pct_dentro_3_min"] = (
            100 * cumplimiento_3min_caseta["dentro_3_min"] / cumplimiento_3min_caseta["vehiculos"]
        ).round(2)
        cumplimiento_3min_caseta[["tec_promedio_min", "tec_p95_min", "tec_max_min"]] = cumplimiento_3min_caseta[
            ["tec_promedio_min", "tec_p95_min", "tec_max_min"]
        ].round(2)
        cumplimiento_3min_caseta["PEAJE"] = cumplimiento_3min_caseta["PEAJE"].map(format_dashboard_dimension)
        cumplimiento_3min_caseta["CASETA"] = cumplimiento_3min_caseta["CASETA"].map(format_dashboard_dimension)
        cumplimiento_3min_caseta["SENTIDO"] = cumplimiento_3min_caseta["SENTIDO"].map(format_dashboard_dimension)
        cumplimiento_3min_caseta = cumplimiento_3min_caseta.merge(
            fugas_por_caseta[["PEAJE", "CASETA", "SENTIDO", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]],
            on=["PEAJE", "CASETA", "SENTIDO"],
            how="left",
        )
        cumplimiento_3min_caseta[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]] = cumplimiento_3min_caseta[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]].fillna(0)
        cumplimiento_3min_caseta["FUGAS_FLUJO"] = cumplimiento_3min_caseta["FUGAS_FLUJO"].astype(int)
        cumplimiento_3min_caseta["FRAGMENTACIONES_FLUJO"] = cumplimiento_3min_caseta["FRAGMENTACIONES_FLUJO"].astype(int)

        top_casetas_tec = descriptivos_caseta.head(10).copy()

        top_bloques_30min = (
            df_resultados_extra.groupby(["PEAJE", "CASETA", "SENTIDO", "BLOQUE_30MIN"], dropna=False)
            .agg(
                vehiculos=("PLACA_FINAL", "size"),
                tec_promedio_min=("T_TEC_FINAL_MINUTOS", "mean"),
                tec_p95_min=("T_TEC_FINAL_MINUTOS", q95),
                cola_promedio_veh=("COLA_ESPERA_USUARIOS", "mean"),
                cola_maxima_veh=("COLA_ESPERA_USUARIOS", "max"),
            )
            .reset_index()
        )
        top_bloques_30min = top_bloques_30min[top_bloques_30min["vehiculos"] >= 5].copy()
        top_bloques_30min[["tec_promedio_min", "tec_p95_min", "cola_promedio_veh"]] = top_bloques_30min[
            ["tec_promedio_min", "tec_p95_min", "cola_promedio_veh"]
        ].round(2)
        top_bloques_30min = top_bloques_30min.sort_values(
            ["tec_promedio_min", "cola_maxima_veh", "vehiculos"],
            ascending=[False, False, False],
        ).head(30)

        acciones_por_peaje = (
            df_export_base.groupby(["PEAJE", "ACCION_REALIZADA"], dropna=False)
            .size()
            .rename("filas")
            .reset_index()
            .sort_values(["PEAJE", "filas"], ascending=[True, False])
        )

        top_peaje = cumplimiento_3min_peaje.sort_values("tec_promedio_min", ascending=False).iloc[0]
        top_caseta = top_casetas_tec.iloc[0]
        texto_sugerido_extra = pd.DataFrame(
            {
                "Texto sugerido": [
                    (
                        "El "
                        f"{round(100 * df_export_base['REGISTRO_AJUSTADO'].mean(), 2):.2f}% "
                        "de los registros de la base limpia requirio algun ajuste o imputacion."
                    ),
                    (
                        "El peaje/sentido con mayor tiempo de espera promedio fue "
                        f"{top_peaje['PEAJE']} - {top_peaje['SENTIDO']}."
                    ),
                    (
                        "La caseta mas critica por TEC promedio fue "
                        f"{top_caseta['PEAJE']} - caseta {int(top_caseta['CASETA'])} - {top_caseta['SENTIDO']}."
                    ),
                    (
                        "Los bloques de 30 minutos permiten identificar periodos puntuales de mayor congestion "
                        "y priorizar revisiones operativas por caseta."
                    ),
                ]
            }
        )

    excel_sheets = {
        "resumen": resumen_complementario,
        "desc_peaje_sentido": descriptivos_peaje_sentido,
        "desc_caseta_tec": descriptivos_caseta,
        "cumpl_3min_peaje": cumplimiento_3min_peaje,
        "cumpl_3min_caseta": cumplimiento_3min_caseta,
        "fugas_peaje": fugas_por_peaje,
        "fugas_caseta": fugas_por_caseta,
        "top_casetas_tec": top_casetas_tec,
        "top_bloques_30m": top_bloques_30min,
        "acciones_por_peaje": acciones_por_peaje,
        "texto_sugerido": texto_sugerido_extra,
    }
    return {
        "excel_sheets": excel_sheets,
        "top_casetas_tec": top_casetas_tec,
        "top_bloques_30m": top_bloques_30min,
    }


def to_excel_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    header_fill = PatternFill(fill_type="solid", fgColor="163564")
    header_font = Font(color="FFFFFF", bold=True)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_alignment = Alignment(vertical="top", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="D7E1F0"),
        right=Side(style="thin", color="D7E1F0"),
        top=Side(style="thin", color="D7E1F0"),
        bottom=Side(style="thin", color="D7E1F0"),
    )

    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for sheet_name, df_sheet in sheets.items():
            safe_name = sheet_name[:31]
            df_sheet.to_excel(writer, sheet_name=safe_name, index=False)
            worksheet = writer.book[safe_name]
            worksheet.freeze_panes = "A2"
            worksheet.auto_filter.ref = worksheet.dimensions
            worksheet.sheet_view.showGridLines = False

            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = header_alignment
                cell.border = thin_border

            for row in worksheet.iter_rows(min_row=2):
                for cell in row:
                    cell.alignment = body_alignment
                    cell.border = thin_border

            for col_idx, column_name in enumerate(df_sheet.columns, start=1):
                series = df_sheet[column_name]
                value_lengths = [len(str(column_name))]
                if not series.empty:
                    value_lengths.extend(series.map(measure_text_width).tolist())
                adjusted_width = min(max(value_lengths) + 2, 50)
                adjusted_width = max(adjusted_width, 10)
                worksheet.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

            worksheet.row_dimensions[1].height = 24
    output.seek(0)
    return output.getvalue()


def write_report_block(
    writer: pd.ExcelWriter,
    sheet_name: str,
    title: str,
    table: pd.DataFrame,
    startrow: int,
) -> int:
    pd.DataFrame({title: []}).to_excel(writer, sheet_name=sheet_name, startrow=startrow, index=False)
    worksheet = writer.sheets[sheet_name]
    worksheet.cell(row=startrow + 1, column=1, value=title)
    table.to_excel(writer, sheet_name=sheet_name, startrow=startrow + 2, index=False)
    return startrow + len(table) + 5


def measure_text_width(value: object) -> int:
    if value is None:
        return 0
    try:
        is_missing = pd.isna(value)
        if isinstance(is_missing, (bool, np.bool_)) and is_missing:
            return 0
    except Exception:
        pass
    text = str(value)
    lines = text.splitlines() or [text]
    return max((len(part) for part in lines), default=0)


def apply_standard_worksheet_format(worksheet, df_sheet: pd.DataFrame | None = None, header_row: int = 1, freeze_panes: str = "A2") -> None:
    header_fill = PatternFill(fill_type="solid", fgColor="163564")
    header_font = Font(color="FFFFFF", bold=True)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_alignment = Alignment(vertical="top", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="D7E1F0"),
        right=Side(style="thin", color="D7E1F0"),
        top=Side(style="thin", color="D7E1F0"),
        bottom=Side(style="thin", color="D7E1F0"),
    )

    worksheet.freeze_panes = freeze_panes
    worksheet.auto_filter.ref = worksheet.dimensions
    worksheet.sheet_view.showGridLines = False

    for cell in worksheet[header_row]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment
        cell.border = thin_border

    for row in worksheet.iter_rows(min_row=header_row + 1):
        for cell in row:
            cell.alignment = body_alignment
            cell.border = thin_border

    if df_sheet is not None:
        for col_idx, column_name in enumerate(df_sheet.columns, start=1):
            series = df_sheet[column_name]
            value_lengths = [len(str(column_name))]
            if not series.empty:
                value_lengths.extend(series.map(measure_text_width).tolist())
            adjusted_width = min(max(value_lengths) + 2, 50)
            adjusted_width = max(adjusted_width, 10)
            worksheet.column_dimensions[get_column_letter(col_idx)].width = adjusted_width
    else:
        for col_idx in range(1, worksheet.max_column + 1):
            max_length = 10
            for row in worksheet.iter_rows(min_col=col_idx, max_col=col_idx, values_only=True):
                value = row[0]
                if value is None:
                    continue
                max_length = max(max_length, measure_text_width(value) + 2)
            worksheet.column_dimensions[get_column_letter(col_idx)].width = min(max_length, 50)

    worksheet.row_dimensions[header_row].height = 24


def apply_report_summary_format(worksheet) -> None:
    title_fill = PatternFill(fill_type="solid", fgColor="0F3D91")
    header_fill = PatternFill(fill_type="solid", fgColor="163564")
    title_font = Font(color="FFFFFF", bold=True)
    header_font = Font(color="FFFFFF", bold=True)
    title_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    body_alignment = Alignment(vertical="top", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="D7E1F0"),
        right=Side(style="thin", color="D7E1F0"),
        top=Side(style="thin", color="D7E1F0"),
        bottom=Side(style="thin", color="D7E1F0"),
    )

    worksheet.sheet_view.showGridLines = False
    worksheet.freeze_panes = "A3"

    previous_blank = True
    for row_idx in range(1, worksheet.max_row + 1):
        values = [worksheet.cell(row=row_idx, column=col_idx).value for col_idx in range(1, worksheet.max_column + 1)]
        non_empty = [value for value in values if value not in (None, "")]
        if not non_empty:
            previous_blank = True
            continue

        is_title_row = len(non_empty) == 1 and values[0] not in (None, "")
        is_header_row = previous_blank and len(non_empty) > 1

        for col_idx in range(1, worksheet.max_column + 1):
            cell = worksheet.cell(row=row_idx, column=col_idx)
            cell.border = thin_border
            if is_title_row:
                cell.fill = title_fill
                cell.font = title_font
                cell.alignment = title_alignment
            elif is_header_row:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = header_alignment
            else:
                cell.alignment = body_alignment

        if is_title_row:
            worksheet.row_dimensions[row_idx].height = 22
        previous_blank = False

    for col_idx in range(1, worksheet.max_column + 1):
        max_length = 10
        for row in worksheet.iter_rows(min_col=col_idx, max_col=col_idx, values_only=True):
            value = row[0]
            if value is None:
                continue
            max_length = max(max_length, measure_text_width(value) + 2)
        worksheet.column_dimensions[get_column_letter(col_idx)].width = min(max_length, 50)


def to_exact_excel_bytes(export_package: dict[str, pd.DataFrame]) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_package["base_limpia"].to_excel(writer, sheet_name="base_limpia", index=False)
        export_package["casos_eliminados"].to_excel(writer, sheet_name="casos_eliminados", index=False)
        export_package["casos_pendientes"].to_excel(writer, sheet_name="casos_pendientes", index=False)
        export_package["revision_placas"].to_excel(writer, sheet_name="revision_placas", index=False)
        export_package["bloques_decision"].to_excel(writer, sheet_name="bloques_decision", index=False)
        export_package["fugas_flujo"].to_excel(writer, sheet_name="fugas_flujo", index=False)
        export_package["fragmentaciones_probables"].to_excel(writer, sheet_name="fragmentaciones", index=False)

        startrow = 0
        startrow = write_report_block(
            writer,
            "reporte_resumen",
            "Resumen general",
            export_package["resumen_general"],
            startrow,
        )
        startrow = write_report_block(
            writer,
            "reporte_resumen",
            "Resumen de acciones en base_limpia",
            export_package["resumen_accion_realizada"],
            startrow,
        )
        startrow = write_report_block(
            writer,
            "reporte_resumen",
            "Resumen de acciones de placa",
            export_package["resumen_acciones_placa"],
            startrow,
        )
        startrow = write_report_block(
            writer,
            "reporte_resumen",
            "Resumen de acciones de tiempo",
            export_package["resumen_acciones_tiempo"],
            startrow,
        )
        write_report_block(
            writer,
            "reporte_resumen",
            "Resumen de casos eliminados",
            export_package["resumen_eliminados"],
            startrow,
        )
        apply_standard_worksheet_format(writer.book["base_limpia"], export_package["base_limpia"])
        apply_standard_worksheet_format(writer.book["casos_eliminados"], export_package["casos_eliminados"])
        apply_standard_worksheet_format(writer.book["casos_pendientes"], export_package["casos_pendientes"])
        apply_standard_worksheet_format(writer.book["revision_placas"], export_package["revision_placas"])
        apply_standard_worksheet_format(writer.book["bloques_decision"], export_package["bloques_decision"])
        apply_standard_worksheet_format(writer.book["fugas_flujo"], export_package["fugas_flujo"])
        apply_standard_worksheet_format(writer.book["fragmentaciones"], export_package["fragmentaciones_probables"])
        apply_report_summary_format(writer.book["reporte_resumen"])
    output.seek(0)
    return output.getvalue()


def to_docx_bytes(report_label: str, informe_package: dict[str, object]) -> bytes:
    doc = Document()
    doc.add_heading(f"Tablas de resultados {report_label}", level=1)
    doc.add_paragraph(
        "Base utilizada: registros finales limpios. La cola de espera se define como los usuarios "
        "que aun esperan llegar a la caseta cuando arriba un nuevo vehiculo."
    )
    agregar_tabla_docx(doc, "Tabla 1. Tiempo de Espera en Cola por caseta y sentido", informe_package["tabla_tec_caseta"])
    agregar_tabla_docx(doc, "Tabla 2. Promedio de tiempo de espera en cola por peaje segun sentido", informe_package["tabla_tec_peaje"])
    agregar_tabla_docx(doc, "Tabla 3. Cola maxima de espera real por caseta", informe_package["tabla_cola_maxima"])
    agregar_tabla_docx(
        doc,
        "Tabla 4. Frecuencia de usuarios segun tamano de cola de espera para la estacion de peaje Paraiso por caseta",
        informe_package["tabla_frecuencia_paraiso"],
    )
    agregar_tabla_docx(
        doc,
        "Tabla 5. Frecuencia de usuarios segun tamano de cola de espera para la estacion de peaje Variante por caseta",
        informe_package["tabla_frecuencia_variante"],
    )
    doc.add_paragraph("Texto de apoyo para conclusiones", style="Heading 2")
    for texto in informe_package["texto_informe"]["Texto sugerido para informe"]:
        doc.add_paragraph(texto, style="List Bullet")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def to_zip_bytes(sheets: dict[str, pd.DataFrame], excel_bytes: bytes) -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("procesamiento_tec_output.xlsx", excel_bytes)
        for name, df_sheet in sheets.items():
            zf.writestr(f"{name}.csv", df_sheet.to_csv(index=False).encode("utf-8-sig"))
    buffer.seek(0)
    return buffer.getvalue()


def build_run_payload(
    source_name: str,
    mapping: dict[str, str | None],
    config: dict,
    export_tables: dict[str, pd.DataFrame],
) -> dict:
    return {
        "source_name": source_name,
        "input_rows": int(export_tables["reporte_resumen"].loc[
            export_tables["reporte_resumen"]["indicador"] == "filas_entrada", "valor"
        ].iloc[0]),
        "clean_rows": len(export_tables["base_limpia"]),
        "deleted_rows": len(export_tables["casos_eliminados"]),
        "pending_rows": len(export_tables["casos_pendientes"]),
        "mapping": {k: v for k, v in mapping.items() if v},
        "config": config,
        "notes": {
            "sheets": list(export_tables.keys()),
        },
    }


def format_dashboard_dimension(value: object) -> str:
    if value is None or pd.isna(value):
        return "Sin dato"
    text = str(value).strip()
    return text if text else "Sin dato"


def format_dashboard_action(value: object) -> str:
    text = format_dashboard_dimension(value)
    if text == "Sin dato":
        return text
    text = text.replace("placa:", "Placa: ")
    text = text.replace("tiempo:", "Tiempo: ")
    text = text.replace("eliminado:", "Eliminado: ")
    text = text.replace("_", " ")
    return text.capitalize()


def prepare_dashboard_dataframe(df_in: pd.DataFrame) -> pd.DataFrame:
    df = df_in.copy()
    for column in ["PEAJE", "CASETA", "SENTIDO", "PLACA"]:
        if column not in df.columns:
            df[column] = pd.NA
        df[column] = df[column].map(format_dashboard_dimension)
    if "FECHA" not in df.columns:
        df["FECHA"] = pd.NaT
    df["FECHA_DIA_DASH"] = pd.to_datetime(df["FECHA"], errors="coerce").dt.strftime("%Y-%m-%d")
    df["FECHA_DIA_DASH"] = df["FECHA_DIA_DASH"].fillna("Sin fecha")
    return df


def detect_raw_fugas(df_in: pd.DataFrame) -> pd.DataFrame:
    df = prepare_dashboard_dataframe(df_in)
    placa_norm = (
        df["PLACA"]
        .fillna("")
        .astype(str)
        .str.upper()
        .str.replace(r"[^A-Z0-9]", "", regex=True)
    )
    fugas = df.loc[placa_norm.str.endswith("X") & (placa_norm.str.len() > 6)].copy()
    fugas["PLACA_NORMALIZADA"] = placa_norm.loc[fugas.index]
    fugas["LONGITUD_NORMALIZADA"] = fugas["PLACA_NORMALIZADA"].str.len()
    fugas = fugas[
        [
            "PEAJE",
            "CASETA",
            "SENTIDO",
            "FECHA_DIA_DASH",
            "PLACA",
            "PLACA_NORMALIZADA",
            "LONGITUD_NORMALIZADA",
        ]
    ].rename(columns={"FECHA_DIA_DASH": "FECHA"})
    return fugas.sort_values(["PEAJE", "SENTIDO", "CASETA", "FECHA", "PLACA_NORMALIZADA"]).reset_index(drop=True)


def canonicalize_plate_for_fuga(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    text = re.sub(r"[^A-Z0-9]", "", str(value).upper())
    translated = []
    for char in text:
        if char in {"0", "O"}:
            translated.append("O")
        elif char in {"1", "I"}:
            translated.append("I")
        elif char in {"5", "S"}:
            translated.append("S")
        elif char in {"8", "B"}:
            translated.append("B")
        elif char in {"2", "Z"}:
            translated.append("Z")
        elif char in {"6", "G"}:
            translated.append("G")
        else:
            translated.append(char)
    return "".join(translated)


def are_plates_similar_for_fuga(plate_a: object, plate_b: object) -> bool:
    canon_a = canonicalize_plate_for_fuga(plate_a)
    canon_b = canonicalize_plate_for_fuga(plate_b)
    if not canon_a or not canon_b:
        return False
    if canon_a == canon_b:
        return True
    if len(canon_a) == len(canon_b):
        diff = sum(char_a != char_b for char_a, char_b in zip(canon_a, canon_b))
        if diff <= 1:
            return True
        if diff == 2 and sorted(canon_a) == sorted(canon_b):
            return True
    return distancia_levenshtein(canon_a, canon_b) <= 1


def classify_fragment_similarity(plate_a: object, plate_b: object, delta_seg: float) -> tuple[bool, str]:
    canon_a = canonicalize_plate_for_fuga(plate_a)
    canon_b = canonicalize_plate_for_fuga(plate_b)
    if not canon_a or not canon_b:
        return False, "sin_placa"
    if canon_a == canon_b:
        return True, "alta"

    if len(canon_a) == len(canon_b):
        diff = sum(char_a != char_b for char_a, char_b in zip(canon_a, canon_b))
        if diff <= 1:
            return True, "alta" if delta_seg <= 90 else "media"
        if diff == 2 and sorted(canon_a) == sorted(canon_b):
            return True, "media"

    if distancia_levenshtein(canon_a, canon_b) <= 1:
        return True, "media"
    return False, "baja"


def classify_fuga_confidence(
    score: int,
    ubicacion: str,
    base_type: str,
    apariciones_contexto: int,
    apariciones_totales: int,
    eventos_patron_placa: int,
) -> tuple[str, str]:
    if ubicacion != "interior":
        return "incompleto_no_concluyente_borde", "baja"
    strong_recurrence = apariciones_contexto >= 2 or apariciones_totales >= 3 or eventos_patron_placa >= 2
    if score >= 9 and strong_recurrence:
        return f"fuga_fuerte_{base_type}", "alta"
    if score >= 5:
        return f"fuga_probable_{base_type}", "media"
    return "incompleto_no_concluyente_interior", "baja"


def build_fuga_detail(score: int, eventos_patron_placa: int, tasa_placa: float, tasa_contexto: float) -> str:
    return (
        f"score={score}; recurrencias_patron={eventos_patron_placa}; "
        f"tasa_placa={tasa_placa:.2%}; tasa_contexto={tasa_contexto:.2%}"
    )


def detect_flow_fuga_candidates(time_result: dict[str, pd.DataFrame]) -> dict[str, pd.DataFrame]:
    df = time_result.get("df_tiempos_analisis_flujo", time_result.get("df_tiempos_final", time_result["df_tiempos_bordes"])).copy()
    if df.empty:
        empty_fugas = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FECHA", "PLACA_FINAL", "TIPO_FUGA", "NIVEL_CONFIANZA", "ES_FUGA_FUERTE", "ES_FUGA_PROBABLE", "ES_INCOMPLETO_NO_CONCLUYENTE", "SCORE_FUGA", "EVENTOS_PATRON_PLACA", "TASA_ANOMALIA_PLACA", "TASA_ANOMALIA_CONTEXTO", "DETALLE"])
        empty_fragmentos = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FECHA", "PLACA_T1", "PLACA_T2T3", "DELTA_SEG", "CONFIANZA", "ES_ALTA_CONFIANZA", "DETALLE"])
        return {"fugas_probables": empty_fugas, "fragmentaciones_probables": empty_fragmentos}

    t1_source = "LLEGADA_COLA_FINAL" if "LLEGADA_COLA_FINAL" in df.columns else "LLEGADA COLA"
    t2_source = "LLEGADA_CASETA_FINAL" if "LLEGADA_CASETA_FINAL" in df.columns else "LLEGADA CASETA"
    t3_source = "SALIDA_CASETA_FINAL" if "SALIDA_CASETA_FINAL" in df.columns else "SALIDA CASETA"
    df["T1_DASH"] = df[t1_source].map(normalizar_hora)
    df["T2_DASH"] = df[t2_source].map(normalizar_hora)
    df["T3_DASH"] = df[t3_source].map(normalizar_hora)
    df["FECHA_DIA_DASH"] = pd.to_datetime(df["FECHA"], errors="coerce").dt.normalize()
    if "TIEMPO_REFERENCIA" not in df.columns:
        df["TIEMPO_REFERENCIA"] = df["T1_DASH"].combine_first(df["T2_DASH"]).combine_first(df["T3_DASH"])
    df = df.sort_values(
        ["PEAJE", "CASETA", "SENTIDO", "FECHA_DIA_DASH", "TIEMPO_REFERENCIA", "T1_DASH", "T2_DASH", "T3_DASH", "_ORDEN_FILA"],
        na_position="last",
    ).copy()

    fugas_candidatas = []
    fragmentaciones_probables = []

    plate_context_counts = (
        df.groupby(["PEAJE", "SENTIDO", "PLACA_FINAL"], dropna=False)
        .size()
        .rename("APARICIONES_PLACA_CONTEXTO")
        .reset_index()
    )
    plate_total_counts = (
        df.groupby(["PLACA_FINAL"], dropna=False)
        .size()
        .rename("APARICIONES_PLACA_TOTAL")
        .reset_index()
    )
    context_totals = (
        df.groupby(["PEAJE", "SENTIDO"], dropna=False)
        .size()
        .rename("TOTAL_CONTEXTO")
        .reset_index()
    )

    for keys, group in df.groupby(["PEAJE", "CASETA", "SENTIDO", "FECHA_DIA_DASH"], dropna=False):
        group = group.reset_index(drop=True)
        complete_mask = group["T1_DASH"].notna() & group["T2_DASH"].notna() & group["T3_DASH"].notna()
        paired_rows = set()
        for idx, row in group.iterrows():
            if idx in paired_rows or complete_mask.iloc[idx]:
                continue

            prev_complete = complete_mask.iloc[:idx].any()
            next_complete = complete_mask.iloc[idx + 1 :].any()
            ubicacion = "interior" if prev_complete and next_complete else "borde"
            tiene_t1 = pd.notna(row["T1_DASH"])
            tiene_t2 = pd.notna(row["T2_DASH"])
            tiene_t3 = pd.notna(row["T3_DASH"])

            if tiene_t1 and (not tiene_t2) and (not tiene_t3):
                pair_found = False
                if ubicacion == "interior":
                    for look_ahead in range(idx + 1, min(idx + 5, len(group))):
                        if look_ahead in paired_rows:
                            continue
                        other = group.iloc[look_ahead]
                        other_is_fragment = pd.isna(other["T1_DASH"]) and pd.notna(other["T2_DASH"]) and pd.notna(other["T3_DASH"])
                        if not other_is_fragment:
                            continue
                        delta_seg = (other["T2_DASH"] - row["T1_DASH"]).total_seconds()
                        if delta_seg < 0 or delta_seg > 180:
                            continue
                        is_match, confidence = classify_fragment_similarity(row.get("PLACA_FINAL"), other.get("PLACA_FINAL"), delta_seg)
                        if not is_match:
                            continue
                        pair_found = True
                        paired_rows.add(look_ahead)
                        fragmentaciones_probables.append(
                            {
                                "PEAJE": format_dashboard_dimension(keys[0]),
                                "CASETA": format_dashboard_dimension(keys[1]),
                                "SENTIDO": format_dashboard_dimension(keys[2]),
                                "FECHA": pd.to_datetime(keys[3]).strftime("%Y-%m-%d") if pd.notna(keys[3]) else "Sin fecha",
                                "PLACA_T1": format_dashboard_dimension(row.get("PLACA_FINAL")),
                                "PLACA_T2T3": format_dashboard_dimension(other.get("PLACA_FINAL")),
                                "DELTA_SEG": int(delta_seg),
                                "CONFIANZA": confidence,
                                "DETALLE": "Registro partido: cola en una fila y caseta/salida en otra fila muy cercana.",
                            }
                        )
                        break

                if not pair_found:
                    fugas_candidatas.append(
                        {
                            "PEAJE_RAW": keys[0],
                            "SENTIDO_RAW": keys[2],
                            "PEAJE": format_dashboard_dimension(keys[0]),
                            "CASETA": format_dashboard_dimension(keys[1]),
                            "SENTIDO": format_dashboard_dimension(keys[2]),
                            "FECHA": pd.to_datetime(keys[3]).strftime("%Y-%m-%d") if pd.notna(keys[3]) else "Sin fecha",
                            "PLACA_FINAL_RAW": row.get("PLACA_FINAL"),
                            "PLACA_FINAL": format_dashboard_dimension(row.get("PLACA_FINAL")),
                            "UBICACION": ubicacion,
                            "BASE_TYPE": "no_inicia_caseta",
                        }
                    )
            elif tiene_t1 and tiene_t2 and (not tiene_t3):
                fugas_candidatas.append(
                    {
                        "PEAJE_RAW": keys[0],
                        "SENTIDO_RAW": keys[2],
                        "PEAJE": format_dashboard_dimension(keys[0]),
                        "CASETA": format_dashboard_dimension(keys[1]),
                        "SENTIDO": format_dashboard_dimension(keys[2]),
                        "FECHA": pd.to_datetime(keys[3]).strftime("%Y-%m-%d") if pd.notna(keys[3]) else "Sin fecha",
                        "PLACA_FINAL_RAW": row.get("PLACA_FINAL"),
                        "PLACA_FINAL": format_dashboard_dimension(row.get("PLACA_FINAL")),
                        "UBICACION": ubicacion,
                        "BASE_TYPE": "no_finaliza_caseta",
                    }
                )

    df_fugas = pd.DataFrame(fugas_candidatas)
    if not df_fugas.empty:
        eventos_patron = (
            df_fugas.groupby(["PEAJE_RAW", "SENTIDO_RAW", "PLACA_FINAL_RAW", "BASE_TYPE"], dropna=False)
            .size()
            .rename("EVENTOS_PATRON_PLACA")
            .reset_index()
        )
        eventos_contexto = (
            df_fugas.groupby(["PEAJE_RAW", "SENTIDO_RAW"], dropna=False)
            .size()
            .rename("EVENTOS_CONTEXTO")
            .reset_index()
        )
        df_fugas = df_fugas.merge(plate_context_counts, left_on=["PEAJE_RAW", "SENTIDO_RAW", "PLACA_FINAL_RAW"], right_on=["PEAJE", "SENTIDO", "PLACA_FINAL"], how="left")
        df_fugas = df_fugas.drop(columns=["PEAJE_y", "SENTIDO_y", "PLACA_FINAL_y"], errors="ignore").rename(columns={"PEAJE_x": "PEAJE", "SENTIDO_x": "SENTIDO", "PLACA_FINAL_x": "PLACA_FINAL"})
        df_fugas = df_fugas.merge(plate_total_counts, left_on=["PLACA_FINAL_RAW"], right_on=["PLACA_FINAL"], how="left")
        df_fugas = df_fugas.drop(columns=["PLACA_FINAL"], errors="ignore").rename(columns={"PLACA_FINAL_x": "PLACA_FINAL"}) if "PLACA_FINAL_x" in df_fugas.columns else df_fugas
        df_fugas = df_fugas.merge(eventos_patron, on=["PEAJE_RAW", "SENTIDO_RAW", "PLACA_FINAL_RAW", "BASE_TYPE"], how="left")
        df_fugas = df_fugas.merge(eventos_contexto, on=["PEAJE_RAW", "SENTIDO_RAW"], how="left")
        df_fugas = df_fugas.merge(context_totals, left_on=["PEAJE_RAW", "SENTIDO_RAW"], right_on=["PEAJE", "SENTIDO"], how="left")
        df_fugas = df_fugas.drop(columns=["PEAJE_y", "SENTIDO_y"], errors="ignore").rename(columns={"PEAJE_x": "PEAJE", "SENTIDO_x": "SENTIDO"})

        for column in ["APARICIONES_PLACA_CONTEXTO", "APARICIONES_PLACA_TOTAL", "EVENTOS_PATRON_PLACA", "EVENTOS_CONTEXTO", "TOTAL_CONTEXTO"]:
            df_fugas[column] = df_fugas[column].fillna(0).astype(int)

        df_fugas["TASA_ANOMALIA_PLACA"] = df_fugas.apply(
            lambda row: (row["EVENTOS_PATRON_PLACA"] / row["APARICIONES_PLACA_CONTEXTO"]) if row["APARICIONES_PLACA_CONTEXTO"] else 0.0,
            axis=1,
        )
        df_fugas["TASA_ANOMALIA_CONTEXTO"] = df_fugas.apply(
            lambda row: (row["EVENTOS_CONTEXTO"] / row["TOTAL_CONTEXTO"]) if row["TOTAL_CONTEXTO"] else 0.0,
            axis=1,
        )

        def score_row(row: pd.Series) -> int:
            score = 0
            if row["UBICACION"] == "interior":
                score += 3
            else:
                score -= 3
            score += 2
            if row["APARICIONES_PLACA_CONTEXTO"] >= 2:
                score += 3
            if row["APARICIONES_PLACA_TOTAL"] >= 3:
                score += 2
            if row["EVENTOS_PATRON_PLACA"] >= 2:
                score += 3
            if row["TASA_ANOMALIA_PLACA"] >= 0.4:
                score += 2
            if row["TASA_ANOMALIA_PLACA"] >= (2 * row["TASA_ANOMALIA_CONTEXTO"]):
                score += 2
            return score

        df_fugas["SCORE_FUGA"] = df_fugas.apply(score_row, axis=1)
        clasificaciones = df_fugas.apply(
            lambda row: classify_fuga_confidence(
                int(row["SCORE_FUGA"]),
                str(row["UBICACION"]),
                str(row["BASE_TYPE"]),
                int(row["APARICIONES_PLACA_CONTEXTO"]),
                int(row["APARICIONES_PLACA_TOTAL"]),
                int(row["EVENTOS_PATRON_PLACA"]),
            ),
            axis=1,
            result_type="expand",
        )
        clasificaciones.columns = ["TIPO_FUGA", "NIVEL_CONFIANZA"]
        df_fugas[["TIPO_FUGA", "NIVEL_CONFIANZA"]] = clasificaciones
        df_fugas["ES_FUGA_FUERTE"] = df_fugas["TIPO_FUGA"].astype(str).str.startswith("fuga_fuerte_")
        df_fugas["ES_FUGA_PROBABLE"] = df_fugas["TIPO_FUGA"].astype(str).str.startswith("fuga_probable_")
        df_fugas["ES_INCOMPLETO_NO_CONCLUYENTE"] = df_fugas["TIPO_FUGA"].astype(str).str.startswith("incompleto_no_concluyente")
        df_fugas["DETALLE"] = df_fugas.apply(
            lambda row: build_fuga_detail(
                int(row["SCORE_FUGA"]),
                int(row["EVENTOS_PATRON_PLACA"]),
                float(row["TASA_ANOMALIA_PLACA"]),
                float(row["TASA_ANOMALIA_CONTEXTO"]),
            ),
            axis=1,
        )
        df_fugas = df_fugas[
            [
                "PEAJE",
                "CASETA",
                "SENTIDO",
                "FECHA",
                "PLACA_FINAL",
                "TIPO_FUGA",
                "NIVEL_CONFIANZA",
                "ES_FUGA_FUERTE",
                "ES_FUGA_PROBABLE",
                "ES_INCOMPLETO_NO_CONCLUYENTE",
                "SCORE_FUGA",
                "EVENTOS_PATRON_PLACA",
                "TASA_ANOMALIA_PLACA",
                "TASA_ANOMALIA_CONTEXTO",
                "DETALLE",
            ]
        ].sort_values(["TIPO_FUGA", "SCORE_FUGA", "PEAJE", "CASETA", "FECHA", "PLACA_FINAL"], ascending=[True, False, True, True, True, True]).reset_index(drop=True)
    else:
        df_fugas = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FECHA", "PLACA_FINAL", "TIPO_FUGA", "NIVEL_CONFIANZA", "ES_FUGA_FUERTE", "ES_FUGA_PROBABLE", "ES_INCOMPLETO_NO_CONCLUYENTE", "SCORE_FUGA", "EVENTOS_PATRON_PLACA", "TASA_ANOMALIA_PLACA", "TASA_ANOMALIA_CONTEXTO", "DETALLE"])

    df_fragmentos = pd.DataFrame(fragmentaciones_probables)
    if not df_fragmentos.empty:
        df_fragmentos["ES_ALTA_CONFIANZA"] = df_fragmentos["CONFIANZA"].eq("alta")
        df_fragmentos = df_fragmentos.sort_values(["PEAJE", "CASETA", "FECHA", "DELTA_SEG"]).reset_index(drop=True)
    else:
        df_fragmentos = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FECHA", "PLACA_T1", "PLACA_T2T3", "DELTA_SEG", "CONFIANZA", "ES_ALTA_CONFIANZA", "DETALLE"])

    return {
        "fugas_probables": df_fugas,
        "fragmentaciones_probables": df_fragmentos,
    }


def build_volume_tables(df_in: pd.DataFrame) -> dict[str, pd.DataFrame]:
    df = prepare_dashboard_dataframe(df_in)
    total = max(len(df), 1)

    por_peaje = df.groupby("PEAJE", dropna=False).size().rename("REGISTROS").reset_index()
    por_peaje["PARTICIPACION_%"] = (por_peaje["REGISTROS"] / total * 100).round(2)
    por_peaje = por_peaje.sort_values(["REGISTROS", "PEAJE"], ascending=[False, True]).reset_index(drop=True)

    por_peaje_sentido = df.groupby(["PEAJE", "SENTIDO"], dropna=False).size().rename("REGISTROS").reset_index()
    por_peaje_sentido = por_peaje_sentido.sort_values(
        ["PEAJE", "REGISTROS", "SENTIDO"],
        ascending=[True, False, True],
    ).reset_index(drop=True)

    por_caseta = df.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False).size().rename("REGISTROS").reset_index()
    por_caseta["PARTICIPACION_%"] = (por_caseta["REGISTROS"] / total * 100).round(2)
    por_caseta["ETIQUETA"] = por_caseta["PEAJE"] + " | C" + por_caseta["CASETA"] + " | " + por_caseta["SENTIDO"]
    por_caseta = por_caseta.sort_values(["REGISTROS", "PEAJE", "CASETA"], ascending=[False, True, True]).reset_index(drop=True)
    return {
        "por_peaje": por_peaje,
        "por_peaje_sentido": por_peaje_sentido,
        "por_caseta": por_caseta,
    }


def build_dashboard_metric_items(items: list[tuple[object, str]]) -> str:
    cards = []
    for value, label in items:
        cards.append(
            f'<div class="metric-tile"><div class="metric-value" style="color:#163564;">{value}</div>'
            f'<div class="metric-label" style="color:#4d6587;">{label}</div></div>'
        )
    return f'<div class="metrics-strip">{"".join(cards)}</div>'


def build_processing_flow_diagram_html() -> str:
    steps = [
        ("1. Recibir el archivo", "Tomamos la hoja elegida y verificamos que la base tenga la informacion minima para empezar."),
        ("2. Ordenar columnas", "Acomodamos nombres, formatos y posiciones para que todas las reglas trabajen sobre la misma estructura."),
        ("3. Marcar el orden real", "Guardamos el orden original de las filas para no perder la historia del archivo mientras se procesa."),
        ("4. Normalizar placas", "Quitamos ruido visual de las placas y armamos una version comparable para detectar errores o parecidos."),
        ("5. Aplicar reglas manuales", "Si existe una excepcion conocida, la usamos aqui para no mezclar casos de prueba con casos reales."),
        ("6. Separar placas observadas", "Las placas que no se pueden corregir con seguridad no se fuerzan: se dejan marcadas para revision."),
        ("7. Preparar tiempos", "Convertimos llegadas y salidas en tiempos comparables para revisar cola, caseta y salida sin ambiguedades."),
        ("8. Detectar bordes", "Identificamos registros en los extremos del periodo porque suelen tener tiempos incompletos o contexto parcial."),
        ("9. Recuperar faltantes", "Intentamos completar tiempos con interpolacion, medianas locales y referencias de filas parecidas."),
        ("10. Pedir ayuda a donantes", "Cuando una fila esta muy incompleta, usamos patrones validos del mismo peaje, sentido o caseta para proponer horas."),
        ("11. Corregir inversiones", "Si el orden de cola, caseta y salida queda cruzado por pocos segundos, lo reajustamos con reglas cortas y controladas."),
        ("12. Unir duplicados cercanos", "Si dos filas parecen ser el mismo vehiculo separado por muy poco tiempo, consolidamos la historia para no contar doble."),
        ("13. Distinguir fuga o fragmentacion", "Comparamos si el caso se parece mas a una fuga real o a un registro partido antes de etiquetarlo."),
        ("14. Medir fuerza del hallazgo", "No todos los casos valen lo mismo: algunos quedan como fuga fuerte, otros como probable y otros como no concluyentes."),
        ("15. Separar salidas", "Con todo revisado, armamos base limpia, eliminados, pendientes y tablas de soporte para auditoria."),
        ("16. Mostrar resultados", "Al final resumimos volumen, hallazgos, teoria de colas y descargas en un tablero que se pueda explicar facil."),
    ]
    cards = []
    for title, copy in steps:
        cards.append(
            '<div style="flex:1 1 220px; min-width:220px; border:1px solid rgba(15,61,145,0.12); '
            'background:linear-gradient(180deg, rgba(255,255,255,0.98), rgba(240,246,255,0.98)); '
            'border-radius:20px; padding:1rem 1rem 0.9rem; box-shadow:0 10px 24px rgba(12,41,90,0.08);">'
            f'<div style="font-size:1rem; font-weight:700; color:#163564; margin-bottom:0.45rem;">{title}</div>'
            f'<div style="font-size:0.92rem; line-height:1.6; color:#4d6587;">{copy}</div>'
            '</div>'
        )
    return (
        '<div style="border:1px solid rgba(15,61,145,0.12); border-radius:26px; '
        'background:linear-gradient(180deg, rgba(255,255,255,0.96), rgba(235,243,255,0.92)); '
        'padding:1.1rem 1.1rem 1.2rem; margin:0.75rem 0 1.15rem;">'
        '<div style="font-size:0.74rem; letter-spacing:0.12em; text-transform:uppercase; font-weight:700; color:#2f6ddc; margin-bottom:0.45rem;">Mapa Del Procedimiento</div>'
        '<div style="font-size:1.18rem; font-weight:700; color:#163564; margin-bottom:0.35rem;">Como trabaja este analisis, paso por paso</div>'
        '<div style="font-size:0.95rem; line-height:1.7; color:#4d6587; margin-bottom:0.9rem;">Piensalo como una revision ordenada: primero acomodamos la base, despues intentamos mejorarla y al final decidimos que casos son fuertes, cuales son solo alerta y cuales todavia necesitan cuidado.</div>'
        '<div style="display:flex; flex-wrap:wrap; gap:0.85rem; align-items:stretch;">'
        + "".join(cards)
        + '</div></div>'
    )


def build_fuga_rate_tables(
    df_original: pd.DataFrame,
    fugas_flujo: pd.DataFrame,
    fragmentaciones_probables: pd.DataFrame | None = None,
) -> dict[str, pd.DataFrame]:
    by_peaje_columns = [
        "PEAJE",
        "REGISTROS_DATA",
        "FUGAS_FLUJO",
        "FUGAS_FLUJO_%",
        "FRAGMENTACIONES_FLUJO",
        "FRAGMENTACIONES_FLUJO_%",
    ]
    by_caseta_columns = [
        "PEAJE",
        "CASETA",
        "SENTIDO",
        "REGISTROS_DATA",
        "FUGAS_FLUJO",
        "FUGAS_FLUJO_%",
        "FRAGMENTACIONES_FLUJO",
        "FRAGMENTACIONES_FLUJO_%",
    ]
    if df_original.empty:
        return {
            "by_peaje": pd.DataFrame(columns=by_peaje_columns),
            "by_caseta": pd.DataFrame(columns=by_caseta_columns),
        }

    base = df_original.copy()
    base["PEAJE"] = base["PEAJE"].map(format_dashboard_dimension)
    base["CASETA"] = base["CASETA"].map(format_dashboard_dimension)
    base["SENTIDO"] = base["SENTIDO"].map(format_dashboard_dimension)

    registros_peaje = (
        base.groupby(["PEAJE"], dropna=False)
        .size()
        .rename("REGISTROS_DATA")
        .reset_index()
    )
    registros_caseta = (
        base.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
        .size()
        .rename("REGISTROS_DATA")
        .reset_index()
    )

    fugas_operativas = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO"])
    if not fugas_flujo.empty:
        fugas_operativas = fugas_flujo[
            fugas_flujo["TIPO_FUGA"].astype(str).str.startswith(("fuga_fuerte_", "fuga_probable_"))
        ][["PEAJE", "CASETA", "SENTIDO"]].copy()

    if fugas_operativas.empty:
        fugas_peaje = pd.DataFrame(columns=["PEAJE", "FUGAS_FLUJO"])
        fugas_caseta = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FUGAS_FLUJO"])
    else:
        fugas_peaje = (
            fugas_operativas.groupby(["PEAJE"], dropna=False)
            .size()
            .rename("FUGAS_FLUJO")
            .reset_index()
        )
        fugas_caseta = (
            fugas_operativas.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
            .size()
            .rename("FUGAS_FLUJO")
            .reset_index()
        )

    fragmentaciones_operativas = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO"])
    if fragmentaciones_probables is not None and not fragmentaciones_probables.empty:
        fragmentaciones_operativas = fragmentaciones_probables[["PEAJE", "CASETA", "SENTIDO"]].copy()

    if fragmentaciones_operativas.empty:
        fragmentaciones_peaje = pd.DataFrame(columns=["PEAJE", "FRAGMENTACIONES_FLUJO"])
        fragmentaciones_caseta = pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FRAGMENTACIONES_FLUJO"])
    else:
        fragmentaciones_peaje = (
            fragmentaciones_operativas.groupby(["PEAJE"], dropna=False)
            .size()
            .rename("FRAGMENTACIONES_FLUJO")
            .reset_index()
        )
        fragmentaciones_caseta = (
            fragmentaciones_operativas.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
            .size()
            .rename("FRAGMENTACIONES_FLUJO")
            .reset_index()
        )

    by_peaje = registros_peaje.merge(fugas_peaje, on=["PEAJE"], how="left")
    by_caseta = registros_caseta.merge(fugas_caseta, on=["PEAJE", "CASETA", "SENTIDO"], how="left")
    by_peaje = by_peaje.merge(fragmentaciones_peaje, on=["PEAJE"], how="left")
    by_caseta = by_caseta.merge(fragmentaciones_caseta, on=["PEAJE", "CASETA", "SENTIDO"], how="left")
    for frame in [by_peaje, by_caseta]:
        frame["FUGAS_FLUJO"] = frame["FUGAS_FLUJO"].fillna(0).astype(int)
        frame["FUGAS_FLUJO_%"] = (100 * frame["FUGAS_FLUJO"] / frame["REGISTROS_DATA"]).round(2)
        frame["FRAGMENTACIONES_FLUJO"] = frame["FRAGMENTACIONES_FLUJO"].fillna(0).astype(int)
        frame["FRAGMENTACIONES_FLUJO_%"] = (100 * frame["FRAGMENTACIONES_FLUJO"] / frame["REGISTROS_DATA"]).round(2)

    by_peaje = by_peaje.sort_values(
        ["FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO_%", "FUGAS_FLUJO", "REGISTROS_DATA"],
        ascending=[False, False, False, False],
    ).reset_index(drop=True)
    by_caseta = by_caseta.sort_values(
        ["FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO_%", "FUGAS_FLUJO", "REGISTROS_DATA"],
        ascending=[False, False, False, False],
    ).reset_index(drop=True)
    return {
        "by_peaje": by_peaje[by_peaje_columns],
        "by_caseta": by_caseta[by_caseta_columns],
    }


def build_queue_theory_dashboard(df_resultados: pd.DataFrame, fuga_rate_tables: dict[str, pd.DataFrame] | None = None) -> dict[str, object]:
    by_peaje_columns = [
        "PEAJE",
        "VEHICULOS",
        "CASETAS",
        "SENTIDOS",
        "FUGAS_FLUJO",
        "FUGAS_FLUJO_%",
        "FRAGMENTACIONES_FLUJO",
        "FRAGMENTACIONES_FLUJO_%",
        "TEC_PROMEDIO_MIN",
        "TEC_P95_MIN",
        "TEC_MAX_MIN",
        "COLA_PROMEDIO_USU",
        "COLA_P95_USU",
        "COLA_MAX_USU",
        "ATENCION_<=3_MIN_%",
    ]
    by_caseta_columns = [
        "PEAJE",
        "CASETA",
        "SENTIDO",
        "VEHICULOS",
        "FUGAS_FLUJO",
        "FUGAS_FLUJO_%",
        "FRAGMENTACIONES_FLUJO",
        "FRAGMENTACIONES_FLUJO_%",
        "TEC_PROMEDIO_MIN",
        "TEC_P95_MIN",
        "TEC_MAX_MIN",
        "COLA_PROMEDIO_USU",
        "COLA_P95_USU",
        "COLA_MAX_USU",
        "ATENCION_<=3_MIN_%",
    ]
    if df_resultados.empty:
        return {
            "general_cards": [
                (0, "vehiculos evaluados"),
                (0, "peajes evaluados"),
                (0, "casetas evaluadas"),
                ("0.00", "TEC promedio global (min)"),
                ("0.00", "TEC p95 global (min)"),
                ("0.00", "cola promedio (usuarios)"),
                ("0.00", "cola p95 (usuarios)"),
                (0, "cola maxima real"),
                ("0.00%", "atencion <= 3 min"),
            ],
            "general_insights": pd.DataFrame(columns=["resultado", "valor"]),
            "top_peajes": pd.DataFrame(columns=by_peaje_columns),
            "by_peaje": pd.DataFrame(columns=by_peaje_columns),
            "by_caseta": pd.DataFrame(columns=by_caseta_columns),
        }

    df_queue = df_resultados.copy()
    df_queue["PEAJE"] = df_queue["PEAJE"].map(format_dashboard_dimension)
    df_queue["CASETA"] = df_queue["CASETA"].map(format_dashboard_dimension)
    df_queue["SENTIDO"] = df_queue["SENTIDO"].map(format_dashboard_dimension)
    fuga_rate_tables = fuga_rate_tables or {}
    fuga_by_peaje = fuga_rate_tables.get(
        "by_peaje",
        pd.DataFrame(columns=["PEAJE", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]),
    )
    fuga_by_caseta = fuga_rate_tables.get(
        "by_caseta",
        pd.DataFrame(columns=["PEAJE", "CASETA", "SENTIDO", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]),
    )

    overall_pct_3 = round(100 * (df_queue["T_TEC_FINAL_MINUTOS"] <= 3).mean(), 2) if len(df_queue) else 0.0
    general_cards = [
        (len(df_queue), "vehiculos evaluados"),
        (int(df_queue["PEAJE"].nunique(dropna=False)), "peajes evaluados"),
        (int(df_queue[["PEAJE", "CASETA", "SENTIDO"]].drop_duplicates().shape[0]), "casetas evaluadas"),
        (f"{df_queue['T_TEC_FINAL_MINUTOS'].mean():.2f}", "TEC promedio global (min)"),
        (f"{q95(df_queue['T_TEC_FINAL_MINUTOS']):.2f}", "TEC p95 global (min)"),
        (f"{df_queue['COLA_ESPERA_USUARIOS'].mean():.2f}", "cola promedio (usuarios)"),
        (f"{q95(df_queue['COLA_ESPERA_USUARIOS']):.2f}", "cola p95 (usuarios)"),
        (int(df_queue["COLA_ESPERA_USUARIOS"].max()), "cola maxima real"),
        (f"{overall_pct_3:.2f}%", "atencion <= 3 min"),
    ]

    by_peaje = (
        df_queue.groupby(["PEAJE"], dropna=False)
        .agg(
            VEHICULOS=("PLACA_FINAL", "size"),
            CASETAS=("CASETA", pd.Series.nunique),
            SENTIDOS=("SENTIDO", pd.Series.nunique),
            TEC_PROMEDIO_MIN=("T_TEC_FINAL_MINUTOS", "mean"),
            TEC_P95_MIN=("T_TEC_FINAL_MINUTOS", q95),
            TEC_MAX_MIN=("T_TEC_FINAL_MINUTOS", "max"),
            COLA_PROMEDIO_USU=("COLA_ESPERA_USUARIOS", "mean"),
            COLA_P95_USU=("COLA_ESPERA_USUARIOS", q95),
            COLA_MAX_USU=("COLA_ESPERA_USUARIOS", "max"),
            DENTRO_3_MIN=("T_TEC_FINAL_MINUTOS", lambda s: int((s <= 3).sum())),
        )
        .reset_index()
    )
    by_peaje["ATENCION_<=3_MIN_%"] = (100 * by_peaje["DENTRO_3_MIN"] / by_peaje["VEHICULOS"]).round(2)
    by_peaje = by_peaje.drop(columns=["DENTRO_3_MIN"])
    by_peaje = by_peaje.merge(
        fuga_by_peaje[["PEAJE", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]],
        on=["PEAJE"],
        how="left",
    )
    by_peaje[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]] = by_peaje[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]].fillna(0)
    by_peaje["FUGAS_FLUJO"] = by_peaje["FUGAS_FLUJO"].astype(int)
    by_peaje["FRAGMENTACIONES_FLUJO"] = by_peaje["FRAGMENTACIONES_FLUJO"].astype(int)
    by_peaje[[
        "TEC_PROMEDIO_MIN",
        "TEC_P95_MIN",
        "TEC_MAX_MIN",
        "COLA_PROMEDIO_USU",
        "COLA_P95_USU",
    ]] = by_peaje[[
        "TEC_PROMEDIO_MIN",
        "TEC_P95_MIN",
        "TEC_MAX_MIN",
        "COLA_PROMEDIO_USU",
        "COLA_P95_USU",
    ]].round(2)
    by_peaje = by_peaje.sort_values(["TEC_PROMEDIO_MIN", "COLA_MAX_USU", "VEHICULOS"], ascending=[False, False, False]).reset_index(drop=True)

    by_caseta = (
        df_queue.groupby(["PEAJE", "CASETA", "SENTIDO"], dropna=False)
        .agg(
            VEHICULOS=("PLACA_FINAL", "size"),
            TEC_PROMEDIO_MIN=("T_TEC_FINAL_MINUTOS", "mean"),
            TEC_P95_MIN=("T_TEC_FINAL_MINUTOS", q95),
            TEC_MAX_MIN=("T_TEC_FINAL_MINUTOS", "max"),
            COLA_PROMEDIO_USU=("COLA_ESPERA_USUARIOS", "mean"),
            COLA_P95_USU=("COLA_ESPERA_USUARIOS", q95),
            COLA_MAX_USU=("COLA_ESPERA_USUARIOS", "max"),
            DENTRO_3_MIN=("T_TEC_FINAL_MINUTOS", lambda s: int((s <= 3).sum())),
        )
        .reset_index()
    )
    by_caseta["ATENCION_<=3_MIN_%"] = (100 * by_caseta["DENTRO_3_MIN"] / by_caseta["VEHICULOS"]).round(2)
    by_caseta = by_caseta.drop(columns=["DENTRO_3_MIN"])
    by_caseta = by_caseta.merge(
        fuga_by_caseta[["PEAJE", "CASETA", "SENTIDO", "FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]],
        on=["PEAJE", "CASETA", "SENTIDO"],
        how="left",
    )
    by_caseta[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]] = by_caseta[["FUGAS_FLUJO", "FUGAS_FLUJO_%", "FRAGMENTACIONES_FLUJO", "FRAGMENTACIONES_FLUJO_%"]].fillna(0)
    by_caseta["FUGAS_FLUJO"] = by_caseta["FUGAS_FLUJO"].astype(int)
    by_caseta["FRAGMENTACIONES_FLUJO"] = by_caseta["FRAGMENTACIONES_FLUJO"].astype(int)
    by_caseta[[
        "TEC_PROMEDIO_MIN",
        "TEC_P95_MIN",
        "TEC_MAX_MIN",
        "COLA_PROMEDIO_USU",
        "COLA_P95_USU",
    ]] = by_caseta[[
        "TEC_PROMEDIO_MIN",
        "TEC_P95_MIN",
        "TEC_MAX_MIN",
        "COLA_PROMEDIO_USU",
        "COLA_P95_USU",
    ]].round(2)
    by_caseta = by_caseta.sort_values(["TEC_PROMEDIO_MIN", "COLA_MAX_USU", "VEHICULOS"], ascending=[False, False, False]).reset_index(drop=True)

    top_peaje = by_peaje.iloc[0]
    top_caseta = by_caseta.iloc[0]
    worst_compliance_peaje = by_peaje.sort_values(["ATENCION_<=3_MIN_%", "TEC_PROMEDIO_MIN"], ascending=[True, False]).iloc[0]
    highest_queue_caseta = by_caseta.sort_values(["COLA_MAX_USU", "TEC_PROMEDIO_MIN"], ascending=[False, False]).iloc[0]
    general_insights = pd.DataFrame(
        [
            {
                "resultado": "Mayor TEC promedio por peaje",
                "valor": f"{top_peaje['PEAJE']} con {top_peaje['TEC_PROMEDIO_MIN']:.2f} min y cola maxima de {int(top_peaje['COLA_MAX_USU'])} usuarios.",
            },
            {
                "resultado": "Caseta mas exigida",
                "valor": f"{top_caseta['PEAJE']} | C{top_caseta['CASETA']} | {top_caseta['SENTIDO']} con TEC promedio de {top_caseta['TEC_PROMEDIO_MIN']:.2f} min.",
            },
            {
                "resultado": "Peaje con menor cumplimiento <= 3 min",
                "valor": f"{worst_compliance_peaje['PEAJE']} con {worst_compliance_peaje['ATENCION_<=3_MIN_%']:.2f}% dentro del umbral.",
            },
            {
                "resultado": "Mayor cola real observada",
                "valor": f"{highest_queue_caseta['PEAJE']} | C{highest_queue_caseta['CASETA']} | {highest_queue_caseta['SENTIDO']} con {int(highest_queue_caseta['COLA_MAX_USU'])} usuarios.",
            },
        ]
    )

    return {
        "general_cards": general_cards,
        "general_insights": general_insights,
        "top_peajes": by_peaje.head(8)[by_peaje_columns],
        "by_peaje": by_peaje[by_peaje_columns],
        "by_caseta": by_caseta[by_caseta_columns],
    }


def plot_volume_by_peaje(df_chart: pd.DataFrame, title: str, color: str) -> plt.Figure | None:
    if df_chart.empty:
        return None
    plot_df = df_chart.sort_values("REGISTROS", ascending=True)
    fig, ax = plt.subplots(figsize=(7.4, max(3.2, len(plot_df) * 0.68)), dpi=160)
    ax.barh(plot_df["PEAJE"], plot_df["REGISTROS"], color=color, edgecolor="#123256", linewidth=0.7)
    ax.set_title(title, fontsize=12, fontweight="bold", loc="left")
    ax.grid(axis="x", color="#dbe5f3", linewidth=0.8)
    ax.set_axisbelow(True)
    ax.set_xlabel("Registros")
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#bfd0e6")
    for y_pos, value in enumerate(plot_df["REGISTROS"]):
        ax.text(value, y_pos, f" {int(value)}", va="center", ha="left", fontsize=9, color="#163564")
    fig.patch.set_facecolor("white")
    return fig


def plot_volume_by_sentido(df_chart: pd.DataFrame, title: str) -> plt.Figure | None:
    if df_chart.empty:
        return None
    pivot = df_chart.pivot_table(index="PEAJE", columns="SENTIDO", values="REGISTROS", aggfunc="sum", fill_value=0)
    if pivot.empty:
        return None
    pivot = pivot.loc[pivot.sum(axis=1).sort_values().index]
    colors = ["#0f3d91", "#2f6ddc", "#7aa2e3", "#b7ccee", "#dce8f8"]
    fig, ax = plt.subplots(figsize=(7.8, max(3.4, len(pivot) * 0.72)), dpi=160)
    left = pd.Series(0, index=pivot.index, dtype=float)
    for idx, sentido in enumerate(pivot.columns):
        values = pivot[sentido]
        ax.barh(
            pivot.index,
            values,
            left=left,
            label=str(sentido),
            color=colors[idx % len(colors)],
            edgecolor="white",
            linewidth=0.6,
        )
        total_max = float(pivot.sum(axis=1).max()) if len(pivot.index) else 0.0
        for y_pos, value in enumerate(values):
            if value <= 0:
                continue
            x_pos = float(left.iloc[y_pos] + value / 2)
            text_color = "white" if value >= max(40, 0.12 * total_max) else "#163564"
            ax.text(x_pos, y_pos, f"{int(value)}", ha="center", va="center", fontsize=8.5, color=text_color, fontweight="bold")
        left = left + values
    totals = pivot.sum(axis=1)
    for y_pos, total in enumerate(totals):
        ax.text(float(total), y_pos, f" {int(total)}", va="center", ha="left", fontsize=8.8, color="#163564")
    ax.set_title(title, fontsize=12, fontweight="bold", loc="left")
    ax.set_xlabel("Registros")
    ax.grid(axis="x", color="#dbe5f3", linewidth=0.8)
    ax.set_axisbelow(True)
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#bfd0e6")
    ax.legend(frameon=False, loc="lower right")
    fig.patch.set_facecolor("white")
    return fig


def plot_top_labels(df_chart: pd.DataFrame, title: str, color: str) -> plt.Figure | None:
    if df_chart.empty:
        return None
    plot_df = df_chart.head(10).iloc[::-1].copy()
    fig, ax = plt.subplots(figsize=(8.6, max(4.0, len(plot_df) * 0.6)), dpi=160)
    ax.barh(plot_df["ETIQUETA"], plot_df["REGISTROS"], color=color, edgecolor="#183a66", linewidth=0.7)
    ax.set_title(title, fontsize=12, fontweight="bold", loc="left")
    ax.set_xlabel("Registros")
    ax.grid(axis="x", color="#dbe5f3", linewidth=0.8)
    ax.set_axisbelow(True)
    for spine in ["top", "right", "left"]:
        ax.spines[spine].set_visible(False)
    ax.spines["bottom"].set_color("#bfd0e6")
    for y_pos, value in enumerate(plot_df["REGISTROS"]):
        ax.text(value, y_pos, f" {int(value)}", va="center", ha="left", fontsize=8.6, color="#163564")
    fig.patch.set_facecolor("white")
    return fig


def render_dashboard_figure(fig: plt.Figure | None) -> None:
    if fig is None:
        st.info("No hay datos suficientes para graficar este bloque.")
        return
    st.pyplot(fig, use_container_width=True)
    plt.close(fig)


def build_processing_dashboard(df_std: pd.DataFrame, result: dict[str, object]) -> dict[str, object]:
    export_tables = result["export_tables"]
    plate_df = result["plate_result"]["df"].copy()
    time_final = result["time_result"]["df_tiempos_final"].copy()
    df_resultados_queue = result["informe_package"]["df_resultados"].copy()
    raw_df = prepare_dashboard_dataframe(df_std)
    clean_df = prepare_dashboard_dataframe(export_tables["base_limpia"])
    fugas_df = detect_raw_fugas(df_std)
    flow_fugas = detect_flow_fuga_candidates({"df_tiempos_bordes": result["time_result"]["df_tiempos_bordes"]})
    fugas_probables_df = flow_fugas["fugas_probables"]
    fragmentaciones_df = flow_fugas["fragmentaciones_probables"]
    fuga_rate_tables = build_fuga_rate_tables(df_std, fugas_probables_df, fragmentaciones_df)
    fragmentaciones_confianza = (
        fragmentaciones_df["CONFIANZA"].value_counts().rename_axis("CONFIANZA").reset_index(name="FILAS")
        if not fragmentaciones_df.empty
        else pd.DataFrame(columns=["CONFIANZA", "FILAS"])
    )

    plate_actions = plate_df["PLACA_ACCION_FINAL"].astype(str).value_counts(dropna=False).rename_axis("ACCION").reset_index(name="FILAS")
    plate_actions["ACCION_UI"] = plate_actions["ACCION"].map(format_dashboard_action)

    time_actions = time_final["TIEMPO_ACCION_CIERRE"].astype(str).value_counts(dropna=False).rename_axis("ACCION").reset_index(name="FILAS")
    time_actions["ACCION_UI"] = time_actions["ACCION"].map(format_dashboard_action)

    plate_df["FECHA_DIA_DASH"] = pd.to_datetime(plate_df["FECHA"], errors="coerce").dt.strftime("%Y-%m-%d").fillna("Sin fecha")
    plate_df["PEAJE"] = plate_df["PEAJE"].map(format_dashboard_dimension)
    plate_df["CASETA"] = plate_df["CASETA"].map(format_dashboard_dimension)
    plate_df["SENTIDO"] = plate_df["SENTIDO"].map(format_dashboard_dimension)
    plate_df["PLACA_FINAL_DECIDIDA"] = plate_df["PLACA_FINAL_DECIDIDA"].map(format_dashboard_dimension)
    caseta_changes = (
        plate_df.groupby(["PEAJE", "SENTIDO", "FECHA_DIA_DASH", "PLACA_FINAL_DECIDIDA"], dropna=False)
        .agg(
            REGISTROS=("PLACA", "size"),
            CASETAS_UNICAS=("CASETA", "nunique"),
            CASETAS=("CASETA", lambda s: ", ".join(sorted(pd.unique(s.astype(str))))),
            ACCIONES=("PLACA_ACCION_FINAL", lambda s: ", ".join(sorted(pd.unique(s.astype(str))))),
        )
        .reset_index()
    )
    caseta_changes = caseta_changes[caseta_changes["CASETAS_UNICAS"] > 1].copy()
    caseta_changes = caseta_changes.rename(columns={"FECHA_DIA_DASH": "FECHA", "PLACA_FINAL_DECIDIDA": "PLACA_FINAL"})
    caseta_changes = caseta_changes.sort_values(["CASETAS_UNICAS", "REGISTROS"], ascending=[False, False]).reset_index(drop=True)

    corrected_plate_rows = int(plate_df["PLACA_ACCION_FINAL"].astype(str).str.startswith("corregir_").sum())
    recovered_time_rows = int(time_final["TIEMPO_ACCION_CIERRE"].astype(str).str.startswith("imputar_").sum())
    final_time_adjustments = int(time_final["TIEMPO_ACCION_CIERRE"].astype(str).str.contains("swap_|consolidar_post_tiempo", regex=True, na=False).sum())

    overview = {
        "raw_rows": len(raw_df),
        "clean_rows": len(clean_df),
        "deleted_rows": len(export_tables["casos_eliminados"]),
        "pending_rows": len(export_tables["casos_pendientes"]),
        "raw_peajes": int(raw_df["PEAJE"].nunique(dropna=False)),
        "raw_casetas": int(raw_df["CASETA"].nunique(dropna=False)),
        "raw_sentidos": int(raw_df["SENTIDO"].nunique(dropna=False)),
        "clean_peajes": int(clean_df["PEAJE"].nunique(dropna=False)),
        "clean_casetas": int(clean_df["CASETA"].nunique(dropna=False)),
        "clean_sentidos": int(clean_df["SENTIDO"].nunique(dropna=False)),
        "fugas_rows": len(fugas_df),
        "fugas_unique": int(fugas_df["PLACA_NORMALIZADA"].nunique()) if not fugas_df.empty else 0,
        "fugas_fuertes": int(fugas_probables_df["TIPO_FUGA"].astype(str).str.startswith("fuga_fuerte_").sum()) if not fugas_probables_df.empty else 0,
        "fugas_probables": int(fugas_probables_df["TIPO_FUGA"].astype(str).str.startswith("fuga_probable_").sum()) if not fugas_probables_df.empty else 0,
        "fugas_no_concluyentes": int(fugas_probables_df["TIPO_FUGA"].astype(str).str.startswith("incompleto_no_concluyente").sum()) if not fugas_probables_df.empty else 0,
        "fragmentaciones_probables": len(fragmentaciones_df),
        "fragmentaciones_alta_confianza": int(fragmentaciones_df["CONFIANZA"].eq("alta").sum()) if not fragmentaciones_df.empty else 0,
        "flagged_plates": len(export_tables["revision_placas"]),
        "corrected_plate_rows": corrected_plate_rows,
        "caseta_change_groups": len(caseta_changes),
        "recovered_time_rows": recovered_time_rows,
        "final_time_adjustments": final_time_adjustments,
        "retention_pct": round((len(clean_df) / len(raw_df) * 100), 2) if len(raw_df) else 0.0,
    }
    return {
        "overview": overview,
        "raw_tables": build_volume_tables(raw_df),
        "clean_tables": build_volume_tables(clean_df),
        "queue_theory": build_queue_theory_dashboard(df_resultados_queue, fuga_rate_tables),
        "fugas_patron_detalle": fugas_df,
        "fugas_probables_detalle": fugas_probables_df,
        "fragmentaciones_detalle": fragmentaciones_df,
        "fragmentaciones_confianza": fragmentaciones_confianza,
        "plate_actions": plate_actions,
        "time_actions": time_actions,
        "caseta_changes": caseta_changes,
    }


def build_fugas_report_sheets(dashboard: dict[str, object]) -> dict[str, pd.DataFrame]:
    overview = dashboard["overview"]
    resumen = pd.DataFrame(
        [
            {"indicador": "filas_patron_x_longitud", "valor": overview["fugas_rows"]},
            {"indicador": "placas_patron_x_unicas", "valor": overview["fugas_unique"]},
            {"indicador": "fugas_fuertes_flujo", "valor": overview["fugas_fuertes"]},
            {"indicador": "fugas_probables_flujo", "valor": overview["fugas_probables"]},
            {"indicador": "incompletos_no_concluyentes", "valor": overview["fugas_no_concluyentes"]},
            {"indicador": "fragmentaciones_probables", "valor": overview["fragmentaciones_probables"]},
            {"indicador": "fragmentaciones_alta_confianza", "valor": overview["fragmentaciones_alta_confianza"]},
            {"indicador": "placas_con_cambio_caseta", "valor": overview["caseta_change_groups"]},
        ]
    )
    return {
        "resumen_fugas": resumen,
        "fugas_patron_placa": dashboard["fugas_patron_detalle"],
        "fugas_probables_flujo": dashboard["fugas_probables_detalle"],
        "fragmentaciones_probables": dashboard["fragmentaciones_detalle"],
        "fragmentaciones_confianza": dashboard["fragmentaciones_confianza"],
        "placas_cambio_caseta": dashboard["caseta_changes"],
    }


def process_pipeline(df_std: pd.DataFrame, config: dict, manual_rules_df: pd.DataFrame) -> dict[str, object]:
    config = {**DEFAULT_CONFIG, **config}
    df_std = df_std.copy()
    df_std["_ORDEN_FILA"] = range(len(df_std))
    if config["aplicar_limpieza_placa"]:
        plate_result = run_plate_cleaning(df_std, config, manual_rules_df)
    else:
        df_passthrough = limpiar_placas_peaje(df_std)
        df_passthrough["_ORDEN_FILA"] = df_std["_ORDEN_FILA"]
        df_passthrough["PLACA_ACCION_FINAL"] = "sin_cambio"
        df_passthrough["PLACA_FINAL_DECIDIDA"] = df_passthrough["PLACA_NORMALIZADA"]
        df_passthrough["PLACA_EXCLUIR_ANALISIS"] = False
        df_passthrough["PLACA_AJUSTE_MANUAL"] = pd.NA
        plate_result = {
            "df": df_passthrough,
            "df_trabajo": df_passthrough.copy(),
            "df_eliminados": pd.DataFrame(columns=df_passthrough.columns),
            "df_revision_placas": pd.DataFrame(),
            "df_bloques_decision": pd.DataFrame(),
            "resumen_acciones_placa": df_passthrough["PLACA_ACCION_FINAL"].value_counts(dropna=False).rename_axis("PLACA_ACCION_FINAL").to_frame("filas"),
        }

    time_result = run_time_cleaning(plate_result["df_trabajo"], config)
    export_tables = build_export_tables(df_std, plate_result, time_result, config, manual_rules_df)
    exact_export = build_exact_export_package(df_std, export_tables)
    informe_package = build_informe_package(export_tables["export_base_detalle"])
    complementary_package = build_complementary_package(
        informe_package["df_resultados"],
        export_tables["export_base_detalle"],
        df_std,
        export_tables["fugas_flujo"],
        export_tables["fragmentaciones_probables"],
    )
    return {
        "input_df": df_std,
        "plate_result": plate_result,
        "time_result": time_result,
        "export_tables": export_tables,
        "exact_export": exact_export,
        "informe_package": informe_package,
        "complementary_package": complementary_package,
    }


def load_streamlit_secrets() -> dict:
    try:
        return dict(st.secrets)
    except StreamlitSecretNotFoundError:
        return {}


def explain_auth_failure(reason: str) -> str:
    messages = {
        "invalid_credentials": "Usuario o contrasena no validos.",
        "disabled": "Tu usuario esta deshabilitado.",
        "not_yet_active": "Tu usuario todavia no entra en vigencia.",
        "expired": "Tu usuario ya vencio.",
        "auth_not_available": "El backend actual no tiene autenticacion habilitada.",
    }
    return messages.get(reason, "No fue posible iniciar sesion.")


def normalize_date_input(value) -> datetime | None:
    if value is None or pd.isna(value):
        return None
    if isinstance(value, pd.Timestamp):
        return value.to_pydatetime()
    if isinstance(value, datetime):
        return value
    return datetime.combine(value, time.min)


def date_to_window_start(value) -> datetime | None:
    parsed = normalize_date_input(value)
    if not parsed:
        return None
    return parsed.replace(hour=0, minute=0, second=0, microsecond=0)


def date_to_window_end(value) -> datetime | None:
    parsed = normalize_date_input(value)
    if not parsed:
        return None
    return parsed.replace(hour=23, minute=59, second=59, microsecond=0)


def date_value_or_none(value):
    parsed = normalize_date_input(value)
    return parsed.date() if parsed else None


def password_is_valid(password: str) -> bool:
    return len(password) >= 8


def build_available_users_df(users_df: pd.DataFrame) -> pd.DataFrame:
    available_users_df = users_df.copy()
    if "is_enabled" in available_users_df.columns:
        available_users_df["estado"] = available_users_df["is_enabled"].map(
            {True: "Habilitado", False: "Deshabilitado"}
        ).fillna("Sin dato")
    for column_name in ["active_from", "active_until", "last_login_at", "created_at"]:
        if column_name in available_users_df.columns:
            available_users_df[column_name] = pd.to_datetime(
                available_users_df[column_name], errors="coerce"
            ).dt.strftime("%Y-%m-%d %H:%M")
            available_users_df[column_name] = available_users_df[column_name].fillna("-")

    visible_columns = [
        column_name
        for column_name in [
            "username",
            "full_name",
            "role_name",
            "estado",
            "email",
            "phone_number",
            "active_from",
            "active_until",
            "last_login_at",
        ]
        if column_name in available_users_df.columns
    ]
    return available_users_df[visible_columns]


def render_bootstrap_admin(storage_backend) -> None:
    _, center_col, _ = st.columns([1, 0.9, 1], gap="large")
    with center_col:
        st.markdown(
            '<section class="auth-card"><div class="auth-card-kicker">Primer acceso</div><div class="auth-card-title">Crear administrador inicial</div><p class="auth-card-copy">Registra la primera cuenta con control total sobre usuarios, accesos e historial.</p></section>',
            unsafe_allow_html=True,
        )
        with st.form("bootstrap_admin_form"):
            full_name = st.text_input("Nombre completo")
            username = st.text_input("Usuario")
            email = st.text_input("Correo electronico")
            phone_number = st.text_input("Celular")
            password = st.text_input("Contrasena", type="password")
            confirm_password = st.text_input("Confirmar contrasena", type="password")
            submitted = st.form_submit_button("Crear administrador", use_container_width=True)

    if not submitted:
        return

    if not full_name.strip() or not username.strip() or not email.strip() or not phone_number.strip():
        st.error("Nombre completo, usuario, correo y celular son obligatorios.")
        return
    if password != confirm_password:
        st.error("Las contrasenas no coinciden.")
        return
    if not password_is_valid(password):
        st.error("La contrasena debe tener al menos 8 caracteres.")
        return

    try:
        storage_backend.create_initial_admin(username, full_name, password, email.strip(), phone_number.strip())
        auth_result = storage_backend.authenticate_user(username, password)
    except Exception as exc:
        st.error(f"No pude crear el administrador inicial: {exc}")
        return

    if auth_result.get("ok"):
        set_authenticated_user(auth_result["user"])
        st.session_state[APP_NAV_KEY] = "Inicio"
        st.query_params.clear()
        st.rerun()


def render_login_gate(storage_backend):
    if not storage_backend.has_users():
        render_bootstrap_admin(storage_backend)
        return None

    users_df = storage_backend.list_users().copy()
    if "id" in users_df.columns:
        users_df["id"] = pd.to_numeric(users_df["id"], errors="coerce")
        users_df = users_df[users_df["id"].notna()].copy()
    user_options = users_df.to_dict("records")

    _, center_col, _ = st.columns([1, 0.85, 1], gap="large")
    with center_col:
        st.markdown(
            '<section class="auth-card"><div class="auth-card-kicker">Ingreso</div><div class="auth-card-title">Iniciar sesion</div><p class="auth-card-copy">Ingresa tus credenciales para acceder al aplicativo.</p></section>',
            unsafe_allow_html=True,
        )
        with st.form("login_form"):
            selected_user = st.selectbox(
                "Usuario",
                options=user_options,
                format_func=lambda row: f"{row['username']} | {row['full_name']} | {row['role_name']}",
            )
            password = st.text_input("Contrasena", type="password")
            submitted = st.form_submit_button("Ingresar", use_container_width=True)

    if not submitted:
        return None

    auth_result = storage_backend.authenticate_user(selected_user["username"], password)
    if not auth_result.get("ok"):
        st.error(explain_auth_failure(auth_result.get("reason", "")))
        return None

    set_authenticated_user(auth_result["user"])
    st.session_state[APP_NAV_KEY] = "Inicio"
    st.query_params.clear()
    st.rerun()
    return None


def render_history_page(storage_backend) -> None:
    render_back_to_home_button("Historial")
    st.markdown(
        build_hero_panel(
            title="Historial de corridas",
            copy="Consulta ejecuciones recientes guardadas en la base de datos para seguimiento operativo y trazabilidad.",
            kicker="Historial",
            metrics=[
                (storage_backend.mode.upper(), "backend"),
                ("50", "maximo visible"),
            ],
        ),
        unsafe_allow_html=True,
    )
    st.dataframe(storage_backend.list_recent_runs(50), use_container_width=True)


def render_user_management_page(storage_backend, current_user: dict) -> None:
    can_manage_users = user_has_permission("manage_users", current_user)
    can_manage_activation = user_has_permission("manage_user_activation", current_user)

    render_back_to_home_button("Usuarios")
    st.markdown(
        build_hero_panel(
            title="Gestion de usuarios",
            copy="Administra altas, roles, vigencias y credenciales de acceso con una vista centralizada para operacion y control.",
            kicker="Administracion",
            metrics=[
                ("3", "roles base"),
                ("Correo + celular", "datos obligatorios"),
                (current_user["role_label"], "tu perfil"),
            ],
        ),
        unsafe_allow_html=True,
    )
    st.caption("Roles disponibles: administrador general, administrador operativo y analista.")

    if not (can_manage_users or can_manage_activation):
        st.error("No tienes permisos para gestionar usuarios.")
        return

    roles_info = pd.DataFrame(
        [
            {
                "rol": role_key,
                "nombre": role_data["label"],
                "descripcion": role_data["description"],
                "permisos": ", ".join(role_data["permissions"]),
            }
            for role_key, role_data in ROLE_DEFINITIONS.items()
        ]
    )
    st.dataframe(roles_info, use_container_width=True, hide_index=True)

    users_df = storage_backend.list_users().copy()
    if "id" in users_df.columns:
        users_df["id"] = pd.to_numeric(users_df["id"], errors="coerce")
        users_df = users_df[users_df["id"].notna()].copy()

    st.subheader("Usuarios actuales")
    st.caption(f"Usuarios disponibles en la base de datos: {len(users_df)}")
    st.dataframe(build_available_users_df(users_df), use_container_width=True, hide_index=True)

    if users_df.empty:
        st.info("Aun no hay usuarios registrados.")
        return

    user_options = {
        f"{row['username']} | {row['role_name']}": row.to_dict()
        for _, row in users_df.iterrows()
    }
    selected_user_label = st.selectbox("Usuario a editar", options=list(user_options.keys()))
    selected_user = user_options[selected_user_label]
    current_user_id = pd.to_numeric(pd.Series([current_user.get("id")]), errors="coerce").iloc[0]
    selected_user_id = int(selected_user["id"])
    selected_is_self = pd.notna(current_user_id) and selected_user_id == int(current_user_id)

    current_active_from = date_value_or_none(selected_user.get("active_from"))
    current_active_until = date_value_or_none(selected_user.get("active_until"))
    can_edit_access = can_manage_users or can_manage_activation

    st.subheader("Actualizar usuario")
    with st.form("edit_user_form"):
        full_name = st.text_input(
            "Nombre completo",
            value=str(selected_user.get("full_name") or ""),
            disabled=not can_manage_users,
        )
        email = st.text_input(
            "Correo electronico",
            value="" if pd.isna(selected_user.get("email")) else str(selected_user.get("email")),
            disabled=not can_manage_users,
        )
        phone_number = st.text_input(
            "Celular",
            value="" if pd.isna(selected_user.get("phone_number")) else str(selected_user.get("phone_number")),
            disabled=not can_manage_users,
        )
        role_keys = list(ROLE_DEFINITIONS.keys())
        role_index = role_keys.index(selected_user["role_key"]) if selected_user["role_key"] in role_keys else 0
        role_key = st.selectbox(
            "Rol",
            options=role_keys,
            index=role_index,
            format_func=lambda key: ROLE_DEFINITIONS[key]["label"],
            disabled=not can_manage_users,
        )
        is_enabled = st.checkbox(
            "Usuario habilitado",
            value=bool(selected_user.get("is_enabled")),
            disabled=not can_edit_access,
        )
        use_active_from = st.checkbox(
            "Restringir fecha inicial",
            value=current_active_from is not None,
            key=f"edit_use_active_from_{selected_user['id']}",
            disabled=not can_edit_access,
        )
        active_from = st.date_input(
            "Activo desde",
            value=current_active_from or datetime.now().date(),
            key=f"edit_active_from_{selected_user['id']}",
            disabled=(not can_edit_access) or (not use_active_from),
        )
        use_active_until = st.checkbox(
            "Restringir fecha final",
            value=current_active_until is not None,
            key=f"edit_use_active_until_{selected_user['id']}",
            disabled=not can_edit_access,
        )
        active_until = st.date_input(
            "Activo hasta",
            value=current_active_until or datetime.now().date(),
            key=f"edit_active_until_{selected_user['id']}",
            disabled=(not can_edit_access) or (not use_active_until),
        )
        new_password = st.text_input(
            "Nueva contrasena",
            type="password",
            disabled=not can_manage_users,
        )
        confirm_password = st.text_input(
            "Confirmar nueva contrasena",
            type="password",
            disabled=not can_manage_users,
        )
        submitted = st.form_submit_button("Guardar cambios", use_container_width=True)

    if submitted:
        if selected_is_self and can_edit_access and not is_enabled:
            st.error("No puedes deshabilitar tu propio usuario desde esta sesion.")
            return
        if selected_is_self and can_manage_users and role_key != selected_user["role_key"]:
            st.error("No puedes cambiar tu propio rol mientras estas autenticado.")
            return
        if can_manage_users and new_password:
            if new_password != confirm_password:
                st.error("Las contrasenas no coinciden.")
                return
            if not password_is_valid(new_password):
                st.error("La nueva contrasena debe tener al menos 8 caracteres.")
                return
        if can_manage_users and (not email.strip() or not phone_number.strip()):
            st.error("Correo electronico y celular son obligatorios.")
            return

        payload = {}
        if can_manage_users:
            payload.update(
                {
                    "full_name": full_name.strip(),
                    "email": email.strip(),
                    "phone_number": phone_number.strip(),
                    "role_key": role_key,
                }
            )
            if new_password:
                payload["password"] = new_password
        if can_edit_access:
            payload.update(
                {
                    "is_enabled": is_enabled,
                    "active_from": date_to_window_start(active_from) if use_active_from else None,
                    "active_until": date_to_window_end(active_until) if use_active_until else None,
                }
            )

        if payload.get("active_from") and payload.get("active_until") and payload["active_from"] > payload["active_until"]:
            st.error("La fecha inicial no puede ser mayor que la fecha final.")
            return

        try:
            storage_backend.update_user(selected_user_id, payload)
        except Exception as exc:
            st.error(f"No pude actualizar el usuario: {exc}")
            return

        if selected_is_self and can_manage_users:
            current_user["full_name"] = full_name.strip()
            current_user["email"] = email.strip()
            current_user["phone_number"] = phone_number.strip()
            set_authenticated_user(current_user)
        st.success("Usuario actualizado.")
        st.rerun()

    if not can_manage_users:
        st.info("Tu rol solo puede habilitar, deshabilitar y definir vigencia de usuarios.")
        return

    st.subheader("Crear usuario")
    with st.form("create_user_form"):
        username = st.text_input("Nuevo usuario")
        full_name = st.text_input("Nombre completo", key="create_full_name")
        email = st.text_input("Correo electronico", key="create_email")
        phone_number = st.text_input("Celular", key="create_phone_number")
        password = st.text_input("Contrasena inicial", type="password", key="create_password")
        confirm_password = st.text_input("Confirmar contrasena", type="password", key="create_confirm_password")
        role_key = st.selectbox(
            "Rol inicial",
            options=list(ROLE_DEFINITIONS.keys()),
            format_func=lambda key: ROLE_DEFINITIONS[key]["label"],
            key="create_role_key",
        )
        is_enabled = st.checkbox("Crear usuario habilitado", value=True, key="create_enabled")
        use_active_from = st.checkbox("Definir fecha inicial", value=False, key="create_use_active_from")
        active_from = st.date_input(
            "Activo desde",
            value=datetime.now().date(),
            key="create_active_from",
            disabled=not use_active_from,
        )
        use_active_until = st.checkbox("Definir fecha final", value=False, key="create_use_active_until")
        active_until = st.date_input(
            "Activo hasta",
            value=datetime.now().date(),
            key="create_active_until",
            disabled=not use_active_until,
        )
        submitted = st.form_submit_button("Crear usuario", use_container_width=True)

    if not submitted:
        return

    if password != confirm_password:
        st.error("Las contrasenas no coinciden.")
        return
    if not password_is_valid(password):
        st.error("La contrasena inicial debe tener al menos 8 caracteres.")
        return
    if not email.strip() or not phone_number.strip():
        st.error("Correo electronico y celular son obligatorios.")
        return

    payload = {
        "username": username,
        "full_name": full_name,
        "email": email.strip(),
        "phone_number": phone_number.strip(),
        "password": password,
        "role_key": role_key,
        "is_enabled": is_enabled,
        "active_from": date_to_window_start(active_from) if use_active_from else None,
        "active_until": date_to_window_end(active_until) if use_active_until else None,
        "created_by": current_user["id"],
    }
    if payload["active_from"] and payload["active_until"] and payload["active_from"] > payload["active_until"]:
        st.error("La fecha inicial no puede ser mayor que la fecha final.")
        return

    try:
        storage_backend.create_user(payload)
    except Exception as exc:
        st.error(f"No pude crear el usuario: {exc}")
        return

    st.success("Usuario creado.")
    st.rerun()


def render_processing_page(storage_backend, current_user: dict | None) -> None:
    render_back_to_home_button("TEC")
    st.markdown(
        build_hero_panel(
            title="Modulo TEC",
            copy="Carga bases, revisa columnas, controla la configuracion del pipeline y revisa un tablero ejecutivo del antes y despues del procesamiento dentro del aplicativo.",
            kicker="Procesamiento activo",
            metrics=[
                ("TEC", "modulo"),
                (storage_backend.mode.upper(), "persistencia"),
                (current_user["role_label"] if current_user else "Libre", "perfil"),
            ],
        ),
        unsafe_allow_html=True,
    )
    can_manage_general_config = not storage_backend.auth_enabled() or user_has_permission(
        "manage_general_config",
        current_user,
    )
    can_view_history = not storage_backend.auth_enabled() or user_has_permission("view_history", current_user)

    uploaded_file = st.file_uploader("Sube tu archivo Excel o CSV", type=["xlsx", "xls", "xlsm", "csv"])
    if not uploaded_file:
        if storage_backend.mode != "none" and can_view_history:
            st.info(f"Persistencia activa: `{storage_backend.mode}`")
        else:
            st.info("Sube una base para comenzar.")
        return

    sheet_options = list_excel_sheets(uploaded_file)
    selected_sheet = None
    if sheet_options:
        selected_sheet = st.selectbox("Hoja a procesar", options=sheet_options, index=0)

    try:
        df_raw = load_input_dataframe(uploaded_file, selected_sheet)
    except Exception as exc:
        st.error(f"No pude leer el archivo: {exc}")
        return

    st.subheader("Vista previa")
    st.dataframe(df_raw.head(10), use_container_width=True)

    columnas = df_raw.columns.tolist()
    suggestions = {
        "PLACA": suggest_column(columnas, ["placa", "plate"]),
        "LLEGADA COLA": suggest_column(columnas, ["llegada cola", "t1", "hora cola"]),
        "LLEGADA CASETA": suggest_column(columnas, ["llegada caseta", "t2", "hora caseta"]),
        "SALIDA CASETA": suggest_column(columnas, ["salida caseta", "t3", "hora salida"]),
        "PEAJE": suggest_column(columnas, ["peaje"]),
        "CASETA": suggest_column(columnas, ["caseta"]),
        "SENTIDO": suggest_column(columnas, ["sentido"]),
        "FECHA": suggest_column(columnas, ["fecha", "date"]),
        "VEHICULO": suggest_column(columnas, ["vehiculo", "vehicle"]),
        "T. TEC": suggest_column(columnas, ["t. tec", "t tec", "tec"]),
        "T. CASETA": suggest_column(columnas, ["t. caseta", "t caseta"]),
    }

    st.subheader("Mapeo de columnas")
    col1, col2, col3 = st.columns(3)
    mapping = {}
    selectors = [
        ("PLACA", col1),
        ("LLEGADA COLA", col1),
        ("LLEGADA CASETA", col1),
        ("SALIDA CASETA", col1),
        ("PEAJE", col2),
        ("CASETA", col2),
        ("SENTIDO", col2),
        ("FECHA", col2),
        ("VEHICULO", col3),
        ("T. TEC", col3),
        ("T. CASETA", col3),
    ]
    options_required = [None] + columnas
    for target, container in selectors:
        default = suggestions[target]
        index = options_required.index(default) if default in options_required else 0
        mapping[target] = container.selectbox(
            f"Columna para {target}",
            options=options_required,
            index=index,
            key=f"map_{target}",
        )

    with st.sidebar:
        st.header("Configuracion")
        st.caption(f"Persistencia: `{storage_backend.mode}`")
        if not can_manage_general_config:
            st.caption("Tu rol no puede cambiar la configuracion general. Se aplican los valores definidos por administracion.")
        st.caption("Si desmarcas una opcion, esa etapa no se aplicara en esta corrida.")
        st.markdown("**Placas**")
        config = {
            "aplicar_limpieza_placa": st.checkbox(
                "Revisar y corregir placas automaticamente",
                value=DEFAULT_CONFIG["aplicar_limpieza_placa"],
                disabled=not can_manage_general_config,
                help="Activa toda la etapa de analisis de placas antes de trabajar los tiempos.",
            ),
            "aplicar_ruido_con_respaldo": st.checkbox(
                "Corregir espacios o simbolos cuando la placa limpia ya aparece en la base",
                value=DEFAULT_CONFIG["aplicar_ruido_con_respaldo"],
                disabled=not can_manage_general_config,
                help="Ejemplo: convertir 'CFQ 004' en 'CFQ004' cuando la version limpia ya existe en la base.",
            ),
            "aplicar_ruido_sin_respaldo": st.checkbox(
                "Corregir espacios o simbolos aunque no haya otra coincidencia",
                value=DEFAULT_CONFIG["aplicar_ruido_sin_respaldo"],
                disabled=not can_manage_general_config,
                help="Aplica una limpieza mas agresiva a placas con guiones, espacios o simbolos.",
            ),
            "aplicar_coincidencia_unica_lista": st.checkbox(
                "Corregir placas cuando hay una unica coincidencia clara en la base",
                value=DEFAULT_CONFIG["aplicar_coincidencia_unica_lista"],
                disabled=not can_manage_general_config,
                help="Busca una sola candidata muy parecida dentro de toda la base y corrige hacia esa placa.",
            ),
            "aplicar_confusion_visual": st.checkbox(
                "Corregir confusiones visuales como O/0, I/1 o S/5",
                value=DEFAULT_CONFIG["aplicar_confusion_visual"],
                disabled=not can_manage_general_config,
                help="Corrige placas que parecen tener letras y numeros confundidos visualmente.",
            ),
            "aplicar_recorte_sufijo_x": st.checkbox(
                "Quitar una X final cuando parece un sufijo extra",
                value=DEFAULT_CONFIG["aplicar_recorte_sufijo_x"],
                disabled=not can_manage_general_config,
                help="Se usa en placas donde la X final parece agregada y la version recortada tiene sentido.",
            ),
            "aplicar_exclusion_placeholders": st.checkbox(
                "Excluir placas de prueba o ejemplo",
                value=DEFAULT_CONFIG["aplicar_exclusion_placeholders"],
                disabled=not can_manage_general_config,
                help="Elimina valores como placas placeholder o ejemplos que no deben analizarse.",
            ),
            "aplicar_reglas_manuales": st.checkbox(
                "Aplicar las reglas manuales de la tabla inferior",
                value=DEFAULT_CONFIG["aplicar_reglas_manuales"],
                disabled=not can_manage_general_config,
                help="Usa las correcciones o exclusiones definidas manualmente por ti.",
            ),
        }
        st.markdown("**Tiempos**")
        config.update(
            {
                "eliminar_bordes_caseta": st.checkbox(
                    "Eliminar registros incompletos al inicio o final de cada flujo",
                    value=DEFAULT_CONFIG["eliminar_bordes_caseta"],
                    disabled=not can_manage_general_config,
                    help="Quita filas incompletas que quedan antes de que el flujo empiece bien o despues de que ya termino.",
                ),
                "aplicar_interpolacion": st.checkbox(
                    "Completar tiempos faltantes usando vehiculos vecinos del mismo flujo",
                    value=DEFAULT_CONFIG["aplicar_interpolacion"],
                    disabled=not can_manage_general_config,
                    help="Reconstruye tiempos internos cuando hay anclas antes y despues dentro del mismo flujo.",
                ),
                "aplicar_mediana_local": st.checkbox(
                    "Hacer una segunda recuperacion usando medianas del mismo flujo",
                    value=DEFAULT_CONFIG["aplicar_mediana_local"],
                    disabled=not can_manage_general_config,
                    help="Intenta completar tiempos pendientes con duraciones tipicas observadas en ese mismo flujo.",
                ),
                "aplicar_donantes": st.checkbox(
                    "Hacer una tercera recuperacion usando casos similares",
                    value=DEFAULT_CONFIG["aplicar_donantes"],
                    disabled=not can_manage_general_config,
                    help="Usa medianas de casetas, dias o sentidos parecidos como apoyo para recuperar tiempos.",
                ),
                "aplicar_swap_tiempos_completos_cortos": st.checkbox(
                    "Corregir inversiones cortas en registros completos",
                    value=DEFAULT_CONFIG["aplicar_swap_tiempos_completos_cortos"],
                    disabled=not can_manage_general_config,
                    help="Intercambia T1-T2 o T2-T3 cuando quedaron invertidos por pocos segundos dentro de un registro completo.",
                ),
                "modo_contraste_estricto": st.checkbox(
                    "Excluir de la base limpia final los casos recuperados por mediana local y donantes",
                    value=DEFAULT_CONFIG["modo_contraste_estricto"],
                    disabled=not can_manage_general_config,
                    help="Mantiene la interpolacion, pero saca de la salida final los casos mas dificiles de contrastar directamente con peaje.",
                ),
                "aplicar_swap_final_t2_t3": st.checkbox(
                    "Corregir un posible intercambio final entre llegada y salida de caseta",
                    value=DEFAULT_CONFIG["aplicar_swap_final_t2_t3"],
                    disabled=not can_manage_general_config,
                    help="Aplica un ajuste final si los tiempos de llegada y salida quedaron invertidos.",
                ),
            }
        )

    st.subheader("Reglas manuales")
    reglas_df = st.data_editor(
        pd.DataFrame(DEFAULT_MANUAL_RULES),
        num_rows="dynamic",
        use_container_width=True,
        disabled=not can_manage_general_config,
        key="manual_rules_editor",
    )

    uploaded_file_bytes = uploaded_file.getvalue()
    current_signature = build_processing_signature(
        uploaded_file.name,
        selected_sheet,
        uploaded_file_bytes,
        mapping,
        config,
        reglas_df,
    )
    processed_payload = st.session_state.get(TEC_RESULT_STATE_KEY)

    if st.button("Procesar base", type="primary", use_container_width=True):
        faltantes = [campo for campo in ["PLACA", "LLEGADA COLA", "LLEGADA CASETA", "SALIDA CASETA"] if not mapping.get(campo)]
        if faltantes:
            st.error(f"Falta mapear estas columnas obligatorias: {', '.join(faltantes)}")
            return
        if config["eliminar_bordes_caseta"] and not mapping.get("FECHA"):
            st.error("Para procesar tiempos necesitas mapear FECHA.")
            return

        df_std = build_standardized_df(df_raw, mapping)
        if df_std["FECHA"].isna().all():
            st.warning("La columna FECHA quedo vacia; los pasos de tiempos pueden perder precision o fallar.")

        try:
            result = process_pipeline(df_std, config, reglas_df)
        except Exception as exc:
            st.exception(exc)
            return

        export_tables = result["export_tables"]
        storage_backend.save_run(
            build_run_payload(uploaded_file.name, mapping, config, export_tables)
        )

        processed_payload = build_processing_artifacts(uploaded_file.name, selected_sheet, current_signature, result)
        st.session_state[TEC_RESULT_STATE_KEY] = processed_payload

    if processed_payload and processed_payload.get("input_signature") == current_signature:
        render_processing_outputs(processed_payload, storage_backend, can_view_history)
    elif processed_payload:
        st.info("Hay resultados guardados de otra base u hoja. Procesa nuevamente para actualizar las descargas.")


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    inject_global_styles()
    storage_backend = build_storage_backend(load_streamlit_secrets())

    if storage_backend.auth_enabled():
        current_user = get_authenticated_user()
        if not current_user:
            with st.sidebar:
                render_contractor_branding()
            render_login_gate(storage_backend)
            return
    else:
        current_user = None

    can_open_user_management = storage_backend.auth_enabled() and (
        user_has_permission("manage_users", current_user)
        or user_has_permission("manage_user_activation", current_user)
    )

    can_view_history = storage_backend.mode != "none" and (
        not storage_backend.auth_enabled() or user_has_permission("view_history", current_user)
    )

    page_options = ["Inicio"]
    if not storage_backend.auth_enabled() or user_has_permission("process_files", current_user):
        page_options.append("TEC")
    page_options.extend(["Relevamientos", "Auditorias", "Satisfaccion", "Flujogramas"])

    requested_page = get_requested_page()
    if requested_page:
        st.session_state[APP_NAV_KEY] = requested_page

    current_page = st.session_state.get(APP_NAV_KEY, "Inicio")
    valid_pages = set(page_options)
    if can_open_user_management:
        valid_pages.add("Usuarios")
    if can_view_history:
        valid_pages.add("Historial")
    if current_page not in valid_pages:
        current_page = "Inicio"
        st.session_state[APP_NAV_KEY] = current_page

    with st.sidebar:
        if current_user:
            if current_page != "TEC":
                render_contractor_branding()
            st.divider()
            st.header("Sesion")
            st.write(current_user["full_name"])
            st.caption(f"Usuario: {current_user['username']}")
            st.caption(f"Rol: {current_user['role_label']}")
            st.caption(f"Version app: {get_runtime_version_label()}")
            st.caption(describe_access_window(current_user.get("active_from"), current_user.get("active_until")))
            action_cols = st.columns([5, 1]) if can_open_user_management else st.columns([1])
            with action_cols[0]:
                if st.button("Cerrar sesion", key="sidebar_logout_button", use_container_width=True):
                    st.session_state[APP_NAV_KEY] = "Inicio"
                    st.query_params.clear()
                    clear_authenticated_user()
                    st.rerun()
            if can_open_user_management:
                with action_cols[1]:
                    if st.button(
                        "⚙",
                        key="sidebar_users_button",
                        help="Usuarios",
                        use_container_width=True,
                        type="primary" if current_page == "Usuarios" else "secondary",
                    ):
                        navigate_to("Usuarios")

            if can_view_history:
                st.markdown(
                    '<div class="sidebar-text-link"><a href="?page=Historial">Historial</a></div>',
                    unsafe_allow_html=True,
                )
            st.divider()
        page = current_page

    if page == "Inicio":
        render_home_page(current_user)
        return
    if page == "Usuarios":
        render_user_management_page(storage_backend, current_user)
        return
    if page == "Historial":
        render_history_page(storage_backend)
        return
    if page in MODULE_PLACEHOLDERS:
        render_placeholder_module_page(page)
        return

    render_processing_page(storage_backend, current_user)


if __name__ == "__main__":
    main()
