from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from unicodedata import normalize
import sys

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


APP_DIR = Path(__file__).resolve().parent
BASE_DIR = APP_DIR.parent
CONTRASTE_DIR = BASE_DIR / "Contraste"
RAW_PATH = BASE_DIR / "Data TEC Norvial.xlsx"
OUTPUT_CLEAN_PATH = CONTRASTE_DIR / "Data TEC Norvial_limpio_ajustado_pruebas.xlsx"
OUTPUT_WORD_PATH = CONTRASTE_DIR / "Informe contraste base limpia vs peaje.docx"
OUTPUT_EXCEL_PATH = CONTRASTE_DIR / "Informe contraste base limpia vs peaje.xlsx"

if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

import app_tec_norvial_streamlit as app


def normalize_text(value: Any) -> str | None:
    if pd.isna(value):
        return None
    text = str(value).strip()
    text = normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    return text.upper()


def parse_clock_value(value: Any) -> pd.Timestamp | pd.NaT:
    if pd.isna(value):
        return pd.NaT
    text = str(value).strip()
    parsed = pd.to_datetime(text, format="%H:%M:%S", errors="coerce")
    if pd.isna(parsed):
        parsed = pd.to_datetime(text, errors="coerce")
    return parsed


def safe_pct(numerator: float, denominator: float) -> float:
    if denominator in (0, 0.0) or pd.isna(denominator):
        return np.nan
    return round(float(numerator) / float(denominator) * 100, 2)


def prep_clean(df: pd.DataFrame, source_label: str) -> pd.DataFrame:
    out = df.copy()
    out["PEAJE_N"] = out["PEAJE"].map(normalize_text)
    out["SENTIDO_N"] = out["SENTIDO"].map(normalize_text)
    out["PLACA_N"] = out["PLACA"].map(normalize_text)
    out["CASETA_TXT"] = out["CASETA"].astype(str).str.strip()
    out["FECHA_D"] = pd.to_datetime(out["FECHA"], errors="coerce").dt.normalize()
    out["TIME_PART"] = out["LLEGADA CASETA"].apply(parse_clock_value)
    out["DATETIME_CASETA"] = out["FECHA_D"] + pd.to_timedelta(out["TIME_PART"].dt.strftime("%H:%M:%S"))
    out["HOUR_BLOCK"] = out["DATETIME_CASETA"].dt.floor("h")
    out["SOURCE_LABEL"] = source_label
    return out


def prep_official(path: Path) -> pd.DataFrame:
    df = pd.read_excel(path, sheet_name="Hoja1", header=2)
    out = df.copy()
    out["PEAJE_N"] = out["Estacion de peaje"].map(normalize_text)
    out["SENTIDO_N"] = out["Sentido del transito"].map(normalize_text)
    out["PLACA_N"] = out["Placa del vehiculo"].map(normalize_text)
    out["FECHA_D"] = pd.to_datetime(out["Fecha"], dayfirst=True, errors="coerce").dt.normalize()
    out["TIME_PART"] = out["Hora"].apply(parse_clock_value)
    out["DATETIME_CASETA"] = out["FECHA_D"] + pd.to_timedelta(out["TIME_PART"].dt.strftime("%H:%M:%S"))
    out["HOUR_BLOCK"] = out["DATETIME_CASETA"].dt.floor("h")
    out["VIA_TXT"] = out["Via"].astype(str).str.strip()
    return out


def match_by_plate_time(official_df: pd.DataFrame, tech_df: pd.DataFrame, tolerance_seconds: int) -> tuple[int, int]:
    left = official_df[["PEAJE_N", "FECHA_D", "SENTIDO_N", "PLACA_N", "DATETIME_CASETA"]].dropna().sort_values("DATETIME_CASETA").copy()
    right = tech_df[["PEAJE_N", "FECHA_D", "SENTIDO_N", "PLACA_N", "DATETIME_CASETA"]].dropna().sort_values("DATETIME_CASETA").copy()
    if left.empty:
        return 0, 0
    if right.empty:
        return 0, len(left)
    left["group_key"] = left[["PEAJE_N", "FECHA_D", "SENTIDO_N", "PLACA_N"]].astype(str).agg("|".join, axis=1)
    right["group_key"] = right[["PEAJE_N", "FECHA_D", "SENTIDO_N", "PLACA_N"]].astype(str).agg("|".join, axis=1)
    merged = pd.merge_asof(
        left.sort_values("DATETIME_CASETA"),
        right.sort_values("DATETIME_CASETA"),
        on="DATETIME_CASETA",
        by="group_key",
        direction="nearest",
        tolerance=pd.Timedelta(seconds=tolerance_seconds),
        suffixes=("_oficial", "_tec"),
    )
    return int(merged["PLACA_N_tec"].notna().sum()), int(len(left))


def build_hourly_compare(official_df: pd.DataFrame, tech_df: pd.DataFrame, label: str) -> pd.DataFrame:
    official_hourly = official_df.groupby("HOUR_BLOCK").size().rename("oficial").reset_index()
    tech_hourly = tech_df.groupby("HOUR_BLOCK").size().rename(label).reset_index()
    comp = official_hourly.merge(tech_hourly, on="HOUR_BLOCK", how="outer").fillna(0)
    comp["oficial"] = comp["oficial"].astype(int)
    comp[label] = comp[label].astype(int)
    comp["brecha_abs"] = comp[label] - comp["oficial"]
    comp["brecha_pct_vs_oficial"] = np.where(comp["oficial"] > 0, comp["brecha_abs"] / comp["oficial"] * 100, np.nan)
    comp["cobertura_pct"] = np.where(comp["oficial"] > 0, comp[label] / comp["oficial"] * 100, np.nan)
    comp["hora"] = comp["HOUR_BLOCK"].dt.strftime("%Y-%m-%d %H:%M")
    return comp[["hora", "oficial", label, "brecha_abs", "brecha_pct_vs_oficial", "cobertura_pct"]].sort_values("hora")


def build_sentido_compare(official_df: pd.DataFrame, tech_df: pd.DataFrame, label: str) -> pd.DataFrame:
    official_group = official_df.groupby(["HOUR_BLOCK", "SENTIDO_N"]).size().rename("oficial").reset_index()
    tech_group = tech_df.groupby(["HOUR_BLOCK", "SENTIDO_N"]).size().rename(label).reset_index()
    comp = official_group.merge(tech_group, on=["HOUR_BLOCK", "SENTIDO_N"], how="outer").fillna(0)
    comp["oficial"] = comp["oficial"].astype(int)
    comp[label] = comp[label].astype(int)
    comp["brecha_abs"] = comp[label] - comp["oficial"]
    comp["brecha_pct_vs_oficial"] = np.where(comp["oficial"] > 0, comp["brecha_abs"] / comp["oficial"] * 100, np.nan)
    comp["cobertura_pct"] = np.where(comp["oficial"] > 0, comp[label] / comp["oficial"] * 100, np.nan)
    comp["hora"] = comp["HOUR_BLOCK"].dt.strftime("%Y-%m-%d %H:%M")
    return comp[["hora", "SENTIDO_N", "oficial", label, "brecha_abs", "brecha_pct_vs_oficial", "cobertura_pct"]].sort_values(["hora", "SENTIDO_N"])


def build_via_compare(official_df: pd.DataFrame, tech_df: pd.DataFrame, label: str) -> pd.DataFrame:
    official_group = official_df.groupby("VIA_TXT").size().rename("oficial").reset_index()
    tech_group = tech_df.groupby("CASETA_TXT").size().rename(label).reset_index().rename(columns={"CASETA_TXT": "VIA_TXT"})
    comp = official_group.merge(tech_group, on="VIA_TXT", how="outer").fillna(0)
    comp["oficial"] = comp["oficial"].astype(int)
    comp[label] = comp[label].astype(int)
    comp["brecha_abs"] = comp[label] - comp["oficial"]
    comp["brecha_pct_vs_oficial"] = np.where(comp["oficial"] > 0, comp["brecha_abs"] / comp["oficial"] * 100, np.nan)
    comp["cobertura_pct"] = np.where(comp["oficial"] > 0, comp[label] / comp["oficial"] * 100, np.nan)
    return comp.sort_values("VIA_TXT")


def build_operating_windows(tech_df: pd.DataFrame) -> pd.DataFrame:
    windows = (
        tech_df.dropna(subset=["DATETIME_CASETA"])
        .groupby(["CASETA_TXT", "SENTIDO_N"], dropna=False)
        .agg(
            inicio_ventana=("DATETIME_CASETA", "min"),
            fin_ventana=("DATETIME_CASETA", "max"),
            vehiculos_tec=("PLACA_N", "size"),
        )
        .reset_index()
        .sort_values(["CASETA_TXT", "SENTIDO_N"])
    )
    return windows


def filter_to_operating_windows(source_df: pd.DataFrame, windows_df: pd.DataFrame, caseta_col: str) -> pd.DataFrame:
    if source_df.empty or windows_df.empty:
        return source_df.iloc[0:0].copy()

    merged = source_df.merge(
        windows_df,
        left_on=[caseta_col, "SENTIDO_N"],
        right_on=["CASETA_TXT", "SENTIDO_N"],
        how="inner",
        suffixes=("", "_VENTANA"),
    )
    filtered = merged[
        merged["DATETIME_CASETA"].between(merged["inicio_ventana"], merged["fin_ventana"], inclusive="both")
    ].copy()
    drop_cols = ["CASETA_TXT_VENTANA"] if "CASETA_TXT_VENTANA" in filtered.columns else []
    return filtered.drop(columns=drop_cols, errors="ignore")


def summarize_time_metrics(df_resultados: pd.DataFrame, group_cols: list[str]) -> pd.DataFrame:
    grouped = (
        df_resultados.groupby(group_cols, dropna=False)
        .agg(
            vehiculos=("PLACA_FINAL", "size"),
            t_cola_min_min=("T_COLA_FINAL_MINUTOS", "min"),
            t_cola_max_min=("T_COLA_FINAL_MINUTOS", "max"),
            t_cola_prom_min=("T_COLA_FINAL_MINUTOS", "mean"),
            t_caseta_min_min=("T_CASETA_FINAL_MINUTOS", "min"),
            t_caseta_max_min=("T_CASETA_FINAL_MINUTOS", "max"),
            t_caseta_prom_min=("T_CASETA_FINAL_MINUTOS", "mean"),
            t_tec_min_min=("T_TEC_FINAL_MINUTOS", "min"),
            t_tec_max_min=("T_TEC_FINAL_MINUTOS", "max"),
            t_tec_prom_min=("T_TEC_FINAL_MINUTOS", "mean"),
            cola_min_veh=("COLA_ESPERA_USUARIOS", "min"),
            cola_max_veh=("COLA_ESPERA_USUARIOS", "max"),
            cola_prom_veh=("COLA_ESPERA_USUARIOS", "mean"),
        )
        .reset_index()
    )
    numeric_cols = [
        "t_cola_min_min",
        "t_cola_max_min",
        "t_cola_prom_min",
        "t_caseta_min_min",
        "t_caseta_max_min",
        "t_caseta_prom_min",
        "t_tec_min_min",
        "t_tec_max_min",
        "t_tec_prom_min",
        "cola_min_veh",
        "cola_max_veh",
        "cola_prom_veh",
    ]
    grouped[numeric_cols] = grouped[numeric_cols].round(2)
    return grouped.sort_values(group_cols)


def build_findings(summary_df: pd.DataFrame, hourly_tables: dict[str, pd.DataFrame], sentido_tables: dict[str, pd.DataFrame]) -> list[str]:
    findings: list[str] = []
    if summary_df.empty:
        return ["No se encontraron archivos oficiales para contrastar la base limpia."]

    for row in summary_df.to_dict("records"):
        sheet_key = row["sheet_key"]
        hourly = hourly_tables[sheet_key]
        sentido = sentido_tables[sheet_key]
        worst_hour = hourly.sort_values("brecha_abs").iloc[0]
        worst_direction = sentido.sort_values("brecha_abs").iloc[0]
        findings.append(
            (
                f"{row['peaje']} {row['fecha']}: usando la ventana real de operacion observada en TEC por caseta y sentido, "
                f"la base limpia cubre {row['cobertura_limpia_pct']:.1f}% del aforo oficial filtrado y la coincidencia por placa y hora de llegada a caseta en ventana de ±60 s alcanza {row['match_60s_pct']:.1f}%."
            )
        )
        findings.append(
            (
                f"El bloque horario mas critico en {row['peaje']} dentro de las ventanas activas ocurre a las {worst_hour['hora'][-5:]}, "
                f"con brecha de {int(worst_hour['brecha_abs'])} vehiculos y cobertura de {worst_hour['cobertura_pct']:.1f}% frente al oficial filtrado."
            )
        )
        findings.append(
            (
                f"La mayor desviacion por sentido en {row['peaje']} se observa en {worst_direction['SENTIDO_N']}, "
                f"bloque {worst_direction['hora'][-5:]}, con brecha de {int(worst_direction['brecha_abs'])} vehiculos."
            )
        )
    return findings


def render_overview_chart(summary_df: pd.DataFrame) -> bytes:
    labels = [f"{row.peaje}\n{row.fecha}" for row in summary_df.itertuples(index=False)]
    x = np.arange(len(labels))
    width = 0.24
    fig, ax = plt.subplots(figsize=(10, 5), dpi=180)
    ax.bar(x - width, summary_df["cobertura_cruda_pct"], width=width, label="Cobertura base cruda", color="#9bb6e8")
    ax.bar(x, summary_df["cobertura_limpia_pct"], width=width, label="Cobertura base limpia", color="#2f6ddc")
    ax.bar(x + width, summary_df["match_60s_pct"], width=width, label="Coincidencia placa+hora ±60 s", color="#123d82")
    ax.set_ylim(0, max(100, float(summary_df[["cobertura_cruda_pct", "cobertura_limpia_pct", "match_60s_pct"]].max().max()) + 10))
    ax.set_ylabel("Porcentaje (%)")
    ax.set_title("Indicadores porcentuales de contraste por peaje y fecha", fontsize=13, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.35)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(frameon=False, loc="upper right")
    for bars in ax.containers:
        ax.bar_label(bars, fmt="%.1f%%", fontsize=7, padding=3)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer.getvalue()


def render_hourly_chart(hourly_df: pd.DataFrame, peaje: str, fecha: str) -> bytes:
    fig, ax = plt.subplots(figsize=(10, 4.6), dpi=180)
    labels = [value[-5:] for value in hourly_df["hora"]]
    ax.plot(labels, hourly_df["oficial"], label="Oficial", color="#1c2f52", linewidth=2.0)
    ax.plot(labels, hourly_df["base_limpia"], label="Base limpia", color="#2f6ddc", linewidth=2.0)
    ax.fill_between(labels, hourly_df["oficial"], hourly_df["base_limpia"], color="#c9d8f4", alpha=0.45)
    ax.set_title(f"Comparacion horaria de aforos - {peaje} {fecha}", fontsize=12, fontweight="bold")
    ax.set_ylabel("Vehiculos por hora")
    ax.set_xlabel("Hora")
    ax.grid(axis="y", linestyle="--", linewidth=0.5, alpha=0.35)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(frameon=False)
    ax.tick_params(axis="x", labelrotation=45, labelsize=8)
    fig.tight_layout()
    buffer = BytesIO()
    fig.savefig(buffer, format="png", facecolor="white", bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer.getvalue()


def add_table(doc: Document, title: str, table_df: pd.DataFrame, max_rows: int | None = None) -> None:
    doc.add_paragraph(title, style="Heading 2")
    if table_df.empty:
        doc.add_paragraph("No se dispone de filas suficientes para esta tabla.")
        return
    visible = table_df.head(max_rows) if max_rows else table_df
    table = doc.add_table(rows=1, cols=len(visible.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, column in enumerate(visible.columns):
        table.rows[0].cells[idx].text = str(column)
    for row in visible.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)


def build_word_report(summary_df: pd.DataFrame, findings: list[str], hourly_tables: dict[str, pd.DataFrame], time_tables: dict[str, pd.DataFrame]) -> None:
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informe de contraste entre base limpia TEC y data oficial de peaje")
    run.bold = True
    run.font.size = Pt(19)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    doc.add_paragraph(
        "Este informe compara la base limpia reconstruida con la pipeline vigente del aplicativo contra los archivos oficiales del peaje, "
        "recortando primero la data oficial y la base cruda a la ventana real de atencion observada en TEC para cada caseta, sentido y fecha. "
        "De ese modo, la comparacion parte del primer registro TEC en llegada a caseta y termina en el ultimo registro TEC de la misma caseta y sentido. "
        "Se incluyen contrastes por volumen total, coincidencia de placa-tiempo, brechas por hora, brechas por sentido y resumen operativo de tiempos TEC.",
    )

    doc.add_paragraph("Criterio de comparacion", style="Heading 1")
    doc.add_paragraph(
        "La comparacion no se realiza por bloque horario calendario completo. En su lugar, para cada peaje, caseta, sentido y fecha, "
        "se identifica la ventana real de atencion observada en la base TEC usando la hora de llegada a caseta."
    )
    doc.add_paragraph(
        "El inicio de la ventana corresponde al primer registro TEC de llegada a caseta y el fin corresponde al ultimo registro TEC de llegada a caseta "
        "para esa misma combinacion operativa.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Una vez definida esa ventana, la data oficial del peaje y la base cruda se filtran para conservar unicamente los registros que caen dentro de ese mismo tramo temporal.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Por tanto, si una caseta TEC inicia su atencion a las 17:43, la comparacion oficial para esa caseta no empieza a las 17:00, sino desde la primera observacion oficial disponible dentro de la ventana 17:43 en adelante.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Este criterio evita sobreestimar las brechas por comparar periodos donde el aforo oficial ya estaba registrando vehiculos pero la caseta aun no habia iniciado atencion en la base TEC.",
        style="List Bullet",
    )

    doc.add_paragraph("Hallazgos ejecutivos", style="Heading 1")
    for finding in findings:
        doc.add_paragraph(finding, style="List Bullet")

    add_table(
        doc,
        "Resumen comparativo principal",
        summary_df[
            [
                "peaje",
                "fecha",
                "oficial_total",
                "base_cruda_total",
                "base_limpia_total",
                "cobertura_cruda_pct",
                "cobertura_limpia_pct",
                "match_60s_pct",
                "mae_horaria_limpia",
                "max_brecha_horaria_limpia",
            ]
        ].rename(
            columns={
                "peaje": "Peaje",
                "fecha": "Fecha",
                "oficial_total": "Oficial",
                "base_cruda_total": "Base cruda",
                "base_limpia_total": "Base limpia",
                "cobertura_cruda_pct": "Cobertura cruda (%)",
                "cobertura_limpia_pct": "Cobertura limpia (%)",
                "match_60s_pct": "Match ±60 s (%)",
                "mae_horaria_limpia": "MAE horario",
                "max_brecha_horaria_limpia": "Brecha horaria maxima",
            }
        ),
    )

    overview_chart = render_overview_chart(summary_df)
    doc.add_paragraph("Graficas de indicadores porcentuales", style="Heading 1")
    doc.add_picture(BytesIO(overview_chart), width=Inches(6.8))

    for row in summary_df.itertuples(index=False):
        doc.add_paragraph(f"Detalle de {row.peaje} {row.fecha}", style="Heading 1")
        doc.add_paragraph(
            f"La base limpia registra {row.base_limpia_total} vehiculos frente a {row.oficial_total} del aforo oficial, "
            f"con cobertura de {row.cobertura_limpia_pct:.1f}% y coincidencia placa-tiempo ±60 s de {row.match_60s_pct:.1f}%.",
        )
        hourly_chart = render_hourly_chart(hourly_tables[row.sheet_key], row.peaje, row.fecha)
        doc.add_picture(BytesIO(hourly_chart), width=Inches(6.8))

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Resultados TEC para soporte del informe", style="Heading 1")
    add_table(
        doc,
        "Resumen por peaje y sentido con minimos, maximos y promedios",
        time_tables["tiempos_peaje_sentido"].rename(
            columns={
                "PEAJE": "Peaje",
                "SENTIDO": "Sentido",
                "vehiculos": "Vehiculos",
                "t_cola_min_min": "T.cola min",
                "t_cola_max_min": "T.cola max",
                "t_cola_prom_min": "T.cola prom",
                "t_caseta_min_min": "T.caseta min",
                "t_caseta_max_min": "T.caseta max",
                "t_caseta_prom_min": "T.caseta prom",
                "t_tec_min_min": "TEC min",
                "t_tec_max_min": "TEC max",
                "t_tec_prom_min": "TEC prom",
            }
        ),
        max_rows=20,
    )

    doc.save(OUTPUT_WORD_PATH)


def style_dataframe_sheet(worksheet, df: pd.DataFrame, title: str, note: str | None = None, start_row: int = 1) -> int:
    title_cell = worksheet.cell(row=start_row, column=1, value=title)
    title_cell.font = Font(name="Aptos Display", size=14, bold=True, color="FFFFFF")
    title_cell.fill = PatternFill("solid", fgColor="173F7A")
    title_cell.alignment = Alignment(horizontal="left")
    end_column = max(1, len(df.columns))
    if end_column > 1:
        worksheet.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=end_column)

    data_start = start_row + 1
    if note:
        note_cell = worksheet.cell(row=start_row + 1, column=1, value=note)
        note_cell.font = Font(name="Aptos", size=10, italic=True, color="5B6579")
        if end_column > 1:
            worksheet.merge_cells(start_row=start_row + 1, start_column=1, end_row=start_row + 1, end_column=end_column)
        data_start = start_row + 2

    header_row = data_start
    for col_idx, column_name in enumerate(df.columns, start=1):
        cell = worksheet.cell(row=header_row, column=col_idx, value=column_name)
        cell.font = Font(name="Aptos", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="3B6DB1")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    thin = Side(style="thin", color="D9E2F1")
    for row_idx, row in enumerate(df.itertuples(index=False), start=header_row + 1):
        fill = PatternFill("solid", fgColor="F7FAFF" if (row_idx - header_row) % 2 else "EDF3FC")
        for col_idx, value in enumerate(row, start=1):
            cell = worksheet.cell(row=row_idx, column=col_idx, value=None if pd.isna(value) else value)
            cell.alignment = Alignment(horizontal="center" if isinstance(value, (int, float, np.number)) else "left", vertical="center")
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.fill = fill
            if isinstance(value, (int, float, np.number)):
                column_lower = str(df.columns[col_idx - 1]).lower()
                if "pct" in column_lower or "%" in column_lower:
                    cell.number_format = '0.00'
                elif "min" in column_lower or "prom" in column_lower or "mae" in column_lower:
                    cell.number_format = '0.00'
                else:
                    cell.number_format = '0'

    worksheet.freeze_panes = worksheet.cell(row=header_row + 1, column=1)
    worksheet.auto_filter.ref = f"A{header_row}:{get_column_letter(end_column)}{max(header_row, header_row + len(df))}"

    for col_idx, column_name in enumerate(df.columns, start=1):
        values = [str(column_name)] + ["" if pd.isna(value) else str(value) for value in df.iloc[:, col_idx - 1].head(150)]
        width = min(max(len(value) for value in values) + 2, 28)
        worksheet.column_dimensions[get_column_letter(col_idx)].width = width

    return header_row + len(df) + 2


def build_excel_report(summary_df: pd.DataFrame, findings: list[str], time_tables: dict[str, pd.DataFrame], hourly_tables: dict[str, pd.DataFrame], sentido_tables: dict[str, pd.DataFrame], via_tables: dict[str, pd.DataFrame]) -> None:
    with pd.ExcelWriter(OUTPUT_EXCEL_PATH, engine="openpyxl") as writer:
        portada = pd.DataFrame({"Hallazgos ejecutivos": findings})
        portada.to_excel(writer, sheet_name="00_portada", index=False, startrow=2)
        ws_portada = writer.sheets["00_portada"]
        style_dataframe_sheet(
            ws_portada,
            portada,
            "Informe de contraste base limpia vs peaje",
            "La base limpia se reconstruyo con la pipeline vigente y se comparo contra el aforo oficial filtrado por la ventana real de atencion en TEC para cada caseta y sentido.",
            start_row=1,
        )

        resumen_sheet = summary_df[
            [
                "peaje",
                "fecha",
                "oficial_total",
                "base_cruda_total",
                "base_limpia_total",
                "cobertura_cruda_pct",
                "cobertura_limpia_pct",
                "match_exacto_pct",
                "match_30s_pct",
                "match_60s_pct",
                "match_300s_pct",
                "mae_horaria_limpia",
                "max_brecha_horaria_limpia",
            ]
        ].rename(
            columns={
                "peaje": "Peaje",
                "fecha": "Fecha",
                "oficial_total": "Oficial",
                "base_cruda_total": "Base cruda",
                "base_limpia_total": "Base limpia",
                "cobertura_cruda_pct": "Cobertura cruda (%)",
                "cobertura_limpia_pct": "Cobertura limpia (%)",
                "match_exacto_pct": "Match exacto (%)",
                "match_30s_pct": "Match ±30 s (%)",
                "match_60s_pct": "Match ±60 s (%)",
                "match_300s_pct": "Match ±300 s (%)",
                "mae_horaria_limpia": "MAE horario limpia",
                "max_brecha_horaria_limpia": "Brecha horaria maxima",
            }
        )
        resumen_sheet.to_excel(writer, sheet_name="01_resumen", index=False, startrow=2)
        style_dataframe_sheet(writer.sheets["01_resumen"], resumen_sheet, "Resumen ejecutivo del contraste", start_row=1)

        ordered_time_sheets = [
            ("02_tiempos_peaje", "Tiempos TEC por peaje", time_tables["tiempos_peaje"]),
            ("03_tiempos_sentido", "Tiempos TEC por peaje y sentido", time_tables["tiempos_peaje_sentido"]),
            ("04_tiempos_caseta", "Tiempos TEC por peaje, caseta y sentido", time_tables["tiempos_caseta"]),
        ]
        for sheet_name, title, table_df in ordered_time_sheets:
            table_df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=2)
            style_dataframe_sheet(
                writer.sheets[sheet_name],
                table_df,
                title,
                "Las columnas de tiempo se presentan con minimo, maximo y promedio para facilitar su uso directo en el informe.",
                start_row=1,
            )

        for idx, row in enumerate(summary_df.itertuples(index=False), start=5):
            hourly = hourly_tables[row.sheet_key].rename(
                columns={
                    "hora": "Hora",
                    "oficial": "Oficial",
                    "base_limpia": "Base limpia",
                    "brecha_abs": "Brecha absoluta",
                    "brecha_pct_vs_oficial": "Brecha vs oficial (%)",
                    "cobertura_pct": "Cobertura limpia (%)",
                }
            )
            sentido = sentido_tables[row.sheet_key].rename(
                columns={
                    "hora": "Hora",
                    "SENTIDO_N": "Sentido",
                    "oficial": "Oficial",
                    "base_limpia": "Base limpia",
                    "brecha_abs": "Brecha absoluta",
                    "brecha_pct_vs_oficial": "Brecha vs oficial (%)",
                    "cobertura_pct": "Cobertura limpia (%)",
                }
            )
            via = via_tables[row.sheet_key].rename(
                columns={
                    "VIA_TXT": "Via/Caseta",
                    "oficial": "Oficial",
                    "base_limpia": "Base limpia",
                    "brecha_abs": "Brecha absoluta",
                    "brecha_pct_vs_oficial": "Brecha vs oficial (%)",
                    "cobertura_pct": "Cobertura limpia (%)",
                }
            )

            sheet_h = f"{idx:02d}_hora_{row.peaje[:4]}"[:31]
            sheet_s = f"{idx:02d}_sent_{row.peaje[:4]}"[:31]
            sheet_v = f"{idx:02d}_via_{row.peaje[:4]}"[:31]
            hourly.to_excel(writer, sheet_name=sheet_h, index=False, startrow=2)
            sentido.to_excel(writer, sheet_name=sheet_s, index=False, startrow=2)
            via.to_excel(writer, sheet_name=sheet_v, index=False, startrow=2)
            style_dataframe_sheet(writer.sheets[sheet_h], hourly, f"Brechas horarias - {row.peaje} {row.fecha}", start_row=1)
            style_dataframe_sheet(writer.sheets[sheet_s], sentido, f"Brechas por sentido - {row.peaje} {row.fecha}", start_row=1)
            style_dataframe_sheet(writer.sheets[sheet_v], via, f"Comparacion por via/caseta - {row.peaje} {row.fecha}", start_row=1)


def regenerate_clean_base() -> tuple[dict[str, pd.DataFrame], pd.DataFrame, pd.DataFrame]:
    raw_df = pd.read_excel(RAW_PATH, sheet_name="peajes")
    mapping = {
        "PEAJE": "PEAJE",
        "CASETA": "CASETA",
        "SENTIDO": "SENTIDO",
        "FECHA": "FECHA",
        "VEHICULO": "VEHICULO",
        "PLACA": "PLACA",
        "LLEGADA COLA": "LLEGADA COLA",
        "LLEGADA CASETA": "LLEGADA CASETA",
        "SALIDA CASETA": "SALIDA CASETA",
        "T. TEC": "T. TEC",
        "T. CASETA": "T. CASETA",
    }
    config = dict(app.DEFAULT_CONFIG)
    manual_rules_df = pd.DataFrame(app.DEFAULT_MANUAL_RULES)
    df_std = app.build_standardized_df(raw_df, mapping)
    result = app.process_pipeline(df_std, config, manual_rules_df)
    export_tables = result["export_tables"]
    with pd.ExcelWriter(OUTPUT_CLEAN_PATH, engine="openpyxl") as writer:
        export_tables["base_limpia"].to_excel(writer, sheet_name="base_limpia", index=False)
        export_tables["casos_eliminados"].to_excel(writer, sheet_name="casos_eliminados", index=False)
        export_tables["reporte_resumen"].to_excel(writer, sheet_name="reporte_resumen", index=False)
    df_resultados = app.build_resultados_dataframe(export_tables["export_base_detalle"])
    return export_tables, raw_df, df_resultados


def build_contrast_outputs(export_tables: dict[str, pd.DataFrame], raw_df: pd.DataFrame, df_resultados: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, pd.DataFrame], dict[str, pd.DataFrame], dict[str, pd.DataFrame], dict[str, pd.DataFrame]]:
    raw_prep = prep_clean(raw_df, "base_cruda")
    clean_prep = prep_clean(export_tables["base_limpia"], "base_limpia")

    summary_rows: list[dict[str, Any]] = []
    hourly_tables: dict[str, pd.DataFrame] = {}
    sentido_tables: dict[str, pd.DataFrame] = {}
    via_tables: dict[str, pd.DataFrame] = {}

    for official_path in sorted(CONTRASTE_DIR.glob("Tra*.xlsx")):
        official_df = prep_official(official_path)
        peaje = official_df["PEAJE_N"].dropna().iloc[0]
        fecha = official_df["FECHA_D"].dropna().iloc[0]
        official_day = official_df[(official_df["PEAJE_N"] == peaje) & (official_df["FECHA_D"] == fecha) & official_df["DATETIME_CASETA"].notna()].copy()
        raw_day = raw_prep[(raw_prep["PEAJE_N"] == peaje) & (raw_prep["FECHA_D"] == fecha) & raw_prep["DATETIME_CASETA"].notna()].copy()
        clean_day = clean_prep[(clean_prep["PEAJE_N"] == peaje) & (clean_prep["FECHA_D"] == fecha) & clean_prep["DATETIME_CASETA"].notna()].copy()

        windows_df = build_operating_windows(clean_day)
        official_f = filter_to_operating_windows(official_day, windows_df, "VIA_TXT")
        raw_f = filter_to_operating_windows(raw_day, windows_df, "CASETA_TXT")
        clean_f = filter_to_operating_windows(clean_day, windows_df, "CASETA_TXT")

        match_exact, total_official = match_by_plate_time(official_f, clean_f, 0)
        match_30, _ = match_by_plate_time(official_f, clean_f, 30)
        match_60, _ = match_by_plate_time(official_f, clean_f, 60)
        match_300, _ = match_by_plate_time(official_f, clean_f, 300)

        sheet_key = f"{peaje}_{fecha.strftime('%Y%m%d')}"
        hourly_tables[sheet_key] = build_hourly_compare(official_f, clean_f, "base_limpia")
        sentido_tables[sheet_key] = build_sentido_compare(official_f, clean_f, "base_limpia")
        via_tables[sheet_key] = build_via_compare(official_f, clean_f, "base_limpia")
        summary_rows.append(
            {
                "sheet_key": sheet_key,
                "archivo_contraste": official_path.name,
                "peaje": peaje,
                "fecha": fecha.strftime("%Y-%m-%d"),
                "casetas_con_ventana": len(windows_df),
                "oficial_total": len(official_f),
                "base_cruda_total": len(raw_f),
                "base_limpia_total": len(clean_f),
                "cobertura_cruda_pct": safe_pct(len(raw_f), len(official_f)),
                "cobertura_limpia_pct": safe_pct(len(clean_f), len(official_f)),
                "match_exacto_pct": safe_pct(match_exact, total_official),
                "match_30s_pct": safe_pct(match_30, total_official),
                "match_60s_pct": safe_pct(match_60, total_official),
                "match_300s_pct": safe_pct(match_300, total_official),
                "mae_horaria_limpia": round(hourly_tables[sheet_key]["brecha_abs"].abs().mean(), 2),
                "max_brecha_horaria_limpia": int(hourly_tables[sheet_key]["brecha_abs"].abs().max()),
            }
        )

    summary_df = pd.DataFrame(summary_rows)
    time_tables = {
        "tiempos_peaje": summarize_time_metrics(df_resultados, ["PEAJE"]),
        "tiempos_peaje_sentido": summarize_time_metrics(df_resultados, ["PEAJE", "SENTIDO"]),
        "tiempos_caseta": summarize_time_metrics(df_resultados, ["PEAJE", "CASETA", "SENTIDO"]),
    }
    return summary_df, time_tables, hourly_tables, sentido_tables, via_tables


def main() -> None:
    export_tables, raw_df, df_resultados = regenerate_clean_base()
    summary_df, time_tables, hourly_tables, sentido_tables, via_tables = build_contrast_outputs(export_tables, raw_df, df_resultados)
    findings = build_findings(summary_df, hourly_tables, sentido_tables)
    build_word_report(summary_df, findings, hourly_tables, time_tables)
    build_excel_report(summary_df, findings, time_tables, hourly_tables, sentido_tables, via_tables)

    print(f"CLEAN_EXPORT={OUTPUT_CLEAN_PATH}")
    print(f"WORD_EXPORT={OUTPUT_WORD_PATH}")
    print(f"EXCEL_EXPORT={OUTPUT_EXCEL_PATH}")
    print(summary_df.to_string(index=False))


if __name__ == "__main__":
    main()